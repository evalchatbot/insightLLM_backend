from docx import Document
import os

def create_refined_rubric():
    doc = Document()
    doc.add_heading('REFINED RUBRIC WITH SELF-EXPLANATORY POINTS', 0)

    points = [
        ("1. INTRODUCTION QUALITY", "The introduction should directly respond to the question instead of offering unrelated background information. It must demonstrate that the student has correctly understood what is being asked and must define the scope of the answer by identifying the key themes, time period, or analytical angle. A good introduction also presents a clear thesis or central argument and briefly outlines how the rest of the answer will be structured."),
        ("2. HEADINGS & SUBHEADINGS", "Each heading and subheading should be self-explanatory and should immediately tell the examiner what point is being discussed. The headings must stay directly relevant to the question and should reflect the exact themes or components mentioned in the question statement. They should guide the reader by showing how each section contributes to answering the question rather than acting as generic slogans or decorative titles."),
        ("3. ARGUMENTATION QUALITY", "The arguments inside each paragraph must be meaningful, logically developed, and directly linked to the main question. A strong paragraph should present a clear claim, explain it, provide evidence or examples, and show cause-effect relationships. Weak arguments include simple narration of events, emotional opinions, or statements that are not tied back to the central thesis. Critical thinking should be visible through comparisons, explanations of why something happened, and how it shaped outcomes."),
        ("4. FACTUAL ACCURACY", "All factual information—including dates, statistics, examples, treaties, policies, quotations, and references—must be correct, updated, and verifiable. The answer should avoid outdated numbers, exaggerated claims, and misattributions. Proper use of recent data from reliable sources (e.g., Economic Survey, UN, SIPRI, IMF reports) demonstrates credibility. Incorrect or outdated facts weaken the analytical strength of the answer."),
        ("5. GRAMMAR & LANGUAGE", "The language used should be grammatically correct, clear, and suitable for formal academic writing. Sentences must be structured properly, spellings should be accurate, and punctuation should be used consistently. The answer should avoid informal expressions, conversational tone, and slang. The writing must enhance clarity rather than distract from the content."),
        ("6. PRESENTATION QUALITY", "The answer should be presented in a readable and organized manner. Handwriting must be legible, and paragraphs should be spaced clearly. Maps, diagrams, charts, and flowcharts should only be included when they directly support the argument and not simply to fill space. Each visual element must be neat, properly labeled, and contextually meaningful."),
        ("7. CONTEMPORARY RELEVANCE", "The answer should connect the discussion to recent developments and contemporary examples rather than relying entirely on old events or outdated references. Using fresh data, recent policies, current geopolitical shifts, or modern scholarly views shows that the student understands the present relevance of the topic. Answers that ignore recent changes appear incomplete and outdated."),
        ("8. LENGTH & COMPLETENESS", "The overall length of the answer should match the marks allocated. For a 20-mark CSS question, the ideal length is generally between 08 to 10 pages. The response must cover all the key demands of the question. Answers that are too short often lack depth, while excessively long answers may contain irrelevant material and lose focus. Completeness means that no part of the question has been left unaddressed."),
        ("9. REPETITIVENESS", "The answer should avoid repeating the same ideas, examples, or arguments in multiple paragraphs. Each paragraph must contribute a new point or add additional depth to the discussion. Rephrasing the same concept with different words or repeating entire themes shows poor structure and wastes space. An excellent answer progresses logically without redundancy.")
    ]

    for title, content in points:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        doc.add_paragraph('_' * 40)

    # Save to backend/ocr/REFINED RUBRIC.docx
    output_path = os.path.join("d:/css_proj/insightLLM_backend/backend/ocr", "REFINED RUBRIC.docx")
    doc.save(output_path)
    print(f"Created {output_path}")

if __name__ == "__main__":
    create_refined_rubric()
