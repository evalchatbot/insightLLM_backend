#!/usr/bin/env python3
"""Parse a complete rubric file to see full structure with indicators."""

from docx import Document
from pathlib import Path
import json

def parse_complete_rubric(docx_path: Path):
    """Parse and display complete rubric structure."""
    doc = Document(docx_path)

    print(f"\n{'='*80}")
    print(f"COMPLETE RUBRIC: {docx_path.stem}")
    print(f"{'='*80}\n")

    # Print all paragraphs
    print("--- ALL CONTENT ---\n")
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            print(f"[{i:03d}] [{style:20s}] {para.text}")

    # Print all tables
    print(f"\n\n{'='*80}")
    print("--- ALL TABLES ---")
    print(f"{'='*80}\n")

    for table_idx, table in enumerate(doc.tables, 1):
        print(f"\nTABLE {table_idx}:")
        print(f"{'─'*80}")

        for row_idx, row in enumerate(table.rows, 1):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"Row {row_idx:02d}: {cells}")

    print(f"\n\n{'='*80}")
    print("END OF RUBRIC")
    print(f"{'='*80}\n")

if __name__ == "__main__":
    # Parse Political Science rubric first
    rubric_file = Path(__file__).parent / "Rubrics" / "Political Science Rubric" / "Political Science.docx"

    if rubric_file.exists():
        parse_complete_rubric(rubric_file)
    else:
        print(f"File not found: {rubric_file}")
