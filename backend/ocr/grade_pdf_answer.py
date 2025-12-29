# grade_pdf_answer.py
#
# Simplified pipeline:
#   1) Google Vision OCR for text + bounding boxes.
#   2) Grok Prompt 1: Headings & structure detection (with PDF images).
#   3) Grok Prompt 2: Subject-wise marking (with PDF images + subject rubric DOCX).
#   4) Grok Prompt 3: Refined rubric annotations (with PDF images + refined rubric DOCX).
#   5) Render subject-wise report pages.
#   6) Render refined-rubric summary page.
#   7) Annotate answer pages according to simplified rules.
#   8) Merge all pages into final PDF.

import argparse
import base64
import datetime
import gc
import io
import json
import os
import re
import sys
import tempfile
import uuid
import shutil
import time
import traceback
import random
from functools import lru_cache
from typing import Any, Dict, List, Tuple, Optional, Set
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError

import requests
from dotenv import load_dotenv
from google.api_core.client_options import ClientOptions
from google.cloud import vision
import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import cv2
import numpy as np
from PyPDF2 import PdfWriter, PdfReader

try:
    from backend.ocr.annotate_pdf_with_rubric import annotate_pdf_answer_pages
    from backend.ocr.progress_tracker import OCRProgressTracker
except ImportError:
    from annotate_pdf_with_rubric import annotate_pdf_answer_pages
    from progress_tracker import OCRProgressTracker


# -----------------------------
# UTILS & ENV
# -----------------------------




def debug_dump_sections(
    sections: List[Dict[str, Any]],
    output_path: str = "debug_sections.json",
) -> None:
    """
    Save detected headings/sections to a JSON file and print a clean summary.

    Each section includes:
      - title
      - level (1 = main heading, 2 = subheading)
      - page_numbers
      - first 200 chars of content (for quick checking)
    """
    light_sections = []
    for idx, sec in enumerate(sections):
        light_sections.append(
            {
                "index": idx,
                "title": sec.get("title"),
                "exact_ocr_heading": sec.get("exact_ocr_heading"),
                "level": sec.get("level"),
                "page_numbers": sec.get("page_numbers"),
                "content_preview": (sec.get("content") or "")[:200],
                "comment": sec.get("comment"),  # Add comment field
            }
        )

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(light_sections, f, ensure_ascii=False, indent=2)

    print("\n==== DETECTED HEADINGS / SECTIONS (from Grok) ====")
    for sec in light_sections:
        comment_display = f" | Comment: {sec['comment']}" if sec.get('comment') else ""
        print(
            f"[{sec['index']}] "
            f"Title: {sec['title']!r} | "
            f"exact_ocr_heading: {sec['exact_ocr_heading']!r} | "
            f"Level: {sec['level']} | "
            f"Pages: {sec['page_numbers']}"
            f"{comment_display}"
        )
    print(f"Saved detailed section info to {output_path}")
    print("=================================================\n")


def _append_log(log_path: Optional[str], level: str, message: str) -> None:
    if not log_path:
        return
    try:
        timestamp = datetime.datetime.utcnow().isoformat() + "Z"
        line = f"{timestamp} [{level}] {message}\n"
        
        # Determine log directory
        log_dir = os.path.dirname(log_path)
        
        # Check if this is an OCR-related log message
        # OCR logs include: upload, steps, timing reports, completion, OCR events
        is_ocr_log = (
            "upload_start" in message or
            "start pdf=" in message or
            " step=" in message or  # Note: space before step= to avoid false matches
            "TIMING_REPORT" in message or
            ("completed" in message and "request=" in message) or
            "report_generated" in message or
            "ocr_" in message or
            "Step " in message  # Timing report steps (e.g., "Step 1: Convert PDF")
        )
        
        # Write to main log file
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(line)
        
        # Write OCR logs to separate OCR log file
        if is_ocr_log:
            ocr_log_path = os.path.join(log_dir, "ocr_log.txt")
            with open(ocr_log_path, "a", encoding="utf-8") as f:
                f.write(line)
        
        # Write errors to separate error log file
        if level == "ERROR":
            error_log_path = os.path.join(log_dir, "errors_log.txt")
            with open(error_log_path, "a", encoding="utf-8") as f:
                f.write(line)
    except Exception:
        # Never fail the pipeline due to logging issues.
        pass


def _format_time(seconds: float) -> str:
    """Format time in seconds to 'X min Y sec' format."""
    total_seconds = int(seconds)
    minutes = total_seconds // 60
    secs = total_seconds % 60
    if minutes > 0:
        return f"{minutes} min {secs} sec"
    else:
        return f"{secs} sec"






def clean_json_from_llm(text: str) -> str:
    """
    Remove markdown code fences and extract JSON from LLM responses.
    Handles various formats: ```json ... ```, ``` ... ```, or plain JSON.
    """
    text = text.strip()
    
    # Remove markdown code fences
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z0-9]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
        text = text.strip()
    
    # Try to extract JSON object/array if wrapped in text
    # Look for JSON object starting with {
    json_match = re.search(r'(\{[\s\S]*\})', text)
    if json_match:
        text = json_match.group(1)
    
    return text.strip()


def repair_json(text: str, error_pos: Optional[int] = None) -> str:
    """
    Attempt to repair common JSON issues:
    - Trailing commas before closing brackets/braces
    - Control characters
    - Unclosed strings (basic handling)
    """
    # Remove trailing commas before } or ]
    text = re.sub(r',(\s*[}\]])', r'\1', text)
    
    # Remove control characters except newlines and tabs
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t\r')
    
    # If error position is known, try to fix around that area
    if error_pos is not None and error_pos < len(text):
        # Try to add missing comma if we're at a position that might need one
        # This is a heuristic - look for patterns like "value }" or "value ]"
        if error_pos > 0:
            # Check context around error position
            start = max(0, error_pos - 50)
            end = min(len(text), error_pos + 50)
            context = text[start:end]
            
            # Try to fix common patterns: "value }" -> "value, }" (but be careful)
            # This is a last resort and might not always work correctly
            pass  # Keep it simple for now - trailing comma removal is the main fix
    
    return text


# -----------------------------
# JSON SCHEMA VALIDATION
# -----------------------------


def validate_refined_summary(summary_list: List[Dict[str, Any]]) -> bool:
    """Validate refined_rubric_summary schema."""
    REQUIRED_KEYS = ["id", "name", "rating", "comment"]
    VALID_RATINGS = {"weak", "average", "good", "excellent"}
    VALID_IDS = {"argumentation_quality", "presentation",
                 "contemporary_relevance", "length_completeness"}

    for idx, item in enumerate(summary_list):
        # Check required keys
        missing = set(REQUIRED_KEYS) - set(item.keys())
        if missing:
            raise ValueError(f"refined_rubric_summary[{idx}] missing fields: {missing}")

        # Validate rating
        if item["rating"].lower() not in VALID_RATINGS:
            print(f"WARNING: Invalid rating in summary[{idx}]: {item['rating']} (expected: weak/average/good/excellent)")

        # Validate ID (warn but don't fail)
        if item["id"] not in VALID_IDS:
            print(f"WARNING: Unexpected summary ID in summary[{idx}]: {item['id']}")

    return True


def validate_annotation(annotation: Dict[str, Any], idx: int = 0) -> bool:
    """Validate single annotation schema."""
    REQUIRED_KEYS = ["type", "rubric_point", "page",
                     "target_word_or_sentence", "context_before",
                     "context_after", "correction", "comment"]

    missing = set(REQUIRED_KEYS) - set(annotation.keys())
    if missing:
        print(f"WARNING: Annotation[{idx}] missing fields: {missing}")
        return False

    # Validate page number
    if not isinstance(annotation["page"], int) or annotation["page"] < 1:
        print(f"WARNING: Invalid page number in annotation[{idx}]: {annotation['page']}")
        return False

    return True


def validate_input_paths(pdf_path: str, output_json_path: str, output_pdf_path: str) -> bool:
    """Validate all input/output paths before processing."""
    # Check PDF exists and is readable
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    try:
        with open(pdf_path, 'rb') as f:
            # Check file is not empty and starts with PDF header
            header = f.read(4)
            if header != b'%PDF':
                raise ValueError(f"File is not a valid PDF: {pdf_path}")
    except Exception as e:
        raise ValueError(f"Cannot read PDF {pdf_path}: {e}")

    # Check output paths are writable
    for path in [output_json_path, output_pdf_path]:
        try:
            dirname = os.path.dirname(path)
            if dirname and not os.path.exists(dirname):
                os.makedirs(dirname, exist_ok=True)
            # Test write access
            with open(path, 'w') as f:
                f.write("")
            os.remove(path)
        except Exception as e:
            raise ValueError(f"Cannot write to {path}: {e}")

    return True


def load_environment() -> Tuple[str, vision.ImageAnnotatorClient]:
    """
    Load environment variables for Grok and Google Vision.

    .env must contain:
      Grok_API=YOUR_GROK_KEY
      Google_cloud_key=YOUR_GOOGLE_CLOUD_API_KEY
    """
    load_dotenv()
    grok_key = os.getenv("Grok_API")
    google_key = os.getenv("Google_cloud_key")
    missing = []
    if not grok_key:
        missing.append("Grok_API")
    if not google_key:
        missing.append("Google_cloud_key")
    if missing:
        raise EnvironmentError(
            f"Missing environment variable(s): {', '.join(missing)}. "
            "Please set them in your .env file."
        )

    # Validate API key formats
    if len(grok_key) < 20:
        raise ValueError(
            f"Invalid Grok_API key format: key is too short ({len(grok_key)} characters). "
            "Expected at least 20 characters."
        )
    if len(google_key) < 20:
        raise ValueError(
            f"Invalid Google_cloud_key format: key is too short ({len(google_key)} characters). "
            "Expected at least 20 characters."
        )

    client_options = ClientOptions(api_key=google_key)
    vision_client = vision.ImageAnnotatorClient(client_options=client_options)
    return grok_key, vision_client


# -----------------------------
# PDF → PAGE IMAGES (for Grok)
# -----------------------------


def pdf_to_page_images_for_grok(
    pdf_path: str,
    max_pages: int = 20,
    max_dim: int = 400,
    base64_cap: int = 12000,
    output_dir: str = "grok_images",
) -> List[Dict[str, Any]]:
    """
    Convert up to `max_pages` pages of the PDF into resized JPEG images (for Grok).
    Optimized to reduce token usage: lower DPI, smaller size, JPEG compression.
    Saves images to output_dir for inspection/debugging.
    Returns a list: [{"page": 1, "image_base64": "...", "file_path": "...", "truncated": bool}, ...]
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    doc = fitz.open(pdf_path)
    try:
        page_images: List[Dict[str, Any]] = []
        
        # Process pages one at a time to avoid memory accumulation
        for idx, page in enumerate(doc):
            if idx >= max_pages:
                break
            
            # Load page image
            pix = None
            pil_img = None
            resized = None
            rgb_img = None
            
            try:
                # Use 200 DPI to match OCR processing (ensures coordinate alignment)
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                pil_img = Image.open(io.BytesIO(img_bytes))
                
                # Process immediately - don't accumulate in memory
                # Reduced max dimension from 1200 to 800 (44% fewer pixels)
                resized = pil_img.copy()
                resized.thumbnail((max_dim, max_dim))

                # Convert to RGB if necessary (JPEG doesn't support transparency)
                if resized.mode in ('RGBA', 'LA', 'P'):
                    rgb_img = Image.new('RGB', resized.size, (255, 255, 255))
                    if resized.mode == 'P':
                        resized = resized.convert('RGBA')
                    rgb_img.paste(resized, mask=resized.split()[-1] if resized.mode in ('RGBA', 'LA') else None)
                    resized = rgb_img
                elif resized.mode != 'RGB':
                    resized = resized.convert('RGB')

                buffer = io.BytesIO()
                # Changed from PNG to JPEG with 60% quality for reduced payload size
                resized.save(buffer, format="JPEG", quality=60, optimize=True)

                # Save image to disk
                file_path = os.path.join(output_dir, f"page_{idx + 1:03d}.jpg")
                resized.save(file_path, format="JPEG", quality=60, optimize=True)
                
                encoded = base64.b64encode(buffer.getvalue()).decode("utf-8")
                truncated = False
                if len(encoded) > base64_cap:
                    encoded = encoded[:base64_cap]
                    truncated = True
                page_images.append(
                    {
                        "page": idx + 1,
                        "image_base64": encoded,
                        "file_path": file_path,
                        "truncated": truncated
                    }
                )
            finally:
                # Explicitly delete and cleanup after each page
                if rgb_img is not None:
                    del rgb_img
                if resized is not None:
                    del resized
                if pil_img is not None:
                    pil_img.close()
                    del pil_img
                if pix is not None:
                    del pix
                # Force garbage collection after each page
                gc.collect()
        
        print(f"Saved {len(page_images)} page images to '{output_dir}/'")
        return page_images
    finally:
        doc.close()  # Always close the document to release file handle


def get_report_page_size(
    pdf_path: str,
    dpi: int = 200,
    margin_ratio: float = 0.40,
    min_height: int = 3500,
    fallback: Tuple[int, int] = (2977, 4211),
) -> Tuple[int, int]:
    """
    Match report page size to annotated answer pages:
    annotated width = orig_w + 2 * (margin_ratio * orig_w), height = orig_h.
    
    This ensures that report pages have the same dimensions as the annotated
    answer pages, creating a consistent document layout.
    
    Args:
        pdf_path: Path to the PDF file
        dpi: DPI for page size calculation (default: 200, matches OCR processing)
        margin_ratio: Ratio of margin to page width (default: 0.40 = 40%)
        min_height: Minimum page height in pixels (default: 3500)
        fallback: Fallback page size if calculation fails (default: A4 at 200 DPI)
    
    Returns:
        Tuple of (width, height) in pixels at the specified DPI
    
    Example:
        >>> size = get_report_page_size("answer.pdf")
        >>> print(size)  # (4167, 4211) - width with margins, height with minimum
    """
    doc = fitz.open(pdf_path)
    try:
        if doc.page_count == 0:
            return fallback
        pix = doc[0].get_pixmap(dpi=dpi)
        orig_w, orig_h = pix.width, pix.height
        margin = int(orig_w * margin_ratio)
        return (orig_w + 2 * margin, max(orig_h, min_height))
    except Exception:
        return fallback
    finally:
        doc.close()


# -----------------------------
# OCR WITH GOOGLE VISION
# -----------------------------


def _bbox_to_tuples(bbox) -> List[Tuple[int, int]]:
    return [(v.x, v.y) for v in bbox.vertices]


def _paragraph_text(paragraph) -> str:
    words = []
    for word in paragraph.words:
        symbols = "".join(symbol.text for symbol in word.symbols)
        words.append(symbols)
    return " ".join(words).strip()


def _is_noise_text(text: str, bbox: List[Tuple[int, int]], page_w: int, page_h: int) -> bool:
    """
    Filter out background noise from OCR results.
    Returns True if the text is likely noise.
    """
    if not text or not bbox:
        return True

    # Filter very short text (1-2 chars) that's likely noise
    if len(text.strip()) <= 2:
        return True

    # Calculate bbox dimensions
    xs = [p[0] for p in bbox]
    ys = [p[1] for p in bbox]
    if not xs or not ys:
        return True

    width = max(xs) - min(xs)
    height = max(ys) - min(ys)

    # Filter extremely small text (likely artifacts)
    if width < 10 or height < 10:
        return True

    # Filter text at extreme edges (often page numbers or noise)
    center_x = (min(xs) + max(xs)) / 2
    center_y = (min(ys) + max(ys)) / 2

    margin = 30  # pixels
    if center_x < margin or center_x > page_w - margin:
        if center_y < margin or center_y > page_h - margin:
            return True

    return False


def _is_retryable_error(
    error: Exception,
    response: Optional[vision.AnnotateImageResponse] = None,
) -> Tuple[bool, str]:
    """
    Determine if an error should be retried based on error type and message.
    
    Args:
        error: The exception that was raised
        response: Optional Vision API response object (for checking response.error)
    
    Returns:
        Tuple of (is_retryable: bool, error_category: str)
        error_category: One of: 'network_error', 'rate_limit', 'server_error', 
                       'timeout', 'auth_error', 'invalid_input', 'not_found', 'unknown'
    """
    error_type = type(error).__name__
    error_msg = str(error).lower()
    
    # Check Google Vision API response errors first (if available)
    if response and hasattr(response, 'error') and response.error.message:
        api_error_msg = response.error.message.lower()
        
        # Rate limit errors
        if any(keyword in api_error_msg for keyword in [
            'resource_exhausted', 'rate limit', 'quota', '429', 'rate limit'
        ]):
            return True, 'rate_limit'
        
        # Server errors
        if any(keyword in api_error_msg for keyword in [
            'unavailable', 'deadline_exceeded', 'internal error', '500', '502', '503', '504'
        ]):
            return True, 'server_error'
        
        # Non-retryable API errors
        if any(keyword in api_error_msg for keyword in [
            'permission_denied', 'invalid_argument', 'invalid_image', 
            'not_found', '401', '403', '400', '404', '422'
        ]):
            if 'permission_denied' in api_error_msg or '401' in api_error_msg or '403' in api_error_msg:
                return False, 'auth_error'
            elif 'invalid' in api_error_msg or '400' in api_error_msg or '422' in api_error_msg:
                return False, 'invalid_input'
            elif 'not_found' in api_error_msg or '404' in api_error_msg:
                return False, 'not_found'
    
    # Check exception types
    # Network/Connection errors (retryable)
    if error_type in ['ConnectionError', 'ConnectionResetError', 'ConnectionAbortedError']:
        return True, 'network_error'
    
    if 'connection' in error_msg or 'network' in error_msg or 'dns' in error_msg or 'socket' in error_msg:
        return True, 'network_error'
    
    # Timeout errors (retryable, but conditional)
    if error_type == 'TimeoutError' or 'timeout' in error_msg or 'deadline' in error_msg:
        return True, 'timeout'
    
    # Rate limit errors (retryable)
    if '429' in error_msg or 'rate limit' in error_msg or 'quota' in error_msg or 'resource_exhausted' in error_msg:
        return True, 'rate_limit'
    
    # Server errors (retryable)
    if any(code in error_msg for code in ['500', '502', '503', '504']):
        return True, 'server_error'
    if 'unavailable' in error_msg or 'internal error' in error_msg or 'gateway' in error_msg:
        return True, 'server_error'
    
    # Authentication/Permission errors (non-retryable)
    if error_type in ['PermissionError']:
        return False, 'auth_error'
    if any(keyword in error_msg for keyword in [
        'auth', 'permission', 'unauthorized', 'forbidden', '401', '403'
    ]):
        return False, 'auth_error'
    
    # Invalid request errors (non-retryable)
    if error_type == 'ValueError':
        return False, 'invalid_input'
    if any(keyword in error_msg for keyword in [
        'invalid', 'bad request', '400', '422', 'invalid_argument', 'invalid_image'
    ]):
        return False, 'invalid_input'
    
    # Not found errors (non-retryable)
    if '404' in error_msg or 'not found' in error_msg:
        return False, 'not_found'
    
    # Format/corruption errors (non-retryable)
    if any(keyword in error_msg for keyword in [
        'format', 'corrupt', 'unsupported', 'too large', 'size limit'
    ]):
        return False, 'invalid_input'
    
    # Default: For unknown errors, be conservative and don't retry
    # This can be adjusted based on observed error patterns
    return False, 'unknown'


def _calculate_backoff_delay(
    attempt_number: int,
    base_delay: float = 1.0,
    max_delay: float = 60.0,
    jitter_range: float = 0.2,
    is_rate_limit: bool = False,
    rate_limit_base_delay: float = 5.0,
    rate_limit_max_delay: float = 300.0,
    retry_after: Optional[float] = None,
) -> float:
    """
    Calculate exponential backoff delay with jitter.
    
    Formula: delay = base_delay * (2 ^ (attempt_number - 1)) + jitter
    - Exponential growth per attempt
    - Capped at max_delay
    - Jitter added to prevent thundering herd
    
    Args:
        attempt_number: Current attempt number (1-indexed, attempt 1 = no wait)
        base_delay: Base delay in seconds (default: 1.0)
        max_delay: Maximum delay in seconds (default: 60.0)
        jitter_range: Jitter range (0.0-1.0) for randomization (default: 0.2 = ±20%)
        is_rate_limit: If True, use rate limit specific delays
        rate_limit_base_delay: Base delay for rate limits (default: 5.0)
        rate_limit_max_delay: Max delay for rate limits (default: 300.0)
        retry_after: Optional Retry-After header value in seconds (takes precedence)
    
    Returns:
        Delay in seconds (float)
    
    Examples:
        # Standard backoff (attempt 2, base=1.0, max=60.0, jitter=0.2)
        # delay = 1.0 * 2^1 = 2.0s, jitter = ±0.4s → 1.6-2.4s
        
        # Rate limit backoff (attempt 3, base=5.0, max=300.0, jitter=0.2)
        # delay = 5.0 * 2^2 = 20.0s, jitter = ±4.0s → 16.0-24.0s
        
        # With Retry-After header
        # delay = retry_after (e.g., 30.0s) - no jitter applied
    """
    # If Retry-After header is provided, use it directly (no jitter)
    if retry_after is not None and retry_after > 0:
        return float(retry_after)
    
    # Use rate limit parameters if this is a rate limit error
    if is_rate_limit:
        effective_base = rate_limit_base_delay
        effective_max = rate_limit_max_delay
    else:
        effective_base = base_delay
        effective_max = max_delay
    
    # Calculate exponential delay: base * (2 ^ (attempt - 1))
    # Attempt 1 = no wait (2^0 = 1, but we don't wait before first attempt)
    # Attempt 2 = base * 2^1 = base * 2
    # Attempt 3 = base * 2^2 = base * 4
    # etc.
    if attempt_number <= 1:
        # First attempt has no backoff (immediate)
        return 0.0
    
    exponential_delay = effective_base * (2 ** (attempt_number - 1))
    
    # Cap at maximum delay
    capped_delay = min(exponential_delay, effective_max)
    
    # Add jitter: random variation of ±jitter_range percentage
    # Example: jitter_range=0.2 means ±20% variation
    # jitter = random.uniform(-0.2, 0.2) * capped_delay
    jitter = random.uniform(-jitter_range, jitter_range) * capped_delay
    final_delay = capped_delay + jitter
    
    # Ensure delay is non-negative
    return max(0.0, final_delay)


def _check_retry_budget(
    elapsed_time: float,
    overall_timeout: Optional[float],
    backoff_delay: float,
    estimated_attempt_time: float,
    safety_margin: float = 5.0,
) -> bool:
    """
    Check if retry budget allows another retry attempt.
    
    Determines if there's enough time remaining in the overall timeout budget
    to perform another retry attempt (including backoff delay and estimated attempt time).
    
    Args:
        elapsed_time: Time already spent on processing (seconds)
        overall_timeout: Overall timeout budget (None = no limit)
        backoff_delay: Calculated backoff delay before retry (seconds)
        estimated_attempt_time: Estimated time for the retry attempt (seconds)
        safety_margin: Safety margin to avoid cutting it too close (default: 5.0 seconds)
    
    Returns:
        True if retry budget allows another attempt, False otherwise
    
    Examples:
        # Overall timeout = 600s, elapsed = 580s, backoff = 2s, attempt = 10s
        # Remaining = 20s, needed = 12s → True (within budget)
        
        # Overall timeout = 600s, elapsed = 590s, backoff = 5s, attempt = 10s
        # Remaining = 10s, needed = 15s → False (exceeds budget)
    """
    # If no overall timeout, always allow retry
    if overall_timeout is None:
        return True
    
    # Calculate remaining budget
    remaining_budget = overall_timeout - elapsed_time
    
    # Calculate total retry cost (backoff + attempt + safety margin)
    retry_cost = backoff_delay + estimated_attempt_time + safety_margin
    
    # Check if retry would exceed budget
    return retry_cost <= remaining_budget


def _estimate_attempt_time(
    per_page_timeout: float,
    previous_attempt_duration: Optional[float] = None,
) -> float:
    """
    Estimate time for a retry attempt.
    
    Uses conservative estimate based on per-page timeout or previous attempt duration.
    
    Args:
        per_page_timeout: Per-page timeout (used as max estimate)
        previous_attempt_duration: Duration of previous attempt (if available)
    
    Returns:
        Estimated attempt time in seconds
    """
    # If we have previous attempt duration, use it as estimate (with small buffer)
    if previous_attempt_duration is not None:
        # Use previous duration + 10% buffer as estimate
        return previous_attempt_duration * 1.1
    
    # Otherwise, use per-page timeout as conservative estimate
    # This assumes worst case: attempt takes full timeout
    return per_page_timeout


def _call_vision_with_retry(
    vision_client: vision.ImageAnnotatorClient,
    vision_image: vision.Image,
    timeout_seconds: float,
    page_number: int,
    max_retries: int = 3,
    retry_base_delay: float = 1.0,
    retry_max_delay: float = 60.0,
    retry_jitter_range: float = 0.2,
    rate_limit_base_delay: float = 5.0,
    rate_limit_max_delay: float = 300.0,
    overall_timeout: Optional[float] = None,
    overall_start_time: Optional[float] = None,
    log_path: Optional[str] = None,
    request_id: Optional[str] = None,
    retry_stats: Optional[Dict[str, Any]] = None,
) -> vision.AnnotateImageResponse:
    """
    Call Google Vision API with timeout protection and retry logic.
    
    Wraps `_call_vision_with_timeout()` with retry logic, exponential backoff,
    and budget management. Retries transient failures while failing fast on
    permanent errors.
    
    Args:
        vision_client: Google Vision client
        vision_image: Image to process
        timeout_seconds: Maximum time per attempt (seconds)
        page_number: Page number for error messages
        max_retries: Maximum retry attempts (default: 3)
        retry_base_delay: Base delay for exponential backoff (default: 1.0)
        retry_max_delay: Maximum delay between retries (default: 60.0)
        retry_jitter_range: Jitter range for randomization (default: 0.2)
        rate_limit_base_delay: Base delay for rate limit errors (default: 5.0)
        rate_limit_max_delay: Max delay for rate limit errors (default: 300.0)
        overall_timeout: Overall timeout budget (None = no limit)
        overall_start_time: Start time for overall timeout tracking (None = current time)
        log_path: Optional path to log file
        request_id: Optional request ID for logging
    
    Returns:
        Vision API response
    
    Raises:
        TimeoutError: If timeout occurs and retries exhausted or budget exceeded
        RuntimeError: If API call fails with non-retryable error or all retries exhausted
    """
    if overall_start_time is None:
        overall_start_time = time.perf_counter()
    
    last_error: Optional[Exception] = None
    last_response: Optional[vision.AnnotateImageResponse] = None
    previous_attempt_duration: Optional[float] = None
    
    # Initialize retry statistics if not provided
    if retry_stats is None:
        retry_stats = {
            "total_attempts": 0,
            "successful_retries": 0,
            "exhausted_retries": 0,
            "rate_limit_events": 0,
            "non_retryable_errors": 0,
            "budget_exceeded": 0,
            "retry_attempts_by_category": {},
        }
    
    for attempt in range(1, max_retries + 1):
        attempt_start_time = time.perf_counter()
        retry_stats["total_attempts"] += 1
        
        try:
            # Make API call with timeout
            response = _call_vision_with_timeout(
                vision_client=vision_client,
                vision_image=vision_image,
                timeout_seconds=timeout_seconds,
                page_number=page_number,
            )
            
            # Store response for potential error classification
            last_response = response
            
            # Success on first attempt
            if attempt == 1:
                return response
            
            # Success after retry
            attempt_duration = time.perf_counter() - attempt_start_time
            retry_stats["successful_retries"] += 1
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_retry_success page={page_number} "
                f"attempt={attempt}/{max_retries} total_attempts={attempt} "
                f"duration_ms={int(attempt_duration * 1000)}",
            )
            # Return response with retry stats (for potential future use)
            return response
            
        except Exception as e:
            last_error = e
            attempt_duration = time.perf_counter() - attempt_start_time
            previous_attempt_duration = attempt_duration
            
            # Note: last_response may not be available if exception occurred before API call
            # Error classification will work with exception type and message patterns
            
            # Classify error
            is_retryable, error_category = _is_retryable_error(e, last_response)
            
            # Non-retryable error: fail fast
            if not is_retryable:
                retry_stats["non_retryable_errors"] += 1
                _append_log(
                    log_path,
                    "ERROR",
                    f"request={request_id} ocr_non_retryable page={page_number} "
                    f"attempt={attempt} error_category={error_category} error={str(e)}",
                )
                raise RuntimeError(
                    f"OCR failed on page {page_number} (non-retryable {error_category}): {str(e)}"
                ) from e
            
            # Check if we have retries remaining
            if attempt >= max_retries:
                # All retries exhausted
                retry_stats["exhausted_retries"] += 1
                # Track retry attempts by category
                if error_category not in retry_stats["retry_attempts_by_category"]:
                    retry_stats["retry_attempts_by_category"][error_category] = 0
                retry_stats["retry_attempts_by_category"][error_category] += 1
                _append_log(
                    log_path,
                    "ERROR",
                    f"request={request_id} ocr_retry_exhausted page={page_number} "
                    f"attempts={max_retries} error_category={error_category} error={str(e)}",
                )
                raise RuntimeError(
                    f"OCR failed on page {page_number} after {max_retries} attempts "
                    f"({error_category}): {str(e)}"
                ) from e
            
            # Track retry attempts by category
            if error_category not in retry_stats["retry_attempts_by_category"]:
                retry_stats["retry_attempts_by_category"][error_category] = 0
            retry_stats["retry_attempts_by_category"][error_category] += 1
            
            # Calculate backoff delay
            is_rate_limit = (error_category == 'rate_limit')
            if is_rate_limit:
                retry_stats["rate_limit_events"] += 1
            backoff_delay = _calculate_backoff_delay(
                attempt_number=attempt + 1,  # Next attempt number
                base_delay=retry_base_delay,
                max_delay=retry_max_delay,
                jitter_range=retry_jitter_range,
                is_rate_limit=is_rate_limit,
                rate_limit_base_delay=rate_limit_base_delay,
                rate_limit_max_delay=rate_limit_max_delay,
                retry_after=None,  # TODO: Extract from response headers if available
            )
            
            # Estimate attempt time for budget check
            estimated_attempt_time = _estimate_attempt_time(
                per_page_timeout=timeout_seconds,
                previous_attempt_duration=previous_attempt_duration,
            )
            
            # Check retry budget
            elapsed_time = time.perf_counter() - overall_start_time
            if not _check_retry_budget(
                elapsed_time=elapsed_time,
                overall_timeout=overall_timeout,
                backoff_delay=backoff_delay,
                estimated_attempt_time=estimated_attempt_time,
            ):
                # Budget exceeded, stop retrying
                retry_stats["budget_exceeded"] += 1
                _append_log(
                    log_path,
                    "WARNING",
                    f"request={request_id} ocr_retry_budget_exceeded page={page_number} "
                    f"attempt={attempt} elapsed_s={elapsed_time:.1f} "
                    f"overall_timeout_s={overall_timeout} backoff_s={backoff_delay:.1f} "
                    f"estimated_s={estimated_attempt_time:.1f}",
                )
                raise TimeoutError(
                    f"OCR retry budget exceeded on page {page_number} at attempt {attempt}: "
                    f"elapsed {elapsed_time:.1f}s, would need {backoff_delay + estimated_attempt_time:.1f}s"
                )
            
            # Log retry attempt
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_retry_attempt page={page_number} "
                f"attempt={attempt + 1}/{max_retries} error_category={error_category} "
                f"wait_s={backoff_delay:.2f} previous_duration_ms={int(previous_attempt_duration * 1000)}",
            )
            
            # Special logging for rate limits
            if is_rate_limit:
                _append_log(
                    log_path,
                    "WARNING",
                    f"request={request_id} ocr_rate_limit page={page_number} "
                    f"attempt={attempt + 1} wait_s={backoff_delay:.2f}",
                )
            
            # Wait before retry
            time.sleep(backoff_delay)
            
            # Continue to next retry attempt
            continue
    
    # Should never reach here, but handle just in case
    if last_error:
        raise RuntimeError(
            f"OCR failed on page {page_number} after {max_retries} attempts: {str(last_error)}"
        ) from last_error
    
    raise RuntimeError(f"OCR failed on page {page_number}: unexpected error")


def _call_vision_with_timeout(
    vision_client: vision.ImageAnnotatorClient,
    vision_image: vision.Image,
    timeout_seconds: float,
    page_number: int,
) -> vision.AnnotateImageResponse:
    """
    Call Google Vision API with timeout protection.
    
    Args:
        vision_client: Google Vision client
        vision_image: Image to process
        timeout_seconds: Maximum time to wait for response
        page_number: Page number for error messages
    
    Returns:
        Vision API response
    
    Raises:
        TimeoutError: If the API call exceeds timeout_seconds
        RuntimeError: If the API call fails with an error
    """
    def _make_api_call():
        """Make the actual API call in a separate thread."""
        return vision_client.document_text_detection(image=vision_image)
    
    # Use ThreadPoolExecutor to enforce timeout
    with ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(_make_api_call)
        try:
            response = future.result(timeout=timeout_seconds)
            
            # Check for API errors
            if response.error.message:
                raise RuntimeError(
                    f"OCR failed on page {page_number}: {response.error.message}"
                )
            
            return response
            
        except FutureTimeoutError:
            # Cancel the future if possible
            future.cancel()
            raise TimeoutError(
                f"OCR timeout on page {page_number}: "
                f"exceeded {timeout_seconds} seconds"
            )
        except Exception as e:
            # Re-raise other exceptions with context
            if isinstance(e, (TimeoutError, RuntimeError)):
                raise
            raise RuntimeError(
                f"OCR error on page {page_number}: {str(e)}"
            ) from e


def _optimize_image_for_ocr(
    img: Image.Image,
    max_dimension: int,
    min_dimension_for_optimization: int,
    enabled: bool = True,
) -> Tuple[Image.Image, float, float]:
    """
    Optimize image for OCR by downscaling if it exceeds maximum dimensions.
    Preserves aspect ratio and returns scale factors for bounding box adjustment.
    
    Args:
        img: PIL Image to optimize
        max_dimension: Maximum width or height (downscale if larger)
        min_dimension_for_optimization: Only optimize if dimension exceeds this
        enabled: Whether optimization is enabled
    
    Returns:
        Tuple of (optimized_image, scale_x, scale_y)
        - optimized_image: Optimized PIL Image (or original if not optimized)
        - scale_x: X-axis scale factor (original_width / optimized_width, 1.0 if not scaled)
        - scale_y: Y-axis scale factor (original_height / optimized_height, 1.0 if not scaled)
    """
    if not enabled:
        return img, 1.0, 1.0
    
    original_w, original_h = img.size
    max_dim = max(original_w, original_h)
    
    # Only optimize if image exceeds minimum dimension threshold
    if max_dim <= min_dimension_for_optimization:
        return img, 1.0, 1.0
    
    # Calculate new dimensions preserving aspect ratio
    if original_w > original_h:
        # Landscape: limit width
        if original_w > max_dimension:
            new_w = max_dimension
            new_h = int(original_h * (max_dimension / original_w))
        else:
            return img, 1.0, 1.0
    else:
        # Portrait or square: limit height
        if original_h > max_dimension:
            new_h = max_dimension
            new_w = int(original_w * (max_dimension / original_h))
        else:
            return img, 1.0, 1.0
    
    # Downscale using high-quality resampling
    optimized_img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
    
    # Calculate scale factors for bounding box adjustment
    scale_x = original_w / new_w
    scale_y = original_h / new_h
    
    return optimized_img, scale_x, scale_y


def _process_single_page_ocr(
    vision_client: vision.ImageAnnotatorClient,
    img: Image.Image,
    page_num: int,
    per_page_timeout: float,
    overall_timeout: Optional[float],
    overall_start_time: float,
    max_retries: int,
    retry_base_delay: float,
    retry_max_delay: float,
    retry_jitter_range: float,
    rate_limit_base_delay: float,
    rate_limit_max_delay: float,
    log_path: Optional[str],
    request_id: Optional[str],
    image_optimization_enabled: bool = True,
    image_max_dimension: int = 2048,
    image_min_dimension_for_optimization: int = 1500,
) -> Tuple[int, Dict[str, Any], Dict[str, Any], Optional[str]]:
    """
    Process a single page OCR. This function is designed to be called in parallel.
    
    Returns:
        Tuple of (page_number, page_output_dict, retry_stats_dict, full_text, error_message)
        - page_output_dict: Page OCR data or error info
        - retry_stats_dict: Retry statistics for this page
        - full_text: Extracted text from page (empty string if error)
        - error_message: None if success, error message if failed
    """
    page_retry_stats = {
        "total_attempts": 0,
        "successful_retries": 0,
        "exhausted_retries": 0,
        "rate_limit_events": 0,
        "non_retryable_errors": 0,
        "budget_exceeded": 0,
        "retry_attempts_by_category": {},
    }
    
    try:
        # Check overall timeout before processing
        if overall_timeout is not None:
            elapsed = time.perf_counter() - overall_start_time
            if elapsed >= overall_timeout:
                error_msg = f"OCR overall timeout exceeded at page {page_num}: {elapsed:.1f}s >= {overall_timeout}s"
                _append_log(
                    log_path,
                    "WARNING",
                    f"request={request_id} ocr_overall_timeout page={page_num} "
                    f"elapsed={elapsed:.1f}s limit={overall_timeout}s",
                )
                return (
                    page_num,
                    {
                        "page_number": page_num,
                        "lines": [],
                        "error": "timeout",
                        "error_message": error_msg,
                    },
                    page_retry_stats,
                    "",  # No text on error
                    error_msg,
                )
        
        # Store original dimensions for bounding box adjustment and noise filtering
        original_page_w, original_page_h = img.size
        
        # Optimize image for OCR (downscale if too large)
        optimized_img, scale_x, scale_y = _optimize_image_for_ocr(
            img=img,
            max_dimension=image_max_dimension,
            min_dimension_for_optimization=image_min_dimension_for_optimization,
            enabled=image_optimization_enabled,
        )
        
        # Log optimization if applied
        if scale_x != 1.0 or scale_y != 1.0:
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_image_optimized page={page_num} "
                f"original_size={original_page_w}x{original_page_h} "
                f"optimized_size={optimized_img.size[0]}x{optimized_img.size[1]} "
                f"scale_x={scale_x:.2f} scale_y={scale_y:.2f}",
            )
        
        # Use optimized image for OCR
        buffer = io.BytesIO()
        optimized_img.save(buffer, format="PNG")
        vision_image = vision.Image(content=buffer.getvalue())
        
        # Use original dimensions for noise filtering (bounding boxes will be adjusted)
        page_w, page_h = original_page_w, original_page_h
        
        # Call Vision API with timeout protection and retry logic
        page_start_time = time.perf_counter()
        response = _call_vision_with_retry(
            vision_client=vision_client,
            vision_image=vision_image,
            timeout_seconds=per_page_timeout,
            page_number=page_num,
            max_retries=max_retries,
            retry_base_delay=retry_base_delay,
            retry_max_delay=retry_max_delay,
            retry_jitter_range=retry_jitter_range,
            rate_limit_base_delay=rate_limit_base_delay,
            rate_limit_max_delay=rate_limit_max_delay,
            overall_timeout=overall_timeout,
            overall_start_time=overall_start_time,
            log_path=log_path,
            request_id=request_id,
            retry_stats=page_retry_stats,
        )
        page_duration = time.perf_counter() - page_start_time
        
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} ocr_page_complete page={page_num} "
            f"duration_ms={int(page_duration * 1000)}",
        )
        
        # Process the response
        page_lines: List[Dict[str, Any]] = []
        annotation = response.full_text_annotation
        full_text = ""
        
        if annotation:
            full_text = annotation.text.strip()
            for page in annotation.pages:
                for block in page.blocks:
                    for paragraph in block.paragraphs:
                        text = _paragraph_text(paragraph)
                        para_bbox = _bbox_to_tuples(paragraph.bounding_box)
                        
                        # Adjust bounding box for image optimization (scale up to original dimensions)
                        if scale_x != 1.0 or scale_y != 1.0:
                            para_bbox = [
                                (int(x * scale_x), int(y * scale_y)) for x, y in para_bbox
                            ]
                        
                        # Filter noise (using original page dimensions)
                        if _is_noise_text(text, para_bbox, page_w, page_h):
                            continue
                        
                        word_entries: List[Dict[str, Any]] = []
                        for word in paragraph.words:
                            w_text = "".join(
                                symbol.text for symbol in word.symbols
                            ).strip()
                            if not w_text:
                                continue
                            word_bbox = _bbox_to_tuples(word.bounding_box)
                            
                            # Adjust bounding box for image optimization (scale up to original dimensions)
                            if scale_x != 1.0 or scale_y != 1.0:
                                word_bbox = [
                                    (int(x * scale_x), int(y * scale_y)) for x, y in word_bbox
                                ]
                            
                            # Filter noise words (using original page dimensions)
                            if _is_noise_text(w_text, word_bbox, page_w, page_h):
                                continue
                            
                            word_entries.append({
                                "text": w_text,
                                "bbox": word_bbox,
                            })
                        
                        if word_entries:  # Only add paragraph if it has valid words
                            page_lines.append({
                                "text": text,
                                "bbox": para_bbox,
                                "words": word_entries,
                            })
        else:
            text_annotations = response.text_annotations
            if text_annotations:
                full_text = text_annotations[0].description.strip()
                for ta in text_annotations[1:]:
                    ta_bbox = _bbox_to_tuples(ta.bounding_poly)
                    
                    # Adjust bounding box for image optimization (scale up to original dimensions)
                    if scale_x != 1.0 or scale_y != 1.0:
                        ta_bbox = [
                            (int(x * scale_x), int(y * scale_y)) for x, y in ta_bbox
                        ]
                    
                    # Filter noise (using original page dimensions)
                    if not _is_noise_text(ta.description, ta_bbox, page_w, page_h):
                        page_lines.append({
                            "text": ta.description,
                            "bbox": ta_bbox,
                            "words": [],
                        })
        
        return (
            page_num,
            {"page_number": page_num, "lines": page_lines},
            page_retry_stats,
            full_text,
            None,  # No error
        )
        
    except TimeoutError as e:
        # Handle per-page timeout
        error_msg = str(e)
        _append_log(
            log_path,
            "ERROR",
            f"request={request_id} ocr_page_timeout page={page_num} "
            f"timeout={per_page_timeout}s error={error_msg}",
        )
        return (
            page_num,
            {
                "page_number": page_num,
                "lines": [],
                "error": "timeout",
                "error_message": error_msg,
            },
            page_retry_stats,
            "",  # No text on error
            error_msg,
        )
        
    except RuntimeError as e:
        # Handle API errors
        error_msg = str(e)
        _append_log(
            log_path,
            "ERROR",
            f"request={request_id} ocr_page_error page={page_num} error={error_msg}",
        )
        return (
            page_num,
            {
                "page_number": page_num,
                "lines": [],
                "error": "api_error",
                "error_message": error_msg,
            },
            page_retry_stats,
            "",  # No text on error
            error_msg,
        )
        
    except Exception as e:
        # Handle unexpected errors
        error_msg = str(e)
        _append_log(
            log_path,
            "ERROR",
            f"request={request_id} ocr_page_unexpected_error page={page_num} "
            f"error={error_msg}",
        )
        return (
            page_num,
            {
                "page_number": page_num,
                "lines": [],
                "error": "unexpected_error",
                "error_message": error_msg,
            },
            page_retry_stats,
            "",  # No text on error
            error_msg,
        )


def run_ocr_on_pdf(
    vision_client: vision.ImageAnnotatorClient,
    pdf_path: str,
    per_page_timeout: float = 120.0,
    overall_timeout: Optional[float] = None,
    log_path: Optional[str] = None,
    request_id: Optional[str] = None,
    max_retries: int = 3,
    progress_tracker: Optional[Any] = None,
    retry_base_delay: float = 1.0,
    retry_max_delay: float = 60.0,
    retry_jitter_range: float = 0.2,
    rate_limit_base_delay: float = 5.0,
    rate_limit_max_delay: float = 300.0,
    concurrent_pages: int = 2,
    batch_size: int = 5,
    batch_failure_threshold: float = 0.5,
    adaptive_concurrency_enabled: bool = True,
    adaptive_min_concurrency: int = 1,
    adaptive_max_concurrency: int = 4,
    adaptive_latency_threshold_ms: float = 90000.0,
    adaptive_stable_batches: int = 2,
    image_optimization_enabled: bool = True,
    image_max_dimension: int = 2048,
    image_min_dimension_for_optimization: int = 1500,
) -> Dict[str, Any]:
    """
    Run Google Cloud Vision DOCUMENT_TEXT_DETECTION on each page of the PDF.
    Filters out background noise and artifacts.
    
    Args:
        vision_client: Google Vision client
        pdf_path: Path to PDF file
        per_page_timeout: Maximum seconds per page OCR call (default: 120)
        overall_timeout: Maximum seconds for entire OCR process (default: None = no limit)
        log_path: Optional path to log file for timeout logging
        request_id: Optional request ID for logging
        max_retries: Maximum retry attempts per page (default: 3)
        retry_base_delay: Base delay in seconds for exponential backoff (default: 1.0)
        retry_max_delay: Maximum delay in seconds between retries (default: 60.0)
        retry_jitter_range: Jitter range (0.0-1.0) for backoff randomization (default: 0.2)
        rate_limit_base_delay: Base delay in seconds for rate limit backoff (default: 5.0)
        rate_limit_max_delay: Maximum delay in seconds for rate limit backoff (default: 300.0)
        concurrent_pages: Number of pages to process in parallel (default: 2)
        batch_size: Number of pages to process per batch (default: 5)
        batch_failure_threshold: Stop processing if batch failure rate exceeds this (0.0-1.0, default: 0.5)
    
    Returns:
        Dictionary with 'pages' (list of page OCR data) and 'full_text' (concatenated text)
    
    Raises:
        TimeoutError: If overall timeout is exceeded or per-page timeout occurs
        RuntimeError: If OCR processing fails
    """
    # Track overall processing time for timeout enforcement
    overall_start_time = time.perf_counter()
    
    doc = fitz.open(pdf_path)
    try:
        images = []
        for page in doc:
            # Check overall timeout before processing each page
            if overall_timeout is not None:
                elapsed = time.perf_counter() - overall_start_time
                if elapsed >= overall_timeout:
                    raise TimeoutError(
                        f"OCR overall timeout exceeded: {elapsed:.1f}s >= {overall_timeout}s"
                    )
            
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            pil_img = Image.open(io.BytesIO(img_bytes))
            images.append(pil_img)

        pages_output: List[Dict[str, Any]] = []
        full_text_parts: List[str] = []
        failed_pages: List[Dict[str, Any]] = []  # Track failed pages for partial success
        
        # Track aggregate retry statistics across all pages
        aggregate_retry_stats = {
            "total_retry_attempts": 0,
            "successful_retries": 0,
            "exhausted_retries": 0,
            "rate_limit_events": 0,
            "non_retryable_errors": 0,
            "budget_exceeded": 0,
            "retry_attempts_by_category": {},
        }

        # CONDITIONAL PARALLEL OCR: Only parallelize when beneficial
        # Small files (≤5 pages): Sequential processing (no parallel overhead)
        # Medium files (6-15 pages): Low concurrency (2 pages)
        # Large files (16+ pages): Higher concurrency (up to 4 pages)
        total_pages = len(images)
        if total_pages <= 5:
            # Small files: Use sequential processing to avoid parallel overhead
            concurrent_pages = 1
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_conditional_parallel pages={total_pages} "
                f"mode=sequential reason=small_file",
            )
        elif total_pages <= 15:
            # Medium files: Use low concurrency (2 pages)
            concurrent_pages = min(2, total_pages)
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_conditional_parallel pages={total_pages} "
                f"mode=low_concurrency concurrent_pages={concurrent_pages}",
            )
        else:
            # Large files: Use higher concurrency (up to 4 pages)
            concurrent_pages = min(4, total_pages)
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_conditional_parallel pages={total_pages} "
                f"mode=high_concurrency concurrent_pages={concurrent_pages}",
            )
        
        # Ensure concurrent_pages is at least 1 and not more than total pages
        concurrent_pages = max(1, min(concurrent_pages, total_pages))
        
        # Create list to store results (maintains order by page number)
        results: List[Tuple[int, Dict[str, Any], Dict[str, Any], str, Optional[str]]] = []
        
        # WARM-UP PHASE: Process page 1 sequentially first to warm up API connections
        # This prevents cold-start overhead when parallel processing begins
        if len(images) > 0:
            warmup_start_time = time.perf_counter()
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_warmup_start page=1",
            )
            
            # Process page 1 sequentially (outside ThreadPoolExecutor)
            page1_result = _process_single_page_ocr(
                vision_client=vision_client,
                img=images[0],
                page_num=1,
                per_page_timeout=per_page_timeout,
                overall_timeout=overall_timeout,
                overall_start_time=overall_start_time,
                max_retries=max_retries,
                retry_base_delay=retry_base_delay,
                retry_max_delay=retry_max_delay,
                retry_jitter_range=retry_jitter_range,
                rate_limit_base_delay=rate_limit_base_delay,
                rate_limit_max_delay=rate_limit_max_delay,
                log_path=log_path,
                request_id=request_id,
                image_optimization_enabled=image_optimization_enabled,
                image_max_dimension=image_max_dimension,
                image_min_dimension_for_optimization=image_min_dimension_for_optimization,
            )
            results.append(page1_result)
            
            # Update progress after page 1
            if progress_tracker:
                total_pages = len(images)
                progress_tracker.update_progress(
                    request_id=request_id,
                    step="OCR Processing",
                    step_number=2,
                    total_steps=11,  # Will be updated dynamically
                    progress_percent=15.0 + (1 / total_pages) * 30.0,  # 15-45% for OCR
                    message=f"Processing page 1 of {total_pages}...",
                    details={"pages_completed": 1, "total_pages": total_pages},
                )
            
            warmup_duration = time.perf_counter() - warmup_start_time
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_warmup_complete page=1 duration_ms={int(warmup_duration * 1000)}",
            )
        
        # BATCH ORCHESTRATION PHASE: Process remaining pages (2+) in batches
        remaining_pages = len(images) - 1  # Exclude page 1 which was already processed
        
        if remaining_pages > 0:
            # Adjust concurrent_pages for remaining pages
            effective_concurrency = max(1, min(concurrent_pages, remaining_pages))
            
            # Initialize adaptive concurrency tracking
            if adaptive_concurrency_enabled:
                # Ensure concurrency is within adaptive bounds
                effective_concurrency = max(adaptive_min_concurrency, min(effective_concurrency, adaptive_max_concurrency))
                stable_batch_count = 0  # Count consecutive stable batches
                previous_concurrency = effective_concurrency
            
            # Calculate number of batches
            num_batches = (remaining_pages + batch_size - 1) // batch_size  # Ceiling division
            
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_batch_orchestration_start total_pages={len(images)} "
                f"remaining_pages={remaining_pages} batch_size={batch_size} num_batches={num_batches} "
                f"concurrent_pages={effective_concurrency} "
                f"adaptive_enabled={adaptive_concurrency_enabled}",
            )
            
            # Process pages in batches
            remaining_indices = list(range(1, len(images)))  # Indices for pages 2+
            batch_num = 0
            should_continue = True
            
            while remaining_indices and should_continue:
                batch_num += 1
                
                # Check overall timeout before starting batch
                if overall_timeout is not None:
                    elapsed = time.perf_counter() - overall_start_time
                    if elapsed >= overall_timeout:
                        _append_log(
                            log_path,
                            "WARNING",
                            f"request={request_id} ocr_batch_timeout_check batch={batch_num} "
                            f"elapsed={elapsed:.1f}s limit={overall_timeout}s stopping",
                        )
                        break
                
                # Get next batch of pages
                batch_indices = remaining_indices[:batch_size]
                batch_pages = [idx + 1 for idx in batch_indices]  # Convert to page numbers (1-indexed)
                batch_start_time = time.perf_counter()
                
                _append_log(
                    log_path,
                    "INFO",
                    f"request={request_id} ocr_batch_start batch={batch_num}/{num_batches} "
                    f"pages={batch_pages} size={len(batch_pages)}",
                )
                
                # Process batch in parallel
                batch_results: List[Tuple[int, Dict[str, Any], Dict[str, Any], str, Optional[str]]] = []
                batch_futures = {}
                
                with ThreadPoolExecutor(max_workers=min(effective_concurrency, len(batch_indices))) as executor:
                    # Submit batch tasks
                    for idx in batch_indices:
                        page_num = idx + 1
                        future = executor.submit(
                            _process_single_page_ocr,
                            vision_client,
                            images[idx],
                            page_num,
                            per_page_timeout,
                            overall_timeout,
                            overall_start_time,
                            max_retries,
                            retry_base_delay,
                            retry_max_delay,
                            retry_jitter_range,
                            rate_limit_base_delay,
                            rate_limit_max_delay,
                            log_path,
                            request_id,
                            image_optimization_enabled,
                            image_max_dimension,
                            image_min_dimension_for_optimization,
                        )
                        batch_futures[future] = page_num
                    
                    # Collect batch results
                    for future in batch_futures:
                        try:
                            result = future.result()
                            batch_results.append(result)
                        except Exception as e:
                            # Handle unexpected executor errors
                            page_num = batch_futures[future]
                            error_msg = f"Executor error on page {page_num}: {str(e)}"
                            _append_log(
                                log_path,
                                "ERROR",
                                f"request={request_id} ocr_executor_error page={page_num} error={error_msg}",
                            )
                            batch_results.append((
                                page_num,
                                {
                                    "page_number": page_num,
                                    "lines": [],
                                    "error": "executor_error",
                                    "error_message": error_msg,
                                },
                                {
                                    "total_attempts": 0,
                                    "successful_retries": 0,
                                    "exhausted_retries": 0,
                                    "rate_limit_events": 0,
                                    "non_retryable_errors": 0,
                                    "budget_exceeded": 0,
                                    "retry_attempts_by_category": {},
                                },
                                "",  # No text on error
                                error_msg,
                            ))
                
                # Analyze batch results
                batch_duration = time.perf_counter() - batch_start_time
                batch_success = sum(1 for r in batch_results if r[4] is None)  # Count successes (error_msg is None)
                batch_failures = len(batch_results) - batch_success
                batch_failure_rate = batch_failures / len(batch_results) if batch_results else 0.0
                
                # Calculate batch metrics for adaptive concurrency
                batch_rate_limit_events = sum(r[2].get("rate_limit_events", 0) for r in batch_results)
                batch_total_attempts = sum(r[2].get("total_attempts", 0) for r in batch_results)
                
                # Calculate average latency per page in batch (for successful pages)
                successful_results = [r for r in batch_results if r[4] is None]
                if successful_results:
                    # Estimate average latency from retry stats (rough approximation)
                    # In practice, we'd track actual page durations, but for now use batch duration / pages
                    avg_latency_ms = (batch_duration / len(successful_results)) * 1000 if successful_results else 0
                else:
                    avg_latency_ms = 0
                
                # Add batch results to overall results
                results.extend(batch_results)
                
                # Update progress after batch completes
                if progress_tracker:
                    total_pages = len(images)
                    pages_completed = len(results)  # Total pages completed so far
                    progress_percent = 15.0 + (pages_completed / total_pages) * 30.0  # 15-45% for OCR
                    progress_tracker.update_progress(
                        request_id=request_id,
                        step="OCR Processing",
                        step_number=2,
                        total_steps=11,
                        progress_percent=progress_percent,
                        message=f"Processing page {pages_completed} of {total_pages}...",
                        details={"pages_completed": pages_completed, "total_pages": total_pages},
                    )
                
                # Remove processed indices
                remaining_indices = remaining_indices[len(batch_indices):]
                
                _append_log(
                    log_path,
                    "INFO",
                    f"request={request_id} ocr_batch_complete batch={batch_num}/{num_batches} "
                    f"pages={batch_pages} duration_ms={int(batch_duration * 1000)} "
                    f"success={batch_success} failures={batch_failures} "
                    f"failure_rate={batch_failure_rate:.2%} "
                    f"rate_limits={batch_rate_limit_events} avg_latency_ms={int(avg_latency_ms)}",
                )
                
                # ADAPTIVE CONCURRENCY: Adjust concurrency based on batch performance
                if adaptive_concurrency_enabled and batch_num > 1:  # Start adapting after first batch
                    concurrency_changed = False
                    new_concurrency = effective_concurrency
                    
                    # Reduce concurrency if rate limits occurred
                    if batch_rate_limit_events > 0:
                        new_concurrency = max(adaptive_min_concurrency, effective_concurrency - 1)
                        if new_concurrency < effective_concurrency:
                            concurrency_changed = True
                            stable_batch_count = 0  # Reset stable count on reduction
                            _append_log(
                                log_path,
                                "WARNING",
                                f"request={request_id} ocr_adaptive_concurrency_reduce batch={batch_num} "
                                f"rate_limits={batch_rate_limit_events} "
                                f"concurrency={effective_concurrency}->{new_concurrency} reason=rate_limits",
                            )
                    
                    # Reduce concurrency if average latency exceeds threshold
                    elif avg_latency_ms > adaptive_latency_threshold_ms and effective_concurrency > adaptive_min_concurrency:
                        new_concurrency = max(adaptive_min_concurrency, effective_concurrency - 1)
                        if new_concurrency < effective_concurrency:
                            concurrency_changed = True
                            stable_batch_count = 0  # Reset stable count on reduction
                            _append_log(
                                log_path,
                                "WARNING",
                                f"request={request_id} ocr_adaptive_concurrency_reduce batch={batch_num} "
                                f"avg_latency_ms={int(avg_latency_ms)} threshold={int(adaptive_latency_threshold_ms)} "
                                f"concurrency={effective_concurrency}->{new_concurrency} reason=high_latency",
                            )
                    
                    # Increase concurrency if stable (no rate limits, low latency)
                    elif (batch_rate_limit_events == 0 and 
                          avg_latency_ms < adaptive_latency_threshold_ms and 
                          effective_concurrency < adaptive_max_concurrency):
                        stable_batch_count += 1
                        if stable_batch_count >= adaptive_stable_batches:
                            new_concurrency = min(adaptive_max_concurrency, effective_concurrency + 1)
                            if new_concurrency > effective_concurrency:
                                concurrency_changed = True
                                stable_count_before_reset = stable_batch_count
                                stable_batch_count = 0  # Reset after increase
                                _append_log(
                                    log_path,
                                    "INFO",
                                    f"request={request_id} ocr_adaptive_concurrency_increase batch={batch_num} "
                                    f"stable_batches={stable_count_before_reset} "
                                    f"concurrency={effective_concurrency}->{new_concurrency} reason=stable_performance",
                                )
                    else:
                        # Stable but not ready to increase yet
                        stable_batch_count += 1
                    
                    # Update concurrency if changed
                    if concurrency_changed:
                        previous_concurrency = effective_concurrency
                        effective_concurrency = new_concurrency
                        # Ensure concurrency doesn't exceed remaining pages
                        effective_concurrency = min(effective_concurrency, len(remaining_indices))
                
                # Check if we should continue based on failure rate
                if batch_failure_rate > batch_failure_threshold:
                    _append_log(
                        log_path,
                        "WARNING",
                        f"request={request_id} ocr_batch_high_failure_rate batch={batch_num} "
                        f"failure_rate={batch_failure_rate:.2%} threshold={batch_failure_threshold:.2%} "
                        f"stopping_processing",
                    )
                    should_continue = False
                    break
                
                # Check overall timeout after batch
                if overall_timeout is not None:
                    elapsed = time.perf_counter() - overall_start_time
                    remaining_time = overall_timeout - elapsed
                    if remaining_time < (per_page_timeout * 2):  # Not enough time for another batch
                        _append_log(
                            log_path,
                            "INFO",
                            f"request={request_id} ocr_batch_timeout_check batch={batch_num} "
                            f"elapsed={elapsed:.1f}s remaining={remaining_time:.1f}s "
                            f"insufficient_time_for_next_batch stopping",
                        )
                        break
            
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_batch_orchestration_complete batches_processed={batch_num} "
                f"remaining_pages={len(remaining_indices)}",
            )
        elif len(images) == 1:
            # Only one page, already processed in warm-up
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_single_page_complete page=1",
            )
        
        # Sort results by page number to maintain order
        results.sort(key=lambda x: x[0])
        
        # Process results and aggregate statistics
        for page_num, page_output, page_retry_stats, page_full_text, error_msg in results:
            # Aggregate retry statistics
            aggregate_retry_stats["total_retry_attempts"] += page_retry_stats["total_attempts"]
            aggregate_retry_stats["successful_retries"] += page_retry_stats["successful_retries"]
            aggregate_retry_stats["exhausted_retries"] += page_retry_stats["exhausted_retries"]
            aggregate_retry_stats["rate_limit_events"] += page_retry_stats["rate_limit_events"]
            aggregate_retry_stats["non_retryable_errors"] += page_retry_stats["non_retryable_errors"]
            aggregate_retry_stats["budget_exceeded"] += page_retry_stats["budget_exceeded"]
            # Merge category stats
            for category, count in page_retry_stats["retry_attempts_by_category"].items():
                if category not in aggregate_retry_stats["retry_attempts_by_category"]:
                    aggregate_retry_stats["retry_attempts_by_category"][category] = 0
                aggregate_retry_stats["retry_attempts_by_category"][category] += count
            
            # Add page output
            pages_output.append(page_output)
            
            # Add full text if available
            if page_full_text:
                full_text_parts.append(page_full_text)
            
            # Track failed pages
            if error_msg:
                failed_pages.append({
                    "page_number": page_num,
                    "error": page_output.get("error", "unknown_error"),
                    "error_message": error_msg,
                })

        # Log summary
        total_duration = time.perf_counter() - overall_start_time
        success_count = len(pages_output) - len(failed_pages)
        total_pages = len(images)
        
        # Calculate retry success rate
        retry_success_rate = 0.0
        if aggregate_retry_stats["total_retry_attempts"] > 0:
            retry_success_rate = (
                aggregate_retry_stats["successful_retries"] / 
                aggregate_retry_stats["total_retry_attempts"]
            ) * 100.0
        
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} ocr_complete total_pages={total_pages} "
            f"success={success_count} failed={len(failed_pages)} "
            f"duration_ms={int(total_duration * 1000)}",
        )
        
        # Log retry statistics summary
        if aggregate_retry_stats["total_retry_attempts"] > 0:
            category_summary = ", ".join([
                f"{cat}={count}" 
                for cat, count in aggregate_retry_stats["retry_attempts_by_category"].items()
            ])
            _append_log(
                log_path,
                "INFO",
                f"request={request_id} ocr_retry_stats "
                f"total_attempts={aggregate_retry_stats['total_retry_attempts']} "
                f"successful_retries={aggregate_retry_stats['successful_retries']} "
                f"retry_success_rate_pct={retry_success_rate:.1f} "
                f"exhausted={aggregate_retry_stats['exhausted_retries']} "
                f"rate_limits={aggregate_retry_stats['rate_limit_events']} "
                f"non_retryable={aggregate_retry_stats['non_retryable_errors']} "
                f"budget_exceeded={aggregate_retry_stats['budget_exceeded']} "
                f"categories=[{category_summary}]",
            )
        
        # If all pages failed, raise an error
        if len(failed_pages) == total_pages:
            raise RuntimeError(
                f"OCR failed on all {total_pages} pages. "
                f"Errors: {[p['error'] for p in failed_pages]}"
            )
        
        # If some pages failed, log warning but return partial results
        if failed_pages:
            failed_page_nums = [p["page_number"] for p in failed_pages]
            _append_log(
                log_path,
                "WARNING",
                f"request={request_id} ocr_partial_success "
                f"failed_pages={failed_page_nums}",
            )

        return {
            "pages": pages_output,
            "full_text": "\n".join(full_text_parts).strip(),
            "metadata": {
                "total_pages": total_pages,
                "successful_pages": success_count,
                "failed_pages": len(failed_pages),
                "failed_page_numbers": [p["page_number"] for p in failed_pages],
                "processing_duration_seconds": total_duration,
                "retry_statistics": {
                    "total_retry_attempts": aggregate_retry_stats["total_retry_attempts"],
                    "successful_retries": aggregate_retry_stats["successful_retries"],
                    "retry_success_rate_percent": round(retry_success_rate, 2),
                    "exhausted_retries": aggregate_retry_stats["exhausted_retries"],
                    "rate_limit_events": aggregate_retry_stats["rate_limit_events"],
                    "non_retryable_errors": aggregate_retry_stats["non_retryable_errors"],
                    "budget_exceeded": aggregate_retry_stats["budget_exceeded"],
                    "retry_attempts_by_category": aggregate_retry_stats["retry_attempts_by_category"],
                },
            }
        }
    finally:
        doc.close()  # Always close the document to release file handle


# -----------------------------
# RUBRIC DOCX HELPERS
# -----------------------------


def _load_docx_text(path: str) -> str:
    try:
        doc = Document(path)
    except Exception as exc:
        return f"[Error reading DOCX at {path}: {exc}]"
    parts: List[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _normalize_subject_key(value: str) -> str:
    """
    Convert any subject identifier (dropdown id, folder name, filename) into a
    comparable slug so "Political Science Rubric" and "political-science" match.
    """
    if not value:
        return ""
    key = value.lower().strip()
    key = re.sub(r"[\s_]+", "-", key)
    key = re.sub(r"[^a-z0-9-]", "", key)
    key = re.sub(r"-{2,}", "-", key)
    return key.strip("-")


def _subject_key_variants(value: str) -> Set[str]:
    """
    Generate slug variants for matching; we also accept subjects without the
    trailing "-rubric" suffix because some folders include that word.
    """
    base = _normalize_subject_key(value)
    variants: Set[str] = set()
    if base:
        variants.add(base)
        if base.endswith("-rubric"):
            trimmed = base[: -len("-rubric")].strip("-")
            if trimmed:
                variants.add(trimmed)
    return variants


def _subject_keys_match(candidate: str, target_variants: Set[str]) -> bool:
    candidate_variants = _subject_key_variants(candidate)
    return any(var in target_variants for var in candidate_variants)


def find_subject_rubric_path(subject: str) -> Optional[str]:
    """
    Search under ./20marks_Rubrics for a .docx whose folder or filename matches
    the provided subject id (case/spacing insensitive).
    """
    base_dir = os.path.dirname(os.path.abspath(__file__))
    rubrics_root = os.path.join(base_dir, "20marks_Rubrics")
    if not os.path.isdir(rubrics_root):
        return None

    target_variants = _subject_key_variants(subject)
    if not target_variants:
        return None

    # Prefer matches where the directory name (subject folder) aligns with the
    # requested id; fall back to filename matches if needed.
    for entry in sorted(os.listdir(rubrics_root)):
        entry_path = os.path.join(rubrics_root, entry)
        if not os.path.isdir(entry_path):
            continue

        docx_files = sorted(
            f for f in os.listdir(entry_path) if f.lower().endswith(".docx")
        )
        if not docx_files:
            continue

        dir_matches = _subject_keys_match(entry, target_variants)
        for fname in docx_files:
            stem = os.path.splitext(fname)[0]
            if dir_matches or _subject_keys_match(stem, target_variants):
                return os.path.join(entry_path, fname)

    # Handle DOCX files that might live directly under the root.
    for fname in sorted(os.listdir(rubrics_root)):
        if not fname.lower().endswith(".docx"):
            continue
        stem = os.path.splitext(fname)[0]
        if _subject_keys_match(stem, target_variants):
            return os.path.join(rubrics_root, fname)

    return None


def load_subject_rubric_text(subject: str) -> Tuple[str, Optional[str]]:
    docx_path = find_subject_rubric_path(subject)
    if not docx_path:
        print(f"WARNING: No subject rubric DOCX found for '{subject}'.")
        return "", None
    text = _load_docx_text(docx_path)
    return text, docx_path


def load_refined_rubric_text() -> Tuple[str, Optional[str]]:
    """
    Load the refined generic rubric DOCX (REFINED RUBRIC (1).docx).
    """
    base_dir = os.path.dirname(os.path.abspath(__file__))
    refined_path = os.path.join(base_dir, "REFINED RUBRIC.docx")
    if not os.path.isfile(refined_path):
        print("WARNING: Refined rubric DOCX not found.")
        return "", None
    text = _load_docx_text(refined_path)
    return text, refined_path


# -----------------------------
# GROK CALL 1: SECTION DETECTION
# -----------------------------

def call_grok_for_section_detection(
    grok_api_key: str,
    ocr_data: Dict[str, Any],
    page_images: List[Dict[str, Any]],
) -> Tuple[List[Dict[str, Any]], Dict[str, int]]:
    """
    Use Grok to detect headings, subheadings and content sections
    from the full OCR text + ALL page images.

    Returns a list of sections in the shape:
      {
        "title": str,
        "level": int,            # 1 = main heading, 2 = subheading
        "page_numbers": [int],
        "content": str,
        "line_indices": []       # kept for compatibility (not used)
      }

    IMPORTANT:
      - Grok is explicitly told to trust the PAGE IMAGES as the primary source
        and only use OCR as a helper.
    """
    system_msg = {
        "role": "system",
        "content": (
            "You are an expert at visually and logically segmenting handwritten exam answers.\n\n"
            "INPUT DATA YOU RECEIVE:\n"
            "- OCR text (approximate, may contain errors)\n"
            "- Per-page OCR text lines (text only, no bounding boxes or coordinates)\n"
            "- Base64-encoded page images of the handwritten script\n\n"
            "PRIMARY RULE:\n"
            "- The PAGE IMAGES are the primary source of truth\n"
            "- OCR text is only a helper for searching or clarifying words\n"
            "- If OCR and image disagree, ALWAYS trust the handwritten image\n\n"
            "YOUR TASK:\n"
            "Segment the answer into logical sections with this structure:\n"
            "Introduction → Main body sections (with optional subsections) → Conclusion\n\n"
            "STRICT REQUIREMENTS FOR SECTION DETECTION:\n\n"
            "1) INTRODUCTION (REQUIRED - MUST BE FIRST):\n"
            "   - If there is an explicit heading like 'Introduction', use that exact text as the title\n"
            "   - If no such heading exists, identify the first paragraph(s) that introduce the topic\n"
            "   - Set title to 'Introduction' for implicit introductions\n"
            "   - For 'exact_ocr_heading': copy the EXACT OCR text of the heading line (or first line if implicit)\n"
            "   - This MUST be the first section in your output\n\n"
            "2) BODY SECTIONS (IDENTIFY ALL MAJOR HEADINGS):\n"
            "   - Examine the page images carefully for ALL headings between Introduction and Conclusion\n"
            "   - Look for these VISUAL CUES in the images:\n"
            "     • Larger or bolder handwriting\n"
            "     • Underlined words or phrases\n"
            "     • Extra spacing above and/or below text\n"
            "     • Numbered headings: 1., 2., 3., or i., ii., iii., or (a), (b), (c)\n"
            "     • Short phrases at the start of a line that label the content below\n"
            "   - For each heading found:\n"
            "     • Create a section with that heading as the title\n"
            "     • Set 'level' = 1 for main topics, 'level' = 2 for subtopics under the previous main heading\n"
            "     • For 'exact_ocr_heading': copy the EXACT OCR text of that heading line (word-for-word, with any typos)\n"
            "     • Include ALL page numbers where this section's content appears\n"
            "   - DO NOT skip headings - find ALL of them\n"
            "   - DO NOT invent headings that don't exist visually\n\n"
            "3) CONCLUSION (REQUIRED - MUST BE LAST):\n"
            "   - If there is an explicit heading like 'Conclusion' or 'In conclusion', use that text\n"
            "   - If no such heading, identify the final paragraph(s) that summarize/wrap up the answer\n"
            "   - Set title to 'Conclusion' for implicit conclusions\n"
            "   - For 'exact_ocr_heading': copy the EXACT OCR text of the heading line (or last paragraph's first line if implicit)\n"
            "   - This MUST be the last section in your output\n\n"
            "SECTION DEFINITIONS:\n"
            "- A section is a continuous block of content belonging together under one heading/topic\n"
            "- Sections must appear in reading order (top to bottom, page by page)\n"
            "- Sections must NOT overlap - each line belongs to only one section\n"
            "- If you cannot confidently assign some lines, leave them out rather than forcing them\n\n"
            "CONTENT_TEXT REQUIREMENTS:\n"
            "- For each section, provide a concise summary of the student's actual content\n"
            "- Use your own words but stay faithful to what is written\n"
            "- Do NOT invent arguments or facts not present in the script\n"
            "- If handwriting is unclear, infer only what is reasonably supported\n\n"
            "HEADING QUALITY EVALUATION (NEW REQUIREMENT):\n"
            "For each heading/subheading (excluding 'Introduction' and 'Conclusion'), provide a 'comment' evaluating its quality:\n"
            "- Start with POSITIVE or NEGATIVE to clearly indicate the assessment\n"
            "- Evaluate based on these criteria:\n"
            "  • Is it self-explanatory? Does it immediately tell what point is being discussed?\n"
            "  • Is it directly relevant to the question and reflects exact themes from the question statement?\n"
            "  • Does it guide the reader by showing how this section contributes to answering the question?\n"
            "  • Or is it generic, vague, or acts as a decorative slogan rather than a meaningful signpost?\n"
            "- Examples:\n"
            "  • POSITIVE: This heading clearly identifies the specific aspect being analyzed and directly addresses a component of the question.\n"
            "  • NEGATIVE: This heading is too generic and doesn't clearly indicate what specific point will be discussed.\n"
            "- Keep comments concise (1-2 sentences)\n"
            "- For 'Introduction' and 'Conclusion' sections, you may set comment to empty string or a brief note\n\n"
            "OUTPUT FORMAT (CRITICAL):\n"
            "- Return ONLY valid JSON with structure: {\"sections\": [...]}\n"
            "- NO markdown formatting, NO code blocks, NO explanations\n"
            "- Each section object must have these EXACT fields:\n"
            "  • 'title': Clean, readable heading text\n"
            "  • 'exact_ocr_heading': EXACT OCR text with any typos/errors (used for precise location matching)\n"
            "  • 'level': integer (1 for main, 2 for sub)\n"
            "  • 'page_numbers': array of integers\n"
            "  • 'content_text': string summary\n"
            "  • 'comment': quality evaluation starting with POSITIVE or NEGATIVE (or empty for intro/conclusion)\n"
            "- NO extra fields, NO top-level keys besides 'sections'\n\n"
            "CONSISTENCY:\n"
            "- For the same input, produce consistent segmentation\n"
            "- Avoid random or arbitrary splits\n"
            "- Prefer fewer, well-justified sections over many uncertain micro-sections\n"
            "- Be thorough but not excessive\n"
        ),
    }


    # Sanitize OCR pages: remove any bounding-box or coordinate data before sending
    raw_pages = ocr_data.get("pages", [])
    sanitized_pages = []
    for p in raw_pages:
        lines = []
        for line in p.get("lines", []):
            # Keep only textual content (line text + list of word texts) — drop bbox info
            line_text = line.get("text", "")
            words = [w.get("text", "") for w in line.get("words", [])]
            lines.append({"text": line_text, "words": words})
        sanitized_pages.append({"page_number": p.get("page_number"), "lines": lines})

    user_payload = {
        "task": (
            "Segment the handwritten exam answer into logical sections using BOTH the page images and OCR.\n\n"
            "INSTRUCTIONS:\n"
            "1) Identify an explicit INTRODUCTION section:\n"
            "   - If there is a heading like 'Introduction', use that as the section title.\n"
            "   - If there is no such heading, treat the first paragraph or group of lines that introduce\n"
            "     the topic as a section titled 'Introduction'.\n"
            "   - This section must always be the first section in the 'sections' array.\n\n"
            "2) Identify an explicit CONCLUSION section:\n"
            "   - If there is a heading like 'Conclusion', 'In conclusion', or similar, use that text as\n"
            "     the section title.\n"
            "   - If there is no such heading, treat the final summarising paragraph or lines that wrap up\n"
            "     the answer as a section titled 'Conclusion'.\n"
            "   - This section must always be the last section in the 'sections' array.\n\n"
            "3) Identify all major body sections between Introduction and Conclusion:\n"
            "   - Look for headings and subheadings using VISUAL CUES from the images:\n"
            "       • underlined or larger/bolder text,\n"
            "       • numbered headings (1., 2., 3.; i., ii., iii.),\n"
            "       • phrases at the start of lines followed by clearly related content,\n"
            "       • extra spacing above/below a short phrase.\n"
            "   - For each such heading:\n"
            "       • create a section with that title,\n"
            "       • assign level = 1 for main topics, level = 2 for clear subtopics under the\n"
            "         most recent main heading (use level = 1 if unsure).\n"
            "   - Include all pages where that section's content continues in 'page_numbers'.\n\n"
            "4) For each section's 'content_text':\n"
            "   - Provide a concise summary of the student's content for that section.\n"
            "   - Do NOT invent new arguments or facts that are not implied by the script.\n"
            "   - If handwriting is partially unclear, infer only what is reasonably supported.\n\n"
            "5) Global constraints for the output:\n"
            "   - The 'sections' array must:\n"
            "       • start with the Introduction section,\n"
            "       • list all body sections and sub-sections in the order they appear across pages,\n"
            "       • end with the Conclusion section.\n"
            "   - Each section must have: 'title', 'level', 'page_numbers', 'content_text'.\n"
            "   - Do NOT include any keys other than these four in each section.\n"
            "   - Do NOT include any top-level keys other than 'sections'.\n\n"
            "6) Output:\n"
            "   - Return ONLY a single JSON object of the form:\n"
            "     {\"sections\": [ {\"title\": ..., \"level\": ..., \"page_numbers\": [...], \"content_text\": ...}, ... ]}\n"
            "   - Do not wrap the JSON in markdown.\n"
            "   - Do not add explanations, comments, or any extra text.\n"
        ),
        # Send the full OCR text (no truncation) but send only sanitized page content
        "ocr_full_text": ocr_data.get("full_text", ""),
        "ocr_pages": sanitized_pages,
        "page_images_base64_png": page_images,
        "output_schema": {
            "sections": [
                {
                    "title": "string",
                    "exact_ocr_heading": "exact text from OCR",
                    "level": 1,
                    "page_numbers": [1],
                    "content_text": "string",
                }
            ]
        },
    }

    user_msg = {
        "role": "user",
        "content": json.dumps(user_payload, ensure_ascii=False),
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {grok_api_key}",
    }

    payload = {
        "model": "grok-4-fast-reasoning",
        "messages": [system_msg, user_msg],
        "temperature": 0.1,
        "max_tokens": 4000,  # Sufficient for structure detection
    }

    max_retries = 3
    for attempt in range(max_retries):
        if attempt > 0:
            # Exponential backoff: 1s, 2s, 4s
            import time
            time.sleep(2 ** (attempt - 1))
            print(f"Retry attempt {attempt + 1}/{max_retries} for section detection...")

        try:
            resp = requests.post(
                "https://api.x.ai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=120,
            )

            if resp.status_code >= 300:
                if attempt < max_retries - 1:
                    print(f"API error {resp.status_code}, retrying...")
                    continue
                raise RuntimeError(f"Section detection API error {resp.status_code}: {resp.text}")

            data = resp.json()

            # Extract token usage from response
            usage = data.get("usage", {})
            token_usage = {
                "input_tokens": usage.get("prompt_tokens", 0),
                "output_tokens": usage.get("completion_tokens", 0),
            }

            # Print token usage for this Grok section-detection prompt
            try:
                total_tokens = int(token_usage.get("input_tokens", 0)) + int(
                    token_usage.get("output_tokens", 0)
                )
            except Exception:
                total_tokens = None
            print(
                f"[Grok Section Detection] tokens — input={token_usage['input_tokens']}, output={token_usage['output_tokens']}, total={total_tokens}"
            )

            # Check for truncation
            finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
            if finish_reason == "length" and attempt < max_retries - 1:
                payload["max_tokens"] = payload.get("max_tokens", 4000) + 2000
                print(f"Response truncated, increasing max_tokens to {payload['max_tokens']} and retrying...")
                continue

            content = data["choices"][0]["message"]["content"]
            cleaned = clean_json_from_llm(content)
            parsed = json.loads(cleaned)

            # Success
            if attempt > 0:
                print(f"Successfully parsed section detection on attempt {attempt + 1}")

            break  # Exit retry loop on success

        except (requests.Timeout, requests.ConnectionError) as e:
            if attempt < max_retries - 1:
                print(f"Network error, retrying...")
                continue
            raise RuntimeError(f"Section detection network error: {e}") from e

        except json.JSONDecodeError as e:
            if attempt < max_retries - 1:
                print(f"JSON parse error, retrying...")
                continue
            # Fallback on final attempt
            return [
                {
                    "title": "Section Detection Error",
                    "level": 1,
                    "page_numbers": [1],
                    "content_text": f"Grok section detection failed: {e}\nRaw: {content[:400]}",
                }
            ], token_usage

        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Unexpected error: {exc}, retrying...")
                continue
            raise RuntimeError(f"Unexpected section detection error: {exc}") from exc

    # If we get here, parsing succeeded

    raw_sections = parsed.get("sections", []) or []

    sections: List[Dict[str, Any]] = []
    for sec in raw_sections:
        title = (sec.get("title") or "UNSPECIFIED").strip()
        exact_ocr_heading = (sec.get("exact_ocr_heading") or title).strip()
        level = sec.get("level") or 1
        pages = sec.get("page_numbers") or []
        content_text = sec.get("content_text") or sec.get("content") or ""

        sections.append(
            {
                "title": title,
                "exact_ocr_heading": exact_ocr_heading,  # Store exact OCR text
                "level": int(level) if isinstance(level, (int, float)) else 1,
                "page_numbers": sorted(
                    set(int(p) for p in pages if isinstance(p, (int, float)))
                ),
                "content": content_text,
                "line_indices": [],  # not used downstream, kept for compatibility
            }
        )

    return sections, token_usage


# -----------------------------
# GROK CALL 2: SUBJECT-WISE MARKING
# -----------------------------


def build_grok_payload_for_grading(
    subject: str,
    subject_rubric_text: str,
    ocr_data: Dict[str, Any],
    sections: List[Dict[str, Any]],
    page_images: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """
    Build payload for Grok subject-wise grading.

    We send:
      - subject,
      - subject rubric text (DOCX),
      - OCR full text,
      - sections structure,
      - page images.
    """
    schema_hint = {
        "subject": subject,
        "max_marks": 20,
        "total_marks_awarded": 0,
        "question_statement": "",
        "question_expectation": [
            "Bullet point 1: Key theme/concept expected",
            "Bullet point 2: Important theory or framework",
            "Bullet point 3: Historical context or period"
        ],
        "criteria": [
            {
                "id": "knowledge_accuracy",
                "name": "Knowledge & Accuracy",
                "max": 8,
                "awarded": 5,
                "strengths": ["..."],
                "weaknesses": ["..."],
            }
        ],
        "high_scoring_outline": {
            "title": "High-Scoring Ideal Outline",
            "outline_points": [
                {
                    "heading": "Section Title",
                    "summary": "2-3 sentence overview of what an excellent response covers in this section.",
                    "key_points": [
                        "Key argument or piece of evidence 1",
                        "Key argument or piece of evidence 2"
                    ],
                }
            ],
        },
        "overall_comment": "",
    }

    instructions = (
        "You are an experienced strict CSS examiner. "
        "Using ONLY the provided subject-wise rubric text, you must grade the student's answer with STRICT marking. "
        "IMPORTANT: Maximum marks awarded should NOT exceed 14 out of 20. "
        "Average/acceptable answers should score LESS than 10 marks. "
        "Only exceptional answers should approach 14 marks. "
        "Derive criteria and marks from the rubric text. Return STRICT JSON.\n\n"
        "Required fields:\n"
        "  - subject\n"
        "  - max_marks: always 20\n"
        "  - total_marks_awarded (cap at 14 maximum)\n"
        "  - question_statement: the exam question as written by the student\n"
        "  - question_expectation: MUST be an array of 3-5 short, specific bullet points describing what an excellent answer should cover according to the subject rubric. Each bullet should be one clear sentence focusing on key themes, concepts, theories, or historical periods expected.\n"
        "  - criteria[]: each criterion with id, name, max, awarded, strengths[], weaknesses[]\n"
        "  - high_scoring_outline: ALWAYS use title as 'High-Scoring Ideal Outline' and return outline_points you design yourself.\n"
        "      * Provide 4-6 ordered sections that read like a model answer plan: introduction, body sections for each era/theme, comparative analysis, and a conclusion/way forward.\n"
        "      * Each outline_points entry must be an object with: heading (section title), summary (2-3 sentences describing what to cover), and key_points (2-4 concise bullets highlighting evidence, theorists, dates, policies, etc.).\n"
        "      * Summaries and key points must focus on how the student should have answered, not meta commentary.\n"
        "  - overall_comment: 3-5 sentence holistic evaluation.\n"
    )

    content_payload = {
        "subject": subject,
        "rubric_text": subject_rubric_text,
        # Send full OCR text (no truncation)
        "ocr_full_text": ocr_data.get("full_text", ""),
        "sections": sections,
        "page_images_base64_png": page_images,
        "output_schema": schema_hint,
    }

    return {
        "model": "grok-4-fast-reasoning",
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an expert CSS examiner. "
                    "You produce detailed, rubric-based marking reports and respond in JSON only."
                ),
            },
            {
                "role": "user",
                "content": instructions
                + "\n\nDATA:\n"
                + json.dumps(content_payload, ensure_ascii=False),
            },
        ],
        "temperature": 0.15,
        "max_tokens": 8000,  # Increased to allow longer responses
    }


def call_grok_for_grading(
    grok_api_key: str,
    subject: str,
    subject_rubric_text: str,
    ocr_data: Dict[str, Any],
    sections: List[Dict[str, Any]],
    page_images: List[Dict[str, Any]],
    max_retries: int = 3,
) -> Tuple[Dict[str, Any], Dict[str, int]]:
    payload = build_grok_payload_for_grading(
        subject, subject_rubric_text, ocr_data, sections, page_images
    )
    headers = {
        "Authorization": f"Bearer {grok_api_key}",
        "Content-Type": "application/json",
    }

    for attempt in range(max_retries):
        if attempt > 0:
            print(f"Retry attempt {attempt + 1}/{max_retries} for Grok grading...")

        resp = requests.post(
            "https://api.x.ai/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=150,
        )
        if resp.status_code >= 300:
            if attempt < max_retries - 1:
                print(f"API error {resp.status_code}, retrying...")
                continue
            raise RuntimeError(f"Grok grading error {resp.status_code}: {resp.text}")

        try:
            data = resp.json()
        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Response JSON parse error, retrying...")
                continue
            raise RuntimeError(f"Grok grading JSON parse error: {exc}") from exc

        # Extract token usage from response
        usage = data.get("usage", {})
        token_usage = {
            "input_tokens": usage.get("prompt_tokens", 0),
            "output_tokens": usage.get("completion_tokens", 0),
        }

        try:
            content = data["choices"][0]["message"]["content"]
        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Missing content in response, retrying...")
                continue
            raise RuntimeError(f"Unexpected Grok grading response: {data}") from exc

        # Check if response was truncated due to finish_reason
        finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
        if finish_reason == "length":
            print(f"WARNING: Grok response was truncated due to max token limit!")
            print(f"Response length: {len(content)} characters")
            if attempt < max_retries - 1:
                # Increase max_tokens and retry
                payload["max_tokens"] = payload.get("max_tokens", 8000) + 2000
                print(f"Increasing max_tokens to {payload['max_tokens']} and retrying...")
                continue

        try:
            cleaned = clean_json_from_llm(content)
            parsed = json.loads(cleaned)
            # Success! Return the result
            if attempt > 0:
                print(f"Successfully parsed JSON on attempt {attempt + 1}")
            return parsed, token_usage
        except json.JSONDecodeError as exc:
            # Try to repair JSON before giving up
            cleaned = clean_json_from_llm(content)  # Re-clean in case it wasn't set
            try:
                repaired = repair_json(cleaned, error_pos=exc.pos)
                parsed = json.loads(repaired)
                print(f"Successfully parsed JSON after repair on attempt {attempt + 1}")
                return parsed, token_usage
            except (json.JSONDecodeError, Exception) as repair_exc:
                # Save malformed JSON to file for debugging (with unique timestamp)
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                error_file = f"grok_error_response_{timestamp}_{attempt + 1}.txt"
                with open(error_file, "w", encoding="utf-8") as f:
                    f.write(f"=== FULL RESPONSE (length: {len(content)} chars) ===\n")
                    f.write(content)
                    f.write(f"\n\n=== CLEANED (length: {len(cleaned)} chars) ===\n")
                    f.write(cleaned)
                    f.write(f"\n\n=== REPAIRED (length: {len(repaired)} chars) ===\n")
                    f.write(repaired)
                    f.write(f"\n\n=== ORIGINAL ERROR ===\n{exc}\n")
                    f.write(f"\n=== REPAIR ERROR ===\n{repair_exc}\n")

                print(f"\nDEBUG: Full content length: {len(content)} characters")
                print(f"DEBUG: Finish reason: {finish_reason}")
                print(f"DEBUG: JSON parse error at position {exc.pos}: {exc.msg}")
                print(f"DEBUG: JSON repair also failed: {repair_exc}")
                print(f"DEBUG: Saved full response to {error_file}")
                print(f"DEBUG: First 300 chars: {content[:300]}")
                print(f"DEBUG: Last 300 chars: {content[-300:]}")

                if attempt < max_retries - 1:
                    print(f"Malformed JSON, retrying...")
                    continue

                raise RuntimeError(
                    f"Grok grading returned malformed JSON after {max_retries} attempts. "
                    f"Error: {exc.msg} at position {exc.pos}. "
                    f"Repair attempt also failed: {repair_exc}. "
                    f"Full response saved to {error_file}"
                ) from exc

    raise RuntimeError(f"Failed to get valid response after {max_retries} attempts")


# -----------------------------
# GROK CALL 3: REFINED RUBRIC ANNOTATIONS
# -----------------------------


def call_grok_for_refined_rubric_annotations(
    grok_api_key: str,
    refined_rubric_text: str,
    ocr_data: Dict[str, Any],
    sections: List[Dict[str, Any]],
    page_images: List[Dict[str, Any]],
    max_retries: int = 3,
) -> Tuple[Dict[str, Any], Dict[str, int]]:
    """
    Ask Grok to apply the refined generic rubric and produce:
      - annotations[] (for drawing boxes / comments),
      - refined_rubric_summary[] (one entry per rubric point).

    Output shape:

      {
        "annotations": [
          { ... see annotate_pdf_with_rubric.py docstring ... }
        ],
        "refined_rubric_summary": [
          {
            "id": "introduction_quality",
            "name": "Introduction Quality",
            "rating": "weak/average/good/excellent",
            "comment": "..."
          },
          ...
        ]
      }
    """
    system_msg = {
        "role": "system",
        "content": (
            "You are an expert examiner and annotation assistant.\n"
            "You apply a refined generic rubric to handwritten exam answers and output:\n"
            "- concrete annotations (where to draw boxes and what feedback to write), and\n"
            "- a short summary for each rubric point.\n\n"
            "PRIMARY INPUTS YOU RECEIVE:\n"
            "- refined_rubric_text: the generic rubric you must apply;\n"
            "- ocr_full_text + ocr_pages: approximate OCR text and per-page structure;\n"
            "- sections: JSON from a previous step that segments the answer into logical sections;\n"
            "- page_images_base64_png: base64 PNG page images (the handwritten script).\n\n"
            "GROUND TRUTH PRIORITY:\n"
            "- The PAGE IMAGES are the ultimate source of truth.\n"
            "- OCR is only a helper for locating text; if OCR and image disagree, trust the image.\n"
            "- The SECTIONS JSON gives you stable titles and page ranges. Use it whenever you need\n"
            "  to refer to headings or sections (via section_id or target_section_id).\n\n"
            "YOUR GOAL:\n"
            "- Produce a JSON object with:\n"
            "    annotations[]: ALL important issues and observations you can detect,\n"
            "    refined_rubric_summary[]: one item per rubric point with rating + comment.\n"
            "- Be especially thorough for:\n"
            "    * incorrect or weak headings (missing clarity / relevance), and\n"
            "    * factual inaccuracies (wrong dates, wrong names, wrong causal claims, etc.).\n"
            "- It is acceptable to be conservative for spelling/grammar due to OCR noise.\n\n"
            "STRICT OUTPUT FORMAT (IMPORTANT):\n"
            "- Return ONLY valid JSON (no markdown, no commentary).\n"
            "- Top-level keys allowed: 'annotations' and 'refined_rubric_summary'.\n"
            "- annotations[] entries can have additional fields, but MUST include the fields\n"
            "  required in the provided output_schema examples for each type.\n"
            "- The JSON must be parseable by a strict JSON parser.\n"
        ),
    }

    instructions = (
        "Use the refined generic rubric text to evaluate this answer strictly. "
        "You must obey these annotation rules:\n\n"
        "⚠️ ABSOLUTE REQUIREMENT - INTRODUCTION COMMENT:\n"
        "- THE VERY FIRST ANNOTATION IN YOUR OUTPUT MUST BE type='introduction_comment'\n"
        "- THIS IS NON-NEGOTIABLE - EVERY ANSWER MUST HAVE EXACTLY ONE INTRODUCTION COMMENT\n"
        "- PLACE IT AS THE FIRST ITEM IN THE annotations[] ARRAY\n"
        "- IF YOU DO NOT INCLUDE THIS, YOUR RESPONSE WILL BE REJECTED\n\n"
        "CRITICAL REQUIREMENT FOR ALL ANNOTATIONS:\n"
        "- For ALL annotation types, you MUST copy text EXACTLY as it appears in the OCR text.\n"
        "- NEVER paraphrase, reword, or correct the text when filling target_word_or_sentence.\n"
        "- ALWAYS provide context_before (3-5 words immediately before) and context_after (3-5 words immediately after).\n"
        "- Copy these contexts EXACTLY from the OCR text with original spelling and punctuation.\n\n"
        "ALL ANNOTATIONS MUST HAVE THIS UNIFIED SCHEMA:\n"
        "  type: string (introduction_comment/heading_issue/factual_error/grammar_language/repetition)\n"
        "  rubric_point: string (e.g., 'introduction_quality', 'headings_subheadings', 'factual_accuracy', 'grammar_language')\n"
        "  page: integer (page number where the annotation appears)\n"
        "  target_word_or_sentence: string (EXACT text from OCR - the word, phrase, or sentence being annotated)\n"
        "  context_before: string (EXACT 3-5 words from OCR that appear immediately before the target)\n"
        "  context_after: string (EXACT 3-5 words from OCR that appear immediately after the target)\n"
        "  correction: string (the correct version, or suggestion for improvement)\n"
        "  comment: string (explanation of the issue)\n"
        "  sentiment: string (optional, for heading_issue: 'positive' or 'negative')\n\n"
        "1) Introduction:\n"
        "   - MANDATORY: You MUST ALWAYS create exactly ONE annotation of type 'introduction_comment'.\n"
        "   - This annotation is REQUIRED for every answer, regardless of quality.\n"
        "   - Decide if introduction is weak/average/good/excellent and be strict about it.\n"
        "   - Create type 'introduction_comment' with:\n"
        "       rubric_point = 'introduction_quality',\n"
        "       page = first page where introduction appears,\n"
        "       target_word_or_sentence = EXACT first sentence or opening phrase from OCR,\n"
        "       context_before = '' (empty for first sentence),\n"
        "       context_after = EXACT next 3-5 words from OCR after the target,\n"
        "       correction = '' (not applicable for introduction comments),\n"
        "       comment = a detailed 3–5 sentence evaluation of ONLY the introduction using the refined generic rubric.\n"
        "   - DO NOT SKIP THIS ANNOTATION. It is MANDATORY.\n\n"
        "2) Headings and subheadings:\n"
        "   - MANDATORY: You MUST evaluate EVERY SINGLE heading and subheading detected in the sections[] array.\n"
        "   - For EACH heading/subheading found, you MUST create ONE 'heading_issue' annotation.\n"
        "   - DO NOT skip any headings. If you detect 5 headings, you MUST create 5 heading_issue annotations.\n"
        "   - IMPORTANT: DO NOT evaluate spelling, grammar, or OCR errors in headings. These are handled separately.\n"
        "   - ONLY evaluate heading CONTENT: relevance, clarity, self-explanatory nature.\n"
        "   - For CORRECT headings (self-explanatory, relevant, clear), add annotation type 'heading_issue' with:\n"
        "       rubric_point = 'headings_subheadings',\n"
        "       page = page number of that heading,\n"
        "       target_word_or_sentence = EXACT heading text from OCR,\n"
        "       context_before = EXACT 3-5 words from OCR before the heading,\n"
        "       context_after = EXACT 3-5 words from OCR after the heading,\n"
        "       sentiment = 'positive',\n"
        "       correction = '' (empty for positive headings),\n"
        "       comment = 'Correct heading' or 'Good heading - clear and relevant'.\n"
        "   - For INCORRECT/PROBLEMATIC headings, create type 'heading_issue' with:\n"
        "       rubric_point = 'headings_subheadings',\n"
        "       page = page number of that heading,\n"
        "       target_word_or_sentence = EXACT heading text from OCR (even if misspelled),\n"
        "       context_before = EXACT 3-5 words from OCR before the heading,\n"
        "       context_after = EXACT 3-5 words from OCR after the heading,\n"
        "       sentiment = 'negative',\n"
        "       correction = a better alternate heading that would be more self-explanatory and relevant,\n"
        "       comment = short explanation of the issue (NEVER mention spelling/grammar/OCR errors).\n"
        "   - Focus ONLY on: not being self-explanatory, not directly relevant, vague, unclear, irrelevant.\n"
        "   - IGNORE: spelling mistakes, grammar errors, OCR misreads in headings.\n\n"
        "3) Factual inaccuracies:\n"
        "   - CRITICAL: Do NOT create annotations for CORRECT facts. ONLY annotate ACTUAL ERRORS.\n"
        "   - Do NOT mark spelling mistakes as factual errors.\n"
        "   - If a date/fact is correct, DO NOT create any annotation for it.\n"
        "   - Only flag actual factual mistakes (wrong dates, wrong facts, incorrect information).\n"
        "   - Keep target_word_or_sentence VERY SHORT (1-10 words max) containing only the error.\n"
        "   - For each ACTUAL factual mistake, create type 'factual_error' with:\n"
        "       rubric_point = 'factual_accuracy',\n"
        "       page = page where the error appears,\n"
        "       target_word_or_sentence = EXACT SHORT PHRASE containing the WRONG fact (e.g., '1944' when it should be '1945'),\n"
        "       context_before = EXACT 3-5 words immediately before the error in OCR,\n"
        "       context_after = EXACT 3-5 words immediately after the error in OCR,\n"
        "       correction = the CORRECT fact (e.g., '1945'),\n"
        "       comment = short explanation (e.g., 'Year should be 1945 not 1944').\n"
        "   - NEVER copy full sentences - only the specific phrase with the error.\n"
        "   - DO NOT create factual_error if target_word_or_sentence = correction (that means it's correct!)\n"
        "   - Examples:\n"
        "       * WRONG: target='1944', correction='1944', comment='correct' (DO NOT DO THIS!)\n"
        "       * WRONG: target='1707', correction='1707', comment='Year is correct' (DO NOT DO THIS!)\n"
        "       * CORRECT: target='1944', correction='1945', comment='Year should be 1945 not 1944'\n"
        "       * CORRECT: target='World War I', correction='World War II', comment='Should be WWII not WWI'\n\n"
        "4) Spelling only (no grammar):\n"
        "   - Focus ONLY on clear spelling mistakes (wrongly spelled words).\n"
        "   - Do NOT correct grammar, sentence structure, style, or phrasing.\n"
        "   - For each spelling issue, create type 'grammar_language' with:\n"
        "       rubric_point = 'grammar_language',\n"
        "       page = page where the misspelled word appears,\n"
        "       target_word_or_sentence = EXACT misspelled word or very short span (1-3 words) from OCR,\n"
        "       context_before = EXACT 3-5 words from OCR immediately before the misspelled word,\n"
        "       context_after = EXACT 3-5 words from OCR immediately after the misspelled word,\n"
        "       correction = the correctly spelled word or very short corrected phrase,\n"
        "       comment = brief note like 'spelling error'.\n"
        "   - Always cross-check using BOTH OCR text AND the page image:\n"
        "       * Use ocr_full_text to locate the word,\n"
        "       * Then visually verify the spelling directly on the page image.\n"
        "       * If OCR and the image disagree, TRUST THE IMAGE and do NOT flag a spelling error.\n"
        "   - Do not send entire paragraphs or long sentences as target_word_or_sentence.\n"
        "   - If the same misspelling occurs multiple times on the same page, create a SEPARATE annotation for EACH occurrence,\n"
        "     with different context_before and context_after for each instance.\n\n"
        "5) Repetition:\n"
        "   - If content is repeated across pages, create type 'repetition' with:\n"
        "       rubric_point = 'repetitiveness',\n"
        "       page = the page where the repeated content appears again,\n"
        "       target_word_or_sentence = EXACT repeated phrase or sentence from OCR,\n"
        "       context_before = EXACT 3-5 words from OCR before the repeated text,\n"
        "       context_after = EXACT 3-5 words from OCR after the repeated text,\n"
        "       correction = suggestion like 'Remove repetition' or 'Already mentioned on page X',\n"
        "       comment = note indicating where it was first mentioned.\n\n"
        "Additionally, build refined_rubric_summary[]:\n"
        "   - ONLY include these 4 rubric points:\n"
        "     1. argumentation_quality (name: 'Argumentation Quality')\n"
        "     2. presentation (name: 'Presentation Quality')\n"
        "     3. contemporary_relevance (name: 'Contemporary Relevance')\n"
        "     4. length_completeness (name: 'Length & Completeness')\n"
        "   - Each entry: id, name, rating (weak/average/good/excellent), comment (1 sentence max, very concise and brief).\n"
    )


    # Sanitize OCR pages (remove bounding boxes) and send full OCR text
    raw_pages = ocr_data.get("pages", [])
    sanitized_pages = []
    for p in raw_pages:
        lines = []
        for line in p.get("lines", []):
            line_text = line.get("text", "")
            words = [w.get("text", "") for w in line.get("words", [])]
            lines.append({"text": line_text, "words": words})
        sanitized_pages.append({"page_number": p.get("page_number"), "lines": lines})

    user_payload = {
        "refined_rubric_text": refined_rubric_text,
        "ocr_full_text": ocr_data.get("full_text", ""),
        "ocr_pages": sanitized_pages,
        "sections": sections,
        "page_images_base64_png": page_images,
        "output_schema": {
            "annotations": [
                {
                    "type": "introduction_comment",
                    "rubric_point": "introduction_quality",
                    "page": 1,
                    "target_word_or_sentence": "First sentence of introduction (EXACT from OCR)",
                    "context_before": "",
                    "context_after": "Next 3-5 words after (EXACT from OCR)",
                    "correction": "",
                    "comment": "Detailed evaluation of introduction quality (3-5 sentences)",
                },
                {
                    "type": "heading_issue",
                    "rubric_point": "headings_subheadings",
                    "page": 1,
                    "target_word_or_sentence": "EXACT heading text from OCR",
                    "context_before": "EXACT 3-5 words before",
                    "context_after": "EXACT 3-5 words after",
                    "correction": "Better heading suggestion or empty if positive",
                    "comment": "Explanation of issue or 'Correct heading'",
                    "sentiment": "positive/negative",
                },
                {
                    "type": "factual_error",
                    "rubric_point": "factual_accuracy",
                    "page": 2,
                    "target_word_or_sentence": "SHORT PHRASE with error (EXACT from OCR)",
                    "context_before": "EXACT 3-5 words before",
                    "context_after": "EXACT 3-5 words after",
                    "correction": "Correct fact",
                    "comment": "Explanation of error",
                }
            ],
            "refined_rubric_summary": [
                {
                    "id": "argumentation_quality",
                    "name": "Argumentation Quality",
                    "rating": "weak/average/good/excellent",
                    "comment": "string",
                },
                {
                    "id": "presentation",
                    "name": "Presentation Quality",
                    "rating": "weak/average/good/excellent",
                    "comment": "string",
                },
                {
                    "id": "contemporary_relevance",
                    "name": "Contemporary Relevance",
                    "rating": "weak/average/good/excellent",
                    "comment": "string",
                },
                {
                    "id": "length_completeness",
                    "name": "Length & Completeness",
                    "rating": "weak/average/good/excellent",
                    "comment": "string",
                }
            ],
        },
    }

    user_msg = {
        "role": "user",
        "content": instructions + "\n\nDATA:\n" + json.dumps(user_payload, ensure_ascii=False),
    }

    headers = {
        "Authorization": f"Bearer {grok_api_key}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": "grok-4-fast-reasoning",
        "messages": [system_msg, user_msg],
        "temperature": 0.1,
        "max_tokens": 6000,  # Increased for refined rubric annotations
    }

    for attempt in range(max_retries):
        if attempt > 0:
            print(f"Retry attempt {attempt + 1}/{max_retries} for refined rubric annotations...")

        resp = requests.post(
            "https://api.x.ai/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=150,
        )

        if resp.status_code >= 300:
            if attempt < max_retries - 1:
                print(f"API error {resp.status_code}, retrying...")
                continue
            # Fallback on final failure
            return {
                "annotations": [
                    {
                        "type": "introduction_comment",
                        "rubric_point": "introduction_quality",
                        "page": 1,
                        "target_word_or_sentence": "",
                        "context_before": "",
                        "context_after": "",
                        "correction": "",
                        "comment": f"Error: API returned {resp.status_code}",
                    }
                ],
                "refined_rubric_summary": [],
            }, {"input_tokens": 0, "output_tokens": 0}

        try:
            data = resp.json()
        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Response JSON parse error, retrying...")
                continue
            # Fallback on final failure
            return {
                "annotations": [
                    {
                        "type": "introduction_comment",
                        "rubric_point": "introduction_quality",
                        "page": 1,
                        "target_word_or_sentence": "",
                        "context_before": "",
                        "context_after": "",
                        "correction": "",
                        "comment": f"Error parsing API response: {exc}",
                    }
                ],
                "refined_rubric_summary": [],
            }, {"input_tokens": 0, "output_tokens": 0}

        # Extract token usage from response
        usage = data.get("usage", {})
        token_usage = {
            "input_tokens": usage.get("prompt_tokens", 0),
            "output_tokens": usage.get("completion_tokens", 0),
        }

        try:
            content = data["choices"][0]["message"]["content"]
        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Missing content in response, retrying...")
                continue
            # Fallback on final failure
            return {
                "annotations": [
                    {
                        "type": "introduction_comment",
                        "rubric_point": "introduction_quality",
                        "page": 1,
                        "target_word_or_sentence": "",
                        "context_before": "",
                        "context_after": "",
                        "correction": "",
                        "comment": f"Error: Unexpected API response structure",
                    }
                ],
                "refined_rubric_summary": [],
            }, token_usage

        # Check if truncated
        finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
        if finish_reason == "length":
            print(f"WARNING: Refined rubric response truncated!")
            if attempt < max_retries - 1:
                payload["max_tokens"] = payload.get("max_tokens", 6000) + 2000
                print(f"Increasing max_tokens to {payload['max_tokens']} and retrying...")
                continue

        try:
            cleaned = clean_json_from_llm(content)
            parsed = json.loads(cleaned)
        except json.JSONDecodeError as exc:
            # Try to repair JSON before giving up
            cleaned = clean_json_from_llm(content)  # Re-clean in case it wasn't set
            try:
                repaired = repair_json(cleaned, error_pos=exc.pos)
                parsed = json.loads(repaired)
                print(f"Successfully parsed refined rubric JSON after repair on attempt {attempt + 1}")
            except (json.JSONDecodeError, Exception) as repair_exc:
                # Save error for debugging (with unique timestamp)
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                error_file = f"grok_refined_error_{timestamp}_{attempt + 1}.txt"
                with open(error_file, "w", encoding="utf-8") as f:
                    f.write(f"=== FULL RESPONSE (length: {len(content)} chars) ===\n")
                    f.write(content)
                    f.write(f"\n\n=== CLEANED (length: {len(cleaned)} chars) ===\n")
                    f.write(cleaned)
                    f.write(f"\n\n=== REPAIRED (length: {len(repaired)} chars) ===\n")
                    f.write(repaired)
                    f.write(f"\n\n=== ORIGINAL ERROR ===\n{exc}\n")
                    f.write(f"\n=== REPAIR ERROR ===\n{repair_exc}\n")

                print(f"\nDEBUG: JSON parse error at position {exc.pos}: {exc.msg}")
                print(f"DEBUG: JSON repair also failed: {repair_exc}")
                print(f"DEBUG: Saved full response to {error_file}")

                if attempt < max_retries - 1:
                    print(f"Malformed JSON, retrying...")
                    continue

            # Fallback on final failure
            return {
                "annotations": [
                    {
                        "type": "introduction_comment",
                        "rubric_point": "introduction_quality",
                        "page": 1,
                        "target_word_or_sentence": "",
                        "context_before": "",
                        "context_after": "",
                        "correction": "",
                        "comment": f"Error parsing refined rubric response: {exc}",
                    }
                ],
                "refined_rubric_summary": [],
            }, token_usage

        # Success! Normalize and validate
        parsed.setdefault("annotations", [])
        parsed.setdefault("refined_rubric_summary", [])

        # ENFORCE: Ensure introduction_comment always exists as first annotation
        annotations = parsed.get("annotations", [])
        has_intro = any(a.get("type") == "introduction_comment" for a in annotations)

        if not has_intro:
            # Inject introduction_comment as first annotation
            intro_text = ocr_data.get("full_text", "")[:200].split('\n')[0] if ocr_data.get("full_text") else "Introduction"
            intro_annotation = {
                "type": "introduction_comment",
                "rubric_point": "introduction_quality",
                "page": 1,
                "target_word_or_sentence": intro_text,
                "context_before": "",
                "context_after": "",
                "correction": "",
                "comment": "Introduction evaluation (auto-generated due to missing annotation)"
            }
            parsed["annotations"].insert(0, intro_annotation)
            print("⚠️  Warning: Introduction comment was missing. Auto-injected.")

        if attempt > 0:
            print(f"Successfully parsed refined rubric annotations on attempt {attempt + 1}")

        return parsed, token_usage

    # Should never reach here
    return {
        "annotations": [
            {
                "type": "introduction_comment",
                "rubric_point": "introduction_quality",
                "page": 1,
                "target_word_or_sentence": "",
                "context_before": "",
                "context_after": "",
                "correction": "",
                "comment": f"Failed after {max_retries} attempts",
            }
        ],
        "refined_rubric_summary": [],
    }, {"input_tokens": 0, "output_tokens": 0}


# -----------------------------
# GROK CALL 4: PAGE-WISE IMPROVEMENT SUGGESTIONS
# -----------------------------


def call_grok_for_page_wise_suggestions(
    grok_api_key: str,
    subject: str,
    subject_rubric_text: str,
    ocr_data: Dict[str, Any],
    sections: List[Dict[str, Any]],) -> Tuple[Dict[str, Any], Dict[str, int]]:
    """
    Ask Grok to provide specific, actionable improvement suggestions for each page.

    Output shape:
      {
        "page_suggestions": [
          {
            "page": 1,
            "suggestions": [
              "Add comparison with Kant's categorical imperative",
              "Include the 1945 UN Charter establishment date",
              "Reference Rawls' theory of justice for stronger argumentation",
              "Add empirical evidence from 2020 Climate Summit"
            ]
          },
          ...
        ]
      }
    """
    system_msg = {
        "role": "system",
        "content": (
            "You are an expert CSS examiner focused on helping students improve their answers.\n"
            "You receive:\n"
            "- The subject rubric (detailed criteria)\n"
            "- OCR text from student's answer (page-wise)\n"
            "- Section structure (headings detected)\n\n"
            "YOUR GOAL:\n"
            "For each page, provide 3-6 specific, actionable suggestions for improvement.\n"
            "Focus on VALUE ADDITIONS that would strengthen the answer.\n\n"
            "TYPES OF SUGGESTIONS TO PROVIDE:\n"
            "1. Theoretical additions: 'Add comparison with X philosopher/theorist'\n"
            "2. Factual additions: 'Include the date: [specific event] occurred in [year]'\n"
            "3. Evidence additions: 'Add empirical data from [specific study/report]'\n"
            "4. Comparative analysis: 'Compare with [country/era/policy]'\n"
            "5. Critical perspectives: 'Include critique from [scholar/school of thought]'\n"
            "6. Contemporary relevance: 'Link to recent event: [specific event in year]'\n\n"
            "STRICT OUTPUT FORMAT:\n"
            "- Return ONLY valid JSON\n"
            "- Top-level key: 'page_suggestions' with array of page objects\n"
            "- Each page object has: 'page' (integer) and 'suggestions' (array of strings)\n"
            "- Each suggestion must be a single, specific, actionable statement\n"
            "- No markdown, no commentary, just JSON\n"
        ),
    }

    instructions = (
        "Analyze this student's answer page by page.\n"
        "For each page, identify 3-6 specific additions that would improve the answer quality.\n\n"
        "REQUIREMENTS:\n"
        "1. Suggestions must be SPECIFIC, not vague\n"
        "   ❌ Bad: 'Add more theories'\n"
        "   ✅ Good: 'Add Foucault's concept of biopower (1976)'\n\n"
        "2. Focus on what to ADD, not what's wrong\n"
        "   ❌ Bad: 'This argument is weak'\n"
        "   ✅ Good: 'Strengthen argument by adding Weber's bureaucracy theory'\n\n"
        "3. Include specific names, dates, events, theories\n"
        "   ❌ Bad: 'Reference a philosopher'\n"
        "   ✅ Good: 'Reference Mill's harm principle (On Liberty, 1859)'\n\n"
        "4. Suggestions should align with the subject rubric criteria\n\n"
        "5. Each page should have 3-6 suggestions maximum\n\n"
        "OUTPUT:\n"
        "Return JSON with this exact structure:\n"
        "{\n"
        "  \"page_suggestions\": [\n"
        "    {\n"
        "      \"page\": 1,\n"
        "      \"suggestions\": [\n"
        "        \"Add comparison with Locke's social contract theory (Two Treatises, 1689)\",\n"
        "        \"Include the Magna Carta signing date (1215) as historical precedent\",\n"
        "        \"Reference the 1948 Universal Declaration of Human Rights\"\n"
        "      ]\n"
        "    },\n"
        "    {\n"
        "      \"page\": 2,\n"
        "      \"suggestions\": [\n"
        "        \"Add Amartya Sen's capability approach for economic analysis\",\n"
        "        \"Include World Bank poverty data (2021 report)\",\n"
        "        \"Compare with China's economic reforms post-1978\"\n"
        "      ]\n"
        "    }\n"
        "  ]\n"
        "}\n"
    )

    # Simplified page data (text only, no images or bounding boxes to reduce tokens)
    ocr_pages_minimal = [
        {
            "page": p.get("page", idx + 1),
            "text": p.get("text", "")[:300000]  # Limit per-page text
        }
        for idx, p in enumerate(ocr_data.get("pages", []))
    ]

    user_payload = {
        "subject": subject,
        "rubric_text": subject_rubric_text,
        "ocr_full_text": ocr_data.get("full_text", "")[:15000],
        "ocr_pages": ocr_pages_minimal,
        "sections": sections,
        "output_schema": {
            "page_suggestions": [
                {
                    "page": 1,
                    "suggestions": ["string", "string"]
                }
            ]
        },
    }

    user_msg = {
        "role": "user",
        "content": instructions + "\n\nDATA:\n" + json.dumps(user_payload, ensure_ascii=False),
    }

    headers = {
        "Authorization": f"Bearer {grok_api_key}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": "grok-4-fast-reasoning",
        "messages": [system_msg, user_msg],
        "temperature": 0.2,
        "max_tokens": 4000,  # Sufficient for page suggestions
    }

    max_retries = 3
    for attempt in range(max_retries):
        if attempt > 0:
            # Exponential backoff: 1s, 2s, 4s
            import time
            time.sleep(2 ** (attempt - 1))
            print(f"Retry attempt {attempt + 1}/{max_retries} for page suggestions...")

        try:
            resp = requests.post(
                "https://api.x.ai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=150,
            )

            if resp.status_code >= 300:
                if attempt < max_retries - 1:
                    print(f"API error {resp.status_code}, retrying...")
                    continue
                # Fallback on final attempt
                return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}

            data = resp.json()

            # Extract token usage
            usage = data.get("usage", {})
            token_usage = {
                "input_tokens": usage.get("prompt_tokens", 0),
                "output_tokens": usage.get("completion_tokens", 0),
            }

            # Check for truncation
            finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
            if finish_reason == "length" and attempt < max_retries - 1:
                payload["max_tokens"] = payload.get("max_tokens", 4000) + 2000
                print(f"Response truncated, increasing max_tokens to {payload['max_tokens']} and retrying...")
                continue

            content = data["choices"][0]["message"]["content"]
            cleaned = clean_json_from_llm(content)
            parsed = json.loads(cleaned)

            # Success
            if attempt > 0:
                print(f"Successfully parsed page suggestions on attempt {attempt + 1}")

            return parsed, token_usage

        except (requests.Timeout, requests.ConnectionError) as e:
            if attempt < max_retries - 1:
                print(f"Network error, retrying...")
                continue
            # Fallback on final attempt
            return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}

        except json.JSONDecodeError as e:
            if attempt < max_retries - 1:
                print(f"JSON parse error, retrying...")
                continue
            # Fallback on final attempt
            return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}

        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Unexpected error: {exc}, retrying...")
                continue
            # Fallback on final attempt
            return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}

    # Should never reach here, but fallback just in case
    return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}


# -----------------------------
# REPORT RENDERING (SUBJECT MARKING)
# -----------------------------


_MODULE_DIR = os.path.dirname(os.path.abspath(__file__))
_FONTS_DIR = os.path.join(_MODULE_DIR, "fonts")
_FONT_CANDIDATES = [
    os.environ.get("OCR_FONT_PATH"),
    os.path.join(_FONTS_DIR, "ReportFont.ttf"),
    os.path.join(_FONTS_DIR, "NotoSans-Regular.ttf"),
    os.path.join(_FONTS_DIR, "DejaVuSans.ttf"),
    "arial.ttf",
    "Arial.ttf",
    "LiberationSans-Regular.ttf",
    "DejaVuSans.ttf",
]


def _iter_font_candidates() -> List[str]:
    seen: Set[str] = set()
    ordered: List[str] = []
    for candidate in _FONT_CANDIDATES:
        if not candidate:
            continue
        norm = os.path.normpath(candidate) if os.path.isabs(candidate) else candidate
        if norm in seen:
            continue
        seen.add(norm)
        ordered.append(candidate)
    return ordered


@lru_cache(maxsize=32)
def _get_font(size: int) -> ImageFont.FreeTypeFont:
    for candidate in _iter_font_candidates():
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _extract_expectation_bullets(expectation_text: Any) -> List[str]:
    bullet_points: List[str] = []
    if isinstance(expectation_text, list):
        bullet_points = [str(p).strip() for p in expectation_text if p]
    elif expectation_text:
        expectation_str = str(expectation_text).strip()
        sentences = re.split(r"(?<=[.!?])\s+", expectation_str)
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence:
                if sentence.endswith((".", "!", "?")):
                    sentence = sentence[:-1].strip()
                bullet_points.append(sentence)
        if not bullet_points:
            comma_parts = [p.strip() for p in expectation_str.split(",")]
            bullet_points = [p for p in comma_parts if p]
    if not bullet_points:
        bullet_points = ["No specific expectations provided"]
    bullet_points = [pt for pt in bullet_points if pt]
    return bullet_points[:4]


def _wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.FreeTypeFont,
    max_width: int,
) -> List[str]:
    words = text.split()
    lines: List[str] = []
    current = ""
    for w in words:
        trial = (current + " " + w).strip()
        if not trial:
            continue
        width = draw.textlength(trial, font=font)
        if width <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines

def render_subject_report_pages(
    grading_result: Dict[str, Any],
    page_size: Tuple[int, int] = (2977, 4211),  # 200 DPI: Width x Height
) -> List[Image.Image]:
    """
    Render the subject-wise marking report on exactly 2 pages max.
    If content exceeds 2 pages, font sizes are reduced until it fits.
      - SUBJECT NAME (at top)
      - TOTAL MARKS
      - QUESTION STATEMENT
      - WHAT QUESTION EXPECTS (right after question statement, as single paragraph)
      - CRITERIA breakdown (with strengths & weaknesses)
      (Note: high-scoring outline is on a separate dedicated page)
    """
    # Try rendering with progressively smaller fonts until it fits in 2 pages
    font_scale = 1.0
    max_attempts = 10

    for attempt in range(max_attempts):
        pages = _render_subject_report_with_scale(grading_result, page_size, font_scale)

        if len(pages) <= 2:
            # Success! Fits in 2 pages or less
            if attempt > 0:
                print(f"Subject report fit in {len(pages)} pages after reducing font to {font_scale:.1%} of original size")
            return pages

        # Too many pages, reduce font size
        print(f"Subject report has {len(pages)} pages (attempt {attempt+1}), reducing font size...")
        font_scale *= 0.85  # Reduce by 15% each iteration

    # If we still can't fit after max attempts, return what we have
    print(f"WARNING: Subject report still has {len(pages)} pages after {max_attempts} attempts to reduce font size")
    return pages


def _render_subject_report_with_scale(
    grading_result: Dict[str, Any],
    page_size: Tuple[int, int],
    font_scale: float = 1.0,
) -> List[Image.Image]:
    """
    Internal helper to render subject report with a given font scale.
    """
    W, H = page_size
    margin = int(W * 0.07)
    line_spacing = 1.4

    # Base font sizes (scaled by font_scale parameter)
    title_font = _get_font(int(90 * font_scale))
    subject_font = _get_font(int(74 * font_scale))
    section_heading_font = _get_font(int(78 * font_scale))
    question_text_font = _get_font(int(62 * font_scale))
    criteria_heading_font = _get_font(int(70 * font_scale))
    body_font = _get_font(int(58 * font_scale))
    total_heading_font = _get_font(int(106 * font_scale))
    total_marks_font = _get_font(int(140 * font_scale))

    pages: List[Image.Image] = []

    def new_page() -> Tuple[Image.Image, ImageDraw.ImageDraw, int]:
        img = Image.new("RGB", (W, H), "white")
        draw = ImageDraw.Draw(img)
        return img, draw, margin

    img, draw, y = new_page()

    def ensure_space(font_obj: ImageFont.FreeTypeFont, needed_lines: int):
        nonlocal img, draw, y
        line_h = font_obj.getbbox("Ag")[3] - font_obj.getbbox("Ag")[1]
        if y + line_h * needed_lines * line_spacing > H - margin:
            pages.append(img)
            img, draw, y = new_page()

    def draw_bold_text(
        text: str,
        font_obj: ImageFont.FreeTypeFont,
        position: Tuple[int, int],
        fill: str,
    ) -> None:
        offsets = [(0, 0), (1, 0), (0, 1), (1, 1)]
        for dx, dy in offsets:
            draw.text((position[0] + dx, position[1] + dy), text, font=font_obj, fill=fill)

    max_text_width = W - 2 * margin

    # SUBJECT NAME (at top)
    subject = grading_result.get("subject", "")
    if subject:
        ensure_space(subject_font, 2)
        draw.text((margin, y), f"Subject: {subject}", font=subject_font, fill="black")
        line_h_subj = subject_font.getbbox("Ag")[3] - subject_font.getbbox("Ag")[1]
        y += int(line_h_subj * line_spacing * 1.5)

    # TOTAL MARKS
    total = grading_result.get("total_marks_awarded", 0)
    maximum = grading_result.get("max_marks", 20)

    ensure_space(total_heading_font, 2)
    draw_bold_text("TOTAL MARKS", total_heading_font, (margin, y), "black")
    line_h = total_heading_font.getbbox("Ag")[3] - total_heading_font.getbbox("Ag")[1]
    y += int(line_h * line_spacing)
    draw_bold_text(f"{total} / {maximum}", total_marks_font, (margin, y), "#B22222")  # Keep red color
    line_h2 = total_marks_font.getbbox("Ag")[3] - total_marks_font.getbbox("Ag")[1]
    y += int(line_h2 * line_spacing * 2)

    # QUESTION STATEMENT
    question = grading_result.get("question_statement", "")
    ensure_space(section_heading_font, 3)
    # Center the heading
    heading_text = "QUESTION STATEMENT"
    heading_width = draw.textlength(heading_text, font=section_heading_font)
    heading_x = (W - heading_width) // 2
    draw_bold_text(heading_text, section_heading_font, (heading_x, y), "black")
    line_h_section = section_heading_font.getbbox("Ag")[3] - section_heading_font.getbbox("Ag")[1]
    y += int(line_h_section * 1.5)  # Increased spacing between heading and question text
    for line in _wrap_text(draw, question, question_text_font, max_text_width):
        ensure_space(question_text_font, 1)
        line_hq = question_text_font.getbbox("Ag")[3] - question_text_font.getbbox("Ag")[1]
        draw_bold_text(line, question_text_font, (margin, y), "black")  # Bold question text
        y += int(line_hq * line_spacing)
    y += int(question_text_font.getbbox("Ag")[3] - question_text_font.getbbox("Ag")[1])

    # WHAT QUESTION EXPECTS (right after question statement, as a single paragraph)
    expectation = grading_result.get("question_expectation", "")
    expectation_bullets = _extract_expectation_bullets(expectation) if expectation else []
    if expectation_bullets:
        ensure_space(section_heading_font, 2)
        # Center the heading
        heading_text = "WHAT THE QUESTION EXPECTS"
        heading_width = draw.textlength(heading_text, font=section_heading_font)
        heading_x = (W - heading_width) // 2
        draw_bold_text(heading_text, section_heading_font, (heading_x, y), "black")
        line_h_section = section_heading_font.getbbox("Ag")[3] - section_heading_font.getbbox("Ag")[1]
        y += int(line_h_section * line_spacing)
        line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
        for bullet in expectation_bullets:
            wrapped = _wrap_text(draw, bullet, body_font, max_text_width - int(0.05 * W))
            ensure_space(body_font, len(wrapped) + 1)
            for idx, line in enumerate(wrapped):
                if idx == 0:
                    draw.text((margin, y), f"- {line}", font=body_font, fill="black")
                else:
                    indent_x = margin + int(0.04 * W)
                    draw.text((indent_x, y), line, font=body_font, fill="black")
                y += int(line_hb * line_spacing)
            y += int(line_hb * 0.5)
        y += int(line_hb)

        # CRITERIA BREAKDOWN
    ensure_space(section_heading_font, 2)
    # Center the heading
    heading_text = "MARKS BREAKDOWN"
    heading_width = draw.textlength(heading_text, font=section_heading_font)
    heading_x = (W - heading_width) // 2
    draw_bold_text(heading_text, section_heading_font, (heading_x, y), "black")
    line_h_section = section_heading_font.getbbox("Ag")[3] - section_heading_font.getbbox("Ag")[1]
    y += int(line_h_section * line_spacing * 1.5)

    for crit in grading_result.get("criteria", []):
        name = crit.get("name", "")
        awarded = crit.get("awarded", 0)
        max_marks = crit.get("max", 0)
        header = f"{name}  —  {awarded}/{max_marks}"

        ensure_space(criteria_heading_font, 2)
        line_h_criteria = criteria_heading_font.getbbox("Ag")[3] - criteria_heading_font.getbbox("Ag")[1]
        draw_bold_text(header, criteria_heading_font, (margin, y), "black")
        y += int(line_h_criteria * line_spacing)

        # Strengths
        strengths = crit.get("strengths") or []
        if strengths:
            ensure_space(body_font, len(strengths) + 1)
            line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
            draw.text((margin, y), "Strengths:", font=body_font, fill="darkgreen")
            y += int(line_hb * line_spacing)
            for s in strengths:
                for line in _wrap_text(draw, f"• {s}", body_font, max_text_width):
                    ensure_space(body_font, 1)
                    draw.text((margin, y), line, font=body_font, fill="black")
                    y += int(line_hb * line_spacing)

        # Weaknesses
        weaknesses = crit.get("weaknesses") or []
        if weaknesses:
            ensure_space(body_font, len(weaknesses) + 1)
            line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
            draw.text((margin, y), "Weaknesses:", font=body_font, fill="darkred")
            y += int(line_hb * line_spacing)
            for wtxt in weaknesses:
                for line in _wrap_text(draw, f"• {wtxt}", body_font, max_text_width):
                    ensure_space(body_font, 1)
                    draw.text((margin, y), line, font=body_font, fill="black")
                    y += int(line_hb * line_spacing)

        y += int(body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1])

    pages.append(img)
    return pages




# -----------------------------------------
# PAGE: WHAT THE QUESTION EXPECTS (BULLETS)
# -----------------------------------------

def render_question_expectations_page(
    expectation_text: Any,
    page_size: Tuple[int, int] = (2480, 3508),
) -> Image.Image:
    """
    Render a dedicated page for "What the Question Expects" broken into bullets.

    The expectation_text can be a string or list; it will be converted to bullet points.
    """
    W, H = page_size
    margin = int(W * 0.07)
    line_spacing = 1.4

    title_font = _get_font(72)  # Increased from 60
    body_font = _get_font(48)  # Increased from 36

    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    y = margin

    max_text_width = W - 2 * margin

    # Title
    draw.text((margin, y), "WHAT THE QUESTION EXPECTS", font=title_font, fill="black")
    line_h_title = title_font.getbbox("Ag")[3] - title_font.getbbox("Ag")[1]
    y += int(line_h_title * line_spacing * 1.5)

    bullet_points = _extract_expectation_bullets(expectation_text)
    line_h_body = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]

    for point in bullet_points:
        wrapped_lines = _wrap_text(draw, point, body_font, max_text_width - int(0.08 * W))
        needed_lines = len(wrapped_lines)
        if y + line_h_body * needed_lines * line_spacing > H - margin:
            break

        for idx, line in enumerate(wrapped_lines):
            if idx == 0:
                draw.text((margin, y), f"- {line}", font=body_font, fill="black")
            else:
                draw.text((margin + int(0.04 * W), y), line, font=body_font, fill="black")
            y += int(line_h_body * line_spacing)

        y += int(line_h_body * 0.5)

    return img

# -----------------------------
# REFINED RUBRIC SUMMARY PAGE
# -----------------------------


def render_refined_rubric_summary_page(
    refined_summary: List[Dict[str, Any]],
    page_size: Tuple[int, int] = (2480, 3508),
) -> Image.Image:
    """
    Render a single page summarizing each refined rubric point:

      - Introduction Quality – rating + comment
      - Headings & Subheadings – rating + comment
      - etc.
    """
    W, H = page_size
    margin = int(W * 0.07)
    line_spacing = 1.4

    title_font = _get_font(72)
    h2_font = _get_font(56)
    body_font = _get_font(44)

    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    y = margin

    max_text_width = W - 2 * margin

    # Title
    draw.text((margin, y), "Refined Rubric Feedback", font=title_font, fill="black")
    y += int((title_font.getbbox("Ag")[3] - title_font.getbbox("Ag")[1]) * line_spacing * 1.5)

    # For each rubric point
    for item in refined_summary:
        rid = item.get("id", "")
        name = item.get("name", rid)
        rating = (item.get("rating") or "").capitalize()
        comment = item.get("comment") or ""

        header = f"{name} — {rating}"
        header_height = h2_font.getbbox("Ag")[3] - h2_font.getbbox("Ag")[1]
        if y + header_height * 4 > H - margin:
            # if page overflow ever needed, here we could create more pages
            # but for now, just stop rendering
            break

        draw.text((margin, y), header, font=h2_font, fill="black")
        y += int(header_height * line_spacing)

        line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
        for line in _wrap_text(draw, comment, body_font, max_text_width):
            draw.text((margin, y), line, font=body_font, fill="black")
            y += int(line_hb * line_spacing)

        y += int(line_hb * 0.8)

    return img


# -----------------------------------------
# PAGE: HIGH-SCORING IDEAL OUTLINE
# -----------------------------------------

def render_high_scoring_outline_page(
    high_scoring_outline: Dict[str, Any],
    page_size: Tuple[int, int] = (2480, 3508),
) -> Image.Image:
    """
    Render a dedicated page for the "High-Scoring Ideal Outline" with structured headings and detailed content.

    The outline supports:
    - Section headings with descriptions (Format: "Heading: Description")
    - Regular bullet points for additional content
    - Proper indentation and visual hierarchy
    - Multi-line text wrapping with appropriate spacing
    """
    W, H = page_size
    margin = int(W * 0.07)
    line_spacing = 1.5

    title_font = _get_font(72)
    h3_font = _get_font(52)
    body_font = _get_font(44)

    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    y = margin

    max_text_width = W - 2 * margin

    # Title (always use fixed title, never extend it)
    outline_title = "High-Scoring Ideal Outline"
    draw.text((margin, y), outline_title, font=title_font, fill="black")
    line_h_title = title_font.getbbox("Ag")[3] - title_font.getbbox("Ag")[1]
    y += int(line_h_title * line_spacing * 1.5)

    outline_points = high_scoring_outline.get("outline_points", [])

    structured_points: List[Dict[str, Any]] = []
    for raw_point in outline_points:
        heading = ""
        summary = ""
        key_points: List[str] = []
        if isinstance(raw_point, dict):
            heading = str(
                raw_point.get("heading")
                or raw_point.get("title")
                or raw_point.get("section")
                or ""
            ).strip()
            summary = str(
                raw_point.get("summary")
                or raw_point.get("description")
                or raw_point.get("overview")
                or ""
            ).strip()
            raw_key_points = (
                raw_point.get("key_points")
                or raw_point.get("bullets")
                or raw_point.get("points")
                or []
            )
            if isinstance(raw_key_points, str):
                if raw_key_points.strip():
                    key_points = [raw_key_points.strip()]
            else:
                key_points = [str(p).strip() for p in raw_key_points if p]
        else:
            text = str(raw_point).strip()
            if ":" in text:
                parts = text.split(":", 1)
                heading = parts[0].strip()
                summary = parts[1].strip()
            else:
                summary = text
        structured_points.append(
            {"heading": heading, "summary": summary, "key_points": key_points}
        )

    structured_points = [
        pt for pt in structured_points if pt["heading"] or pt["summary"] or pt["key_points"]
    ]

    if not structured_points:
        line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
        draw.text((margin, y), "No outline provided", font=body_font, fill="gray")
        return img

    line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
    h2_font = _get_font(56)
    line_h2 = h2_font.getbbox("Ag")[3] - h2_font.getbbox("Ag")[1]

    for point in structured_points:
        heading_text = point.get("heading", "")
        summary_text = point.get("summary", "")
        bullet_items = point.get("key_points", [])

        if heading_text:
            if y + line_h2 * 2 > H - margin:
                break
            draw.text((margin, y), heading_text, font=h2_font, fill="darkblue")
            y += int(line_h2 * line_spacing)

        if summary_text:
            wrapped_lines = _wrap_text(
                draw, summary_text, body_font, max_text_width - int(0.05 * W)
            )
            for line in wrapped_lines:
                if y + line_hb > H - margin:
                    break
                draw.text((margin + int(0.03 * W), y), line, font=body_font, fill="black")
                y += int(line_hb * line_spacing)

        if bullet_items:
            for bullet in bullet_items:
                bullet_lines = _wrap_text(
                    draw, str(bullet).strip(), body_font, max_text_width - int(0.08 * W)
                )
                for idx, line in enumerate(bullet_lines):
                    if y + line_hb > H - margin:
                        break
                    if idx == 0:
                        draw.text(
                            (margin + int(0.04 * W), y), f"• {line}", font=body_font, fill="black"
                        )
                    else:
                        draw.text(
                            (margin + int(0.08 * W), y), line, font=body_font, fill="black"
                        )
                    y += int(line_hb * line_spacing)

        y += int(line_hb * 0.7)

    return img


# -----------------------------
# MAIN PIPELINE
# -----------------------------


def grade_pdf_answer(
    pdf_path: str,
    subject: str,
    output_json_path: str,
    output_pdf_path: str,
    user_id: Optional[str] = None,
    log_path: Optional[str] = None,
    request_id: Optional[str] = None,
    progress_tracker: Optional[Any] = None,
) -> None:
    start_ts = time.perf_counter()
    request_id = request_id or uuid.uuid4().hex[:8]
    _append_log(
        log_path,
        "INFO",
        f"request={request_id} start pdf={os.path.basename(pdf_path)} subject={subject}",
    )

    # Initialize progress tracker if not provided
    if progress_tracker is None:
        logs_dir = os.path.dirname(log_path) if log_path else None
        progress_tracker = OCRProgressTracker(logs_dir=logs_dir)
    
    # Dictionary to store step timings
    step_timings: Dict[str, float] = {}
    
    # Total steps in pipeline
    TOTAL_STEPS = 11

    # Validate all inputs before processing
    print("Validating inputs...")
    validate_input_paths(pdf_path, output_json_path, output_pdf_path)

    # Validate subject name
    if not subject or len(subject.strip()) == 0:
        raise ValueError("Subject name cannot be empty")

    # Create unique output directory per request to prevent concurrent process conflicts
    output_dir = os.path.join(tempfile.gettempdir(), f"grok_images_{request_id}")
    os.makedirs(output_dir, exist_ok=True)
    print(f"Created unique temp directory: {output_dir}")

    try:
        grok_key, vision_client = load_environment()

        print("Step 1: Converting PDF pages to images (for Grok)...")
        progress_tracker.update_progress(
            request_id=request_id,
            step="Converting PDF to images",
            step_number=1,
            total_steps=TOTAL_STEPS,
            progress_percent=5.0,
            message="Converting PDF pages to images...",
        )
        step_start = time.perf_counter()
        page_images = pdf_to_page_images_for_grok(pdf_path, output_dir=output_dir)
        step_duration = time.perf_counter() - step_start
        step_timings["Step 1: Convert PDF to images"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=1 name=convert_pdf_images duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        print("Step 2: Running OCR on PDF (Google Vision)...")
        progress_tracker.update_progress(
            request_id=request_id,
            step="OCR Processing",
            step_number=2,
            total_steps=TOTAL_STEPS,
            progress_percent=15.0,
            message="Running OCR on PDF pages...",
            details={"pages_completed": 0, "total_pages": len(page_images)},
        )
        step_start = time.perf_counter()
        
        # Configure timeouts (can be made configurable via environment variables)
        per_page_timeout = float(os.getenv("OCR_PER_PAGE_TIMEOUT", "120.0"))  # 2 minutes per page
        overall_timeout = float(os.getenv("OCR_OVERALL_TIMEOUT", "600.0"))  # 10 minutes total
        
        # Configure retry parameters (for retry logic implementation)
        max_retries = int(os.getenv("OCR_MAX_RETRIES", "3"))  # Maximum retry attempts per page
        retry_base_delay = float(os.getenv("OCR_RETRY_BASE_DELAY", "1.0"))  # Base delay in seconds
        retry_max_delay = float(os.getenv("OCR_RETRY_MAX_DELAY", "60.0"))  # Maximum delay in seconds
        retry_jitter_range = float(os.getenv("OCR_RETRY_JITTER_RANGE", "0.2"))  # Jitter range (0.0-1.0)
        rate_limit_base_delay = float(os.getenv("OCR_RATE_LIMIT_BASE_DELAY", "5.0"))  # Base delay for rate limits
        rate_limit_max_delay = float(os.getenv("OCR_RATE_LIMIT_MAX_DELAY", "300.0"))  # Max delay for rate limits (5 minutes)
        
        # Configure parallel processing
        concurrent_pages = int(os.getenv("OCR_CONCURRENT_PAGES", "2"))  # Number of pages to process in parallel
        
        # Configure batch orchestration
        batch_size = int(os.getenv("OCR_BATCH_SIZE", "5"))  # Number of pages per batch
        batch_failure_threshold = float(os.getenv("OCR_BATCH_FAILURE_THRESHOLD", "0.5"))  # Stop if batch failure rate exceeds this
        
        # Configure adaptive concurrency
        adaptive_concurrency_enabled = os.getenv("OCR_ADAPTIVE_CONCURRENCY_ENABLED", "true").lower() in ("true", "1", "yes", "on")
        adaptive_min_concurrency = int(os.getenv("OCR_ADAPTIVE_MIN_CONCURRENCY", "1"))
        adaptive_max_concurrency = int(os.getenv("OCR_ADAPTIVE_MAX_CONCURRENCY", "4"))
        adaptive_latency_threshold_ms = float(os.getenv("OCR_ADAPTIVE_LATENCY_THRESHOLD_MS", "90000.0"))
        adaptive_stable_batches = int(os.getenv("OCR_ADAPTIVE_STABLE_BATCHES", "2"))
        
        # Configure image optimization
        image_optimization_enabled = os.getenv("OCR_IMAGE_OPTIMIZATION_ENABLED", "true").lower() in ("true", "1", "yes", "on")
        image_max_dimension = int(os.getenv("OCR_IMAGE_MAX_DIMENSION", "2048"))
        image_min_dimension_for_optimization = int(os.getenv("OCR_IMAGE_MIN_DIMENSION_FOR_OPTIMIZATION", "1500"))
        
        ocr_data = run_ocr_on_pdf(
            vision_client=vision_client,
            pdf_path=pdf_path,
            per_page_timeout=per_page_timeout,
            overall_timeout=overall_timeout,
            log_path=log_path,
            request_id=request_id,
            max_retries=max_retries,
            retry_base_delay=retry_base_delay,
            retry_max_delay=retry_max_delay,
            retry_jitter_range=retry_jitter_range,
            rate_limit_base_delay=rate_limit_base_delay,
            rate_limit_max_delay=rate_limit_max_delay,
            concurrent_pages=concurrent_pages,
            batch_size=batch_size,
            batch_failure_threshold=batch_failure_threshold,
            adaptive_concurrency_enabled=adaptive_concurrency_enabled,
            adaptive_min_concurrency=adaptive_min_concurrency,
            adaptive_max_concurrency=adaptive_max_concurrency,
            adaptive_latency_threshold_ms=adaptive_latency_threshold_ms,
            adaptive_stable_batches=adaptive_stable_batches,
            image_optimization_enabled=image_optimization_enabled,
            image_max_dimension=image_max_dimension,
            image_min_dimension_for_optimization=image_min_dimension_for_optimization,
            progress_tracker=progress_tracker,
        )
        step_duration = time.perf_counter() - step_start
        step_timings["Step 2: Google Vision OCR"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=2 name=google_vision_ocr duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")
        
        # Update progress after OCR completes
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="OCR Processing",
                step_number=2,
                total_steps=TOTAL_STEPS,
                progress_percent=45.0,
                message="OCR processing complete",
            )

        print("Step 3: Detecting sections/headings with Grok...")
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Detecting sections",
                step_number=3,
                total_steps=TOTAL_STEPS,
                progress_percent=50.0,
                message="Detecting sections and headings...",
            )
        step_start = time.perf_counter()
        sections, section_token_usage = call_grok_for_section_detection(
            grok_api_key=grok_key,
            ocr_data=ocr_data,
            page_images=page_images,
        )
        step_duration = time.perf_counter() - step_start
        step_timings["Step 3: Grok section detection"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=3 name=grok_section_detection duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        # Debug dump (only if DEBUG_SECTIONS environment variable is set)
        if os.getenv("DEBUG_SECTIONS", "").lower() in ("true", "1", "yes"):
            debug_dump_sections(sections, output_path="debug_sections.json")

        # Track total token usage
        total_input_tokens = section_token_usage.get("input_tokens", 0)
        total_output_tokens = section_token_usage.get("output_tokens", 0)

        print("Step 4: Loading subject-wise rubric DOCX...")
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Loading rubric",
                step_number=4,
                total_steps=TOTAL_STEPS,
                progress_percent=55.0,
                message="Loading subject rubric...",
            )
        step_start = time.perf_counter()
        subject_rubric_text, subject_rubric_path = load_subject_rubric_text(subject)
        step_duration = time.perf_counter() - step_start
        step_timings["Step 4: Load subject rubric"] = step_duration
        if subject_rubric_path:
            print(f"Using subject rubric file: {subject_rubric_path}")
        else:
            print("Warning: No subject rubric file found; grading will be weaker.")
            _append_log(
                log_path,
                "WARNING",
                f"request={request_id} missing_subject_rubric subject={subject}",
            )
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=4 name=load_subject_rubric duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        print("Step 5: Calling Grok for subject-wise grading...")
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Subject grading",
                step_number=5,
                total_steps=TOTAL_STEPS,
                progress_percent=60.0,
                message="Grading with subject rubric...",
            )
        step_start = time.perf_counter()
        grading_result, grading_token_usage = call_grok_for_grading(
            grok_api_key=grok_key,
            subject=subject,
            subject_rubric_text=subject_rubric_text,
            ocr_data=ocr_data,
            sections=sections,
            page_images=page_images,
        )
        step_duration = time.perf_counter() - step_start
        step_timings["Step 5: Grok subject grading"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=5 name=grok_subject_grading duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        # Accumulate token usage
        total_input_tokens += grading_token_usage.get("input_tokens", 0)
        total_output_tokens += grading_token_usage.get("output_tokens", 0)

        grading_result.setdefault("subject", subject)
        if not grading_result.get("max_marks"):
            grading_result["max_marks"] = 20

        # Add token usage to grading result before saving
        grading_result["token_usage"] = {
            "input_tokens": total_input_tokens,
            "output_tokens": total_output_tokens,
            "total_tokens": total_input_tokens + total_output_tokens
        }
        print(f"OCR completed. Token usage: Input: {total_input_tokens}, Output: {total_output_tokens}")

        print("Saving grading JSON...")
        with open(output_json_path, "w", encoding="utf-8") as f:
            json.dump(grading_result, f, ensure_ascii=False, indent=2)
        print(f"Grading JSON saved to {output_json_path}")

        print("Step 6: Rendering subject-wise report pages...")
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Rendering report",
                step_number=6,
                total_steps=TOTAL_STEPS,
                progress_percent=70.0,
                message="Rendering subject report pages...",
            )
        step_start = time.perf_counter()
        report_page_size = get_report_page_size(pdf_path)
        subject_report_pages = render_subject_report_pages(
            grading_result,
            page_size=report_page_size,
        )
        step_duration = time.perf_counter() - step_start
        step_timings["Step 6: Render subject report pages"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=6 name=render_subject_report_pages duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        print("Step 7: Loading refined rubric DOCX...")
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Loading refined rubric",
                step_number=7,
                total_steps=TOTAL_STEPS,
                progress_percent=75.0,
                message="Loading refined rubric...",
            )
        step_start = time.perf_counter()
        refined_rubric_text, refined_rubric_path = load_refined_rubric_text()
        step_duration = time.perf_counter() - step_start
        step_timings["Step 7: Load refined rubric"] = step_duration
        if refined_rubric_path:
            print(f"Using refined rubric file: {refined_rubric_path}")
        else:
            print("Warning: No refined rubric file found; annotations will be weaker.")
            _append_log(
                log_path,
                "WARNING",
                f"request={request_id} missing_refined_rubric",
            )
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=7 name=load_refined_rubric duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        print("Step 8: Calling Grok for refined rubric annotations...")
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Refined annotations",
                step_number=8,
                total_steps=TOTAL_STEPS,
                progress_percent=80.0,
                message="Generating refined annotations...",
            )
        step_start = time.perf_counter()
        refined_result, refined_token_usage = call_grok_for_refined_rubric_annotations(
            grok_api_key=grok_key,
            refined_rubric_text=refined_rubric_text,
            ocr_data=ocr_data,
            sections=sections,
            page_images=page_images,
        )
        step_duration = time.perf_counter() - step_start
        step_timings["Step 8: Grok refined annotations"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=8 name=grok_refined_annotations duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        # Accumulate token usage
        total_input_tokens += refined_token_usage.get("input_tokens", 0)
        total_output_tokens += refined_token_usage.get("output_tokens", 0)

        annotations = refined_result.get("annotations", []) or []
        refined_summary = refined_result.get("refined_rubric_summary", []) or []

        # Validate refined summary schema
        if refined_summary:
            try:
                validate_refined_summary(refined_summary)
                print(f"Validated {len(refined_summary)} refined rubric summary items")
            except ValueError as e:
                print(f"WARNING: Refined summary validation failed: {e}")
                _append_log(
                    log_path,
                    "WARNING",
                    f"request={request_id} refined_summary_validation_failed error={e}",
                )

        # Validate annotations schema
        valid_annotations = []
        for idx, ann in enumerate(annotations):
            if validate_annotation(ann, idx):
                valid_annotations.append(ann)
        if len(valid_annotations) < len(annotations):
            skipped_count = len(annotations) - len(valid_annotations)
            print(f"WARNING: {skipped_count} annotations failed validation and were skipped")
            _append_log(
                log_path,
                "WARNING",
                f"request={request_id} annotation_validation_skipped count={skipped_count}",
            )
        annotations = valid_annotations

        print("Step 9: Calling Grok for page-wise improvement suggestions...")
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Page suggestions",
                step_number=9,
                total_steps=TOTAL_STEPS,
                progress_percent=85.0,
                message="Generating page-wise suggestions...",
            )
        step_start = time.perf_counter()
        page_suggestions_result, page_suggestions_token_usage = call_grok_for_page_wise_suggestions(
            grok_api_key=grok_key,
            subject=subject,
            subject_rubric_text=subject_rubric_text,
            ocr_data=ocr_data,
            sections=sections,
        )
        step_duration = time.perf_counter() - step_start
        step_timings["Step 9: Grok page suggestions"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=9 name=grok_page_suggestions duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        # Accumulate token usage
        total_input_tokens += page_suggestions_token_usage.get("input_tokens", 0)
        total_output_tokens += page_suggestions_token_usage.get("output_tokens", 0)

        page_suggestions = page_suggestions_result.get("page_suggestions", []) or []

        print("Step 10: Annotating answer pages with improvement suggestions...")
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Annotating pages",
                step_number=10,
                total_steps=TOTAL_STEPS,
                progress_percent=90.0,
                message="Annotating answer pages...",
            )
        step_start = time.perf_counter()
        annotated_answer_pages = annotate_pdf_answer_pages(
            pdf_path=pdf_path,
            ocr_data=ocr_data,
            sections=sections,
            annotations=annotations,
            page_suggestions=page_suggestions,
            refined_summary=refined_summary,
        )
        step_duration = time.perf_counter() - step_start
        step_timings["Step 10: Annotate answer pages"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=10 name=annotate_answer_pages duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")

        #print("Step 11: Rendering refined rubric summary page...")
        #refined_summary_page = render_refined_rubric_summary_page(refined_summary)
        # Assemble final PDF incrementally to avoid memory accumulation:
        #   1) Subject report pages (includes marks table and question expectations)
        #   2) Annotated answer pages (with left-side improvement suggestions)
        #   3) Refined rubric summary page

        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Writing PDF",
                step_number=11,
                total_steps=TOTAL_STEPS,
                progress_percent=95.0,
                message="Writing final PDF...",
            )
        step_start = time.perf_counter()
        print(f"Step 11: Writing final PDF incrementally to {output_pdf_path} ...")
        
        # Create PDF writer for incremental writing
        pdf_writer = PdfWriter()
        
        # Helper function to convert PIL Image to PDF bytes and add to writer
        def add_image_to_pdf(img: Image.Image) -> None:
            """Convert PIL Image to PDF bytes and add to PDF writer incrementally."""
            buffer = io.BytesIO()
            img.save(buffer, format="PDF", resolution=300.0)
            buffer.seek(0)
            pdf_reader = PdfReader(buffer)
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)
        
        # Add subject report pages (small, 1-2 pages, acceptable to keep in memory)
        for page in subject_report_pages:
            add_image_to_pdf(page)
        
        # Add annotated answer pages incrementally (don't accumulate in memory)
        for page in annotated_answer_pages:
            add_image_to_pdf(page)
        
        # Write final PDF
        with open(output_pdf_path, "wb") as output_file:
            pdf_writer.write(output_file)
        
        step_duration = time.perf_counter() - step_start
        step_timings["Step 11: Write final PDF"] = step_duration
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} step=11 name=merge_and_write_pdf duration_ms={int(step_duration * 1000)}",
        )
        print(f"  ✓ Completed in {_format_time(step_duration)}")
        
        # Update progress to complete
        if progress_tracker:
            progress_tracker.update_progress(
                request_id=request_id,
                step="Complete",
                step_number=11,
                total_steps=TOTAL_STEPS,
                progress_percent=100.0,
                message="✅ Evaluation complete!",
            )
        
        # Calculate total time
        total_duration = time.perf_counter() - start_ts
        
        # Print comprehensive timing report
        print("\n" + "="*70)
        print("TIMING REPORT")
        print("="*70)
        for step_name, duration in step_timings.items():
            print(f"  {step_name:.<50} {_format_time(duration)}")
        print("-"*70)
        print(f"  {'TOTAL TIME':.<50} {_format_time(total_duration)}")
        print("="*70 + "\n")
        
        # Log the timing report
        timing_report_lines = [
            f"request={request_id} TIMING_REPORT_START",
        ]
        for step_name, duration in step_timings.items():
            timing_report_lines.append(
                f"request={request_id} {step_name} duration={_format_time(duration)} duration_sec={duration:.2f}"
            )
        timing_report_lines.append(
            f"request={request_id} TOTAL_TIME duration={_format_time(total_duration)} duration_sec={total_duration:.2f}"
        )
        timing_report_lines.append(f"request={request_id} TIMING_REPORT_END")
        
        for line in timing_report_lines:
            _append_log(log_path, "INFO", line)
        
        _append_log(
            log_path,
            "INFO",
            f"request={request_id} completed total_duration_ms={int(total_duration * 1000)}",
        )
    except Exception as exc:
        _append_log(
            log_path,
            "ERROR",
            f"request={request_id} error={exc} traceback={traceback.format_exc().strip()}",
        )
        raise

    finally:
        # Clean up unique temp directory
        if os.path.exists(output_dir):
            try:
                shutil.rmtree(output_dir)
                print(f"Cleaned up temp directory: {output_dir}")
            except Exception as e:
                print(f"WARNING: Failed to remove temp directory {output_dir}: {e}")
                _append_log(
                    log_path,
                    "WARNING",
                    f"request={request_id} temp_dir_cleanup_failed error={e}",
                )

# -----------------------------
# CLI
# -----------------------------


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Grade a handwritten PDF answer using Grok + Google Vision OCR, "
        "generate a subject-wise report, annotate the answer pages, "
        "and add a refined-rubric summary page."
    )
    parser.add_argument("--pdf_path", required=True, help="Path to the PDF file to grade.")
    parser.add_argument("--subject", required=True, help="Subject name, e.g., 'British History'.")
    parser.add_argument(
        "--output_json",
        default="grading_result.json",
        help="Path to write the grading JSON output.",
    )
    parser.add_argument(
        "--output_pdf",
        default="result.pdf",
        help="Path to write the final PDF (report + annotated pages + rubric summary).",
    )
    return parser.parse_args()


if __name__ == "__main__":
    try:
        args = parse_args()
        grade_pdf_answer(
            pdf_path=args.pdf_path,
            subject=args.subject,
            output_json_path=args.output_json,
            output_pdf_path=args.output_pdf,
        )
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        sys.exit(1)
