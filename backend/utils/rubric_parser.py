#!/usr/bin/env python3
"""
Rubric Parser - Extract evaluation criteria from Word doc rubrics.

Parses CSS subject rubric .docx files to extract:
- Dimensions (criteria) with weights and marks
- Indicators (specific checklist items)
- Objectives (what each criterion assesses)
- Bot integration instructions
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import re
from functools import lru_cache

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

import logging

logger = logging.getLogger(__name__)


# ======================== Data Classes ========================

@dataclass
class RubricIndicator:
    """Single indicator/checklist item for a dimension."""
    text: str
    dimension_name: str


@dataclass
class RubricDimension:
    """One evaluation criterion (e.g., Understanding & Relevance)."""
    name: str
    weight_percent: float
    max_marks: float
    objective: str
    indicators: List[str] = field(default_factory=list)
    assessment_focus: str = ""  # For report table "Assessment Focus" column

    def __post_init__(self):
        """Generate assessment_focus from objective if not provided."""
        if not self.assessment_focus and self.objective:
            # Use first sentence of objective as assessment focus
            self.assessment_focus = self.objective.split('.')[0].strip()


@dataclass
class RubricStructure:
    """Complete rubric for a subject."""
    subject: str
    subject_display_name: str
    total_marks: int
    dimensions: List[RubricDimension] = field(default_factory=list)
    bot_instructions: str = ""
    all_indicators: List[RubricIndicator] = field(default_factory=list)
    rubric_file_path: Optional[Path] = None

    def __post_init__(self):
        """Build flattened indicator list for easy access."""
        if not self.all_indicators:
            for dim in self.dimensions:
                for indicator_text in dim.indicators:
                    self.all_indicators.append(
                        RubricIndicator(
                            text=indicator_text,
                            dimension_name=dim.name
                        )
                    )

    @property
    def total_indicators(self) -> int:
        """Count of all indicators across all dimensions."""
        return len(self.all_indicators)

    @property
    def dimension_names(self) -> List[str]:
        """List of all dimension names."""
        return [dim.name for dim in self.dimensions]


# ======================== Parser ========================

class RubricParser:
    """Parse Word doc rubrics and extract structured evaluation criteria."""

    def __init__(self, rubrics_base_dir: Optional[Path] = None):
        """
        Initialize parser.

        Args:
            rubrics_base_dir: Base directory containing subject rubric folders.
                            Defaults to backend/Rubrics/
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )

        if rubrics_base_dir is None:
            # Default: backend/Rubrics/ relative to this file
            # This file is in: backend/utils/rubric_parser.py
            # Rubrics is at: backend/Rubrics/
            backend_dir = Path(__file__).resolve().parents[1]
            rubrics_base_dir = backend_dir.parent / "Rubrics"

        self.rubrics_base_dir = Path(rubrics_base_dir)

        if not self.rubrics_base_dir.exists():
            logger.warning(f"Rubrics directory not found: {self.rubrics_base_dir}")

        # Subject name normalization mapping
        self.subject_mapping = self._build_subject_mapping()

    def _build_subject_mapping(self) -> Dict[str, Path]:
        """Build mapping from normalized subject names to rubric file paths."""
        mapping = {}

        if not self.rubrics_base_dir.exists():
            return mapping

        # Find all .docx files in Rubrics subdirectories
        for docx_file in self.rubrics_base_dir.glob("**/*.docx"):
            # Extract subject name from file stem
            subject_name = docx_file.stem

            # Normalize: lowercase, replace spaces/underscores with hyphens
            normalized = self._normalize_subject_name(subject_name)

            mapping[normalized] = docx_file

            logger.debug(f"Mapped '{normalized}' -> {docx_file}")

        logger.info(f"Found {len(mapping)} rubric files")
        return mapping

    @staticmethod
    def _normalize_subject_name(name: str) -> str:
        """
        Normalize subject name for matching.

        Examples:
            "Political Science" -> "political-science"
            "International Relations" -> "international-relations"
            "Indo-Pak History" -> "indo-pak-history"
        """
        normalized = name.lower().strip()
        normalized = re.sub(r'[_\s]+', '-', normalized)
        normalized = re.sub(r'[^a-z0-9\-]', '', normalized)
        normalized = re.sub(r'-+', '-', normalized)
        return normalized.strip('-')

    @lru_cache(maxsize=32)
    def parse_rubric(self, subject: str) -> RubricStructure:
        """
        Parse rubric for a subject.

        Args:
            subject: Subject name (e.g., "Political Science", "political-science")

        Returns:
            RubricStructure with all dimensions, indicators, objectives

        Raises:
            FileNotFoundError: If rubric file not found for subject
            ValueError: If rubric parsing fails
        """
        normalized_subject = self._normalize_subject_name(subject)

        # Find rubric file
        if normalized_subject not in self.subject_mapping:
            available = ", ".join(self.subject_mapping.keys())
            raise FileNotFoundError(
                f"No rubric found for subject '{subject}' (normalized: '{normalized_subject}'). "
                f"Available subjects: {available}"
            )

        rubric_file = self.subject_mapping[normalized_subject]

        logger.info(f"Parsing rubric: {rubric_file}")

        # Parse Word document
        try:
            doc = Document(rubric_file)
        except Exception as e:
            raise ValueError(f"Failed to read Word doc {rubric_file}: {e}")

        # Extract rubric structure
        rubric = self._extract_rubric_structure(doc, rubric_file)

        logger.info(
            f"Parsed rubric for {rubric.subject}: "
            f"{len(rubric.dimensions)} dimensions, "
            f"{rubric.total_indicators} indicators"
        )

        return rubric

    def _extract_rubric_structure(
        self, doc: Document, file_path: Path
    ) -> RubricStructure:
        """Extract complete rubric structure from Word document."""

        # Extract subject name from first heading or file name
        subject_display_name = self._extract_subject_name(doc, file_path)

        # Extract dimensions from tables and content
        dimensions = self._extract_dimensions(doc)

        # Extract bot instructions
        bot_instructions = self._extract_bot_instructions(doc)

        return RubricStructure(
            subject=self._normalize_subject_name(subject_display_name),
            subject_display_name=subject_display_name,
            total_marks=20,  # Standard for CSS 20-mark questions
            dimensions=dimensions,
            bot_instructions=bot_instructions,
            rubric_file_path=file_path
        )

    def _extract_subject_name(self, doc: Document, file_path: Path) -> str:
        """Extract subject display name from document."""
        # Try to find in first few paragraphs
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            # Look for pattern: "CSS [Subject] Evaluation Rubric"
            match = re.search(r'CSS\s+(.+?)\s+Evaluation\s+Rubric', text, re.IGNORECASE)
            if match:
                subject = match.group(1).strip()
                # Remove "(For Bot Assessment...)" part
                subject = re.sub(r'\(For Bot.*?\)', '', subject, flags=re.IGNORECASE).strip()
                return subject

        # Fallback to file name
        return file_path.stem

    def _extract_dimensions(self, doc: Document) -> List[RubricDimension]:
        """Extract all dimensions with indicators from document."""
        dimensions = []

        # First, try to extract from the summary table (usually first table)
        dimension_info = self._extract_dimensions_from_table(doc)

        # Then, extract detailed indicators from headings and content
        detailed_dims = self._extract_dimension_details(doc)

        # Merge: use table for weights/marks, use details for indicators/objectives
        for dim_name, (weight, marks) in dimension_info.items():
            # Find matching detailed dimension
            detailed = detailed_dims.get(dim_name)

            if detailed:
                dimensions.append(RubricDimension(
                    name=dim_name,
                    weight_percent=weight,
                    max_marks=marks,
                    objective=detailed.get('objective', ''),
                    indicators=detailed.get('indicators', []),
                    assessment_focus=detailed.get('objective', '').split('.')[0].strip()
                ))
            else:
                # Create basic dimension from table only
                dimensions.append(RubricDimension(
                    name=dim_name,
                    weight_percent=weight,
                    max_marks=marks,
                    objective=f"Assess {dim_name.lower()}",
                    indicators=[],
                    assessment_focus=f"Measures {dim_name.lower()}"
                ))

        return dimensions

    def _extract_dimensions_from_table(self, doc: Document) -> Dict[str, Tuple[float, float]]:
        """
        Extract dimension names, weights, and marks from summary table.

        Returns:
            Dict mapping dimension_name -> (weight_percent, max_marks)
        """
        dimension_info = {}

        # Look for table with columns: Dimension | Weight (%) | Max Marks
        for table in doc.tables:
            if len(table.rows) < 2:
                continue

            # Check if this is the summary table
            header_row = [cell.text.strip().lower() for cell in table.rows[0].cells]

            if 'dimension' in ' '.join(header_row) and 'marks' in ' '.join(header_row):
                # This is likely the summary table
                for row in table.rows[1:]:  # Skip header
                    cells = [cell.text.strip() for cell in row.cells]

                    if len(cells) < 3:
                        continue

                    dim_name = cells[0]
                    weight_text = cells[1]
                    marks_text = cells[2]

                    # Skip empty or header-like rows
                    if not dim_name or 'total' in dim_name.lower():
                        continue

                    # Extract weight percentage (e.g., "20%" -> 20.0)
                    weight_match = re.search(r'(\d+(?:\.\d+)?)', weight_text)
                    weight = float(weight_match.group(1)) if weight_match else 0.0

                    # Extract max marks (e.g., "4" -> 4.0)
                    marks_match = re.search(r'(\d+(?:\.\d+)?)', marks_text)
                    marks = float(marks_match.group(1)) if marks_match else 0.0

                    # Clean dimension name (remove numbering like "1. ")
                    dim_name_clean = re.sub(r'^\d+\.\s*', '', dim_name).strip()

                    dimension_info[dim_name_clean] = (weight, marks)

                break  # Found the table, stop searching

        return dimension_info

    def _extract_dimension_details(self, doc: Document) -> Dict[str, Dict]:
        """
        Extract objectives and indicators for each dimension from document content.

        Returns:
            Dict mapping dimension_name -> {objective: str, indicators: List[str]}
        """
        details = {}
        current_dimension = None
        current_objective = None
        current_indicators = []
        in_indicators_section = False

        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else "Normal"

            if not text:
                continue

            # Check if this is a dimension heading (e.g., "1. Understanding of Question...")
            # Pattern: number + dimension name + marks in parentheses
            dim_match = re.match(
                r'^\d+\.\s*(.+?)\s*\((\d+(?:\.\d+)?)\s*[Mm]arks?\)',
                text
            )

            if dim_match or (style and 'heading' in style.lower() and 'marks' in text.lower()):
                # Save previous dimension
                if current_dimension:
                    details[current_dimension] = {
                        'objective': current_objective or '',
                        'indicators': current_indicators
                    }

                # Start new dimension
                if dim_match:
                    current_dimension = dim_match.group(1).strip()
                else:
                    # Extract from heading text
                    current_dimension = re.sub(r'\(.*?\)', '', text).strip()
                    current_dimension = re.sub(r'^\d+\.\s*', '', current_dimension).strip()

                current_objective = None
                current_indicators = []
                in_indicators_section = False
                continue

            # Check for "Objective:" line
            if text.lower().startswith('objective:'):
                objective_text = text[len('objective:'):].strip()
                current_objective = objective_text
                continue

            # Check for "Indicators:" line
            if text.lower().startswith('indicators:'):
                in_indicators_section = True
                continue

            # If we're in indicators section, collect indicators
            if in_indicators_section and current_dimension:
                # Indicators are typically bullet points or numbered items
                # Clean up bullet/number prefixes
                indicator = re.sub(r'^[\d•\-\*]+\.?\s*', '', text).strip()

                if indicator and len(indicator) > 10:  # Skip very short lines
                    current_indicators.append(indicator)

        # Save last dimension
        if current_dimension:
            details[current_dimension] = {
                'objective': current_objective or '',
                'indicators': current_indicators
            }

        return details

    def _extract_bot_instructions(self, doc: Document) -> str:
        """Extract 'For Bot Integration' section from document."""
        bot_instructions = []
        in_bot_section = False

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Check for "For Bot Integration" heading
            if 'bot integration' in text.lower() or 'bot assessment' in text.lower():
                in_bot_section = True
                continue

            # Stop at next major heading after bot section
            if in_bot_section:
                style = para.style.name if para.style else "Normal"

                # Stop if we hit another major heading
                if style and 'heading 1' in style.lower():
                    break
                if style and 'heading 2' in style.lower() and 'bot' not in text.lower():
                    break

                # Collect bot instruction text
                # Skip "Your AI evaluator should:" header
                if 'evaluator should' not in text.lower():
                    # Clean bullet/number prefixes
                    instruction = re.sub(r'^[\d•\-\*]+\.?\s*', '', text).strip()
                    if instruction:
                        bot_instructions.append(instruction)

        return '\n'.join(bot_instructions)

    def get_available_subjects(self) -> List[str]:
        """Get list of all available subjects with rubrics."""
        return sorted(self.subject_mapping.keys())

    def get_subject_display_name(self, subject: str) -> str:
        """Get display name for a normalized subject name."""
        normalized = self._normalize_subject_name(subject)

        if normalized in self.subject_mapping:
            rubric_file = self.subject_mapping[normalized]
            return rubric_file.stem

        return subject


# ======================== Convenience Functions ========================

_global_parser: Optional[RubricParser] = None


def get_rubric_parser() -> RubricParser:
    """Get global rubric parser instance (singleton)."""
    global _global_parser

    if _global_parser is None:
        _global_parser = RubricParser()

    return _global_parser


def load_rubric(subject: str) -> RubricStructure:
    """
    Load rubric for a subject (convenience function).

    Args:
        subject: Subject name (e.g., "Political Science" or "political-science")

    Returns:
        RubricStructure

    Raises:
        FileNotFoundError: If rubric not found
    """
    parser = get_rubric_parser()
    return parser.parse_rubric(subject)


def get_available_subjects() -> List[str]:
    """Get list of all subjects with rubrics."""
    parser = get_rubric_parser()
    return parser.get_available_subjects()
