#!/usr/bin/env python3
"""
Helpers for loading rubric documents and listing available subjects.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document

RUBRICS_DIR_NAME = "20marks_Rubrics"


@dataclass
class SubjectRubric:
    subject_id: str
    display_name: str
    doc_path: Path


def _normalize_subject_name(name: str) -> str:
    normalized = name.lower().strip()
    normalized = re.sub(r"[\s_]+", "-", normalized)  # collapse whitespace/underscores to single hyphen
    normalized = re.sub(r"[^a-z0-9-]", "", normalized)  # strip everything except alphanumerics and hyphen
    normalized = re.sub(r"-{2,}", "-", normalized)
    return normalized.strip("-")


def _rubrics_root() -> Path:
    backend_dir = Path(__file__).resolve().parents[1]
    candidate = backend_dir / RUBRICS_DIR_NAME
    if not candidate.exists():
        raise FileNotFoundError(
            f"Rubrics directory '{RUBRICS_DIR_NAME}' not found relative to backend utils."
        )
    return candidate


def list_subject_rubrics() -> List[SubjectRubric]:
    """
    Enumerate rubric documents from the 20marks_Rubrics directory.
    """

    subjects: List[SubjectRubric] = []
    root = _rubrics_root()

    for entry in sorted(root.iterdir()):
        if entry.is_dir():
            doc_files = list(entry.glob("*.docx"))
            if not doc_files:
                continue
            doc_path = doc_files[0]
            display = entry.name
            subject_id = _normalize_subject_name(display)
            subjects.append(SubjectRubric(subject_id=subject_id, display_name=display, doc_path=doc_path))

    return subjects


@lru_cache(maxsize=32)
def load_rubric_text(subject_id: str) -> str:
    """
    Load the rubric docx text for a subject and return a plain-text representation.
    """

    normalized = _normalize_subject_name(subject_id)
    lookup: Dict[str, SubjectRubric] = {rub.subject_id: rub for rub in list_subject_rubrics()}
    if normalized not in lookup:
        available = ", ".join(sorted(lookup.keys()))
        raise FileNotFoundError(
            f"No rubric found for subject '{subject_id}'. Available subjects: {available}"
        )

    doc = Document(str(lookup[normalized].doc_path))
    lines: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    return "\n".join(lines)


@lru_cache(maxsize=1)
def load_feedback_template_text() -> str:
    """
    Load the shared 20-marks feedback template docx as plain text.
    """

    root = _rubrics_root()
    template_path: Optional[Path] = None
    for child in root.iterdir():
        if child.is_file() and child.suffix.lower() == ".docx" and "feedback" in child.stem.lower():
            template_path = child
            break
    if not template_path:
        raise FileNotFoundError(
            "20 Marks Question Feedback Template.docx not found in rubrics directory."
        )

    doc = Document(str(template_path))
    content: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            content.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                content.append(" | ".join(cells))

    return "\n".join(content)


def list_available_subjects() -> List[Dict[str, str]]:
    """
    Convenience helper for API responses.
    """

    subjects = list_subject_rubrics()
    return [
        {"id": subject.subject_id, "display_name": subject.display_name}
        for subject in subjects
    ]


@lru_cache(maxsize=64)
def get_subject_display_name(subject_id: str) -> str:
    normalized = _normalize_subject_name(subject_id)
    subjects = list_subject_rubrics()
    mapping = {sub.subject_id: sub.display_name for sub in subjects}
    return mapping.get(normalized, subject_id)
