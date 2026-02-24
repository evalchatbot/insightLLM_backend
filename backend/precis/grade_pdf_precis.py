import argparse
import base64
import io
import importlib.util
import json
import os
import re
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Dict, List, Optional, Tuple
import xml.etree.ElementTree as ET

try:
    import pymupdf as fitz  # PyMuPDF (preferred)
except ModuleNotFoundError:
    import fitz  # type: ignore
import requests
from PIL import Image
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import HttpResponseError
from dotenv import load_dotenv

try:
    from backend.eng_essay.annotate_pdf_with_essay_rubric import annotate_pdf_essay_pages  # type: ignore
except Exception:
    try:
        # Fallback: try loading from eng_essay sibling directory
        _precis_dir = os.path.dirname(os.path.abspath(__file__))
        _backend_dir = os.path.dirname(_precis_dir)
        ann_path = os.path.join(_backend_dir, "eng_essay", "annotate_pdf_with_essay_rubric.py")
        if os.path.exists(ann_path):
            spec = importlib.util.spec_from_file_location("annotate_pdf_with_essay_rubric", ann_path)
            mod = importlib.util.module_from_spec(spec) if spec else None
            if spec and spec.loader and mod:
                spec.loader.exec_module(mod)
                annotate_pdf_essay_pages = mod.annotate_pdf_essay_pages  # type: ignore
            else:
                annotate_pdf_essay_pages = None  # type: ignore
        else:
            annotate_pdf_essay_pages = None  # type: ignore
    except Exception:
        annotate_pdf_essay_pages = None  # type: ignore


DEFAULT_PRECIS_CRITERIA: List[Dict[str, Any]] = [
    {"id": "comprehension", "criterion": "Comprehension & Understanding of Passage", "marks_allocated": 3},
    {"id": "clarity_expression", "criterion": "Clarity, Expression & Language", "marks_allocated": 3},
    {"id": "brevity", "criterion": "Brevity & Conciseness", "marks_allocated": 2},
    {"id": "organization", "criterion": "Organization & Coherence", "marks_allocated": 2},
    {"id": "tone_meaning", "criterion": "Original Tone & Meaning", "marks_allocated": 2},
    {"id": "originality", "criterion": "Originality & Paraphrasing", "marks_allocated": 2},
    {"id": "grammar_presentation", "criterion": "Grammar & Presentation", "marks_allocated": 1},
    {"id": "title", "criterion": "Title", "marks_allocated": 5},
]

DEFAULT_MODELS: Dict[str, Dict[str, Any]] = {
    "grading": {"model": "grok-4-1-fast-reasoning", "temperature": 0.10},
    "annotations": {"model": "grok-4-1-fast-reasoning", "temperature": 0.15},
    "json_repair": {"model": "grok-4-1-fast-reasoning", "temperature": 0.00},
}

# Keep report text aligned with annotation-style readable text size.
REPORT_BASE_TEXT_SIZE = 12.0
PRECIS_DIR = os.path.dirname(os.path.abspath(__file__))
DEBUG_LLM_DIR = os.path.join(PRECIS_DIR, "debug_llm")
GROK_IMAGES_DIR = os.path.join(PRECIS_DIR, "grok_images_precis")


def _format_duration(seconds: float) -> str:
    if seconds < 60:
        return f"{seconds:.2f}s"
    m = int(seconds // 60)
    s = seconds - m * 60
    return f"{m}m {s:.2f}s"


def clean_json_from_llm(text: str) -> str:
    text = (text or "").strip()
    if not text:
        return ""
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z0-9]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
    return text.strip()


def _extract_json_candidate(text: str) -> str:
    s = clean_json_from_llm(text)
    if not s:
        return s
    if s.startswith("{") and s.endswith("}"):
        return s
    if "{" in s and "}" in s:
        start = s.find("{")
        end = s.rfind("}")
        if end > start:
            return s[start : end + 1]
    return s


def _grok_chat(
    grok_api_key: str,
    messages: List[Dict[str, str]],
    model: str = "grok-4-1-fast-reasoning",
    temperature: float = 0.12,
    max_tokens: Optional[int] = None,
    timeout: int = 180,
    max_retries: int = 8,
) -> Dict[str, Any]:
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {grok_api_key}",
    }
    payload: Dict[str, Any] = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
    }
    if max_tokens is not None:
        payload["max_tokens"] = max_tokens

    last_err: Optional[Exception] = None
    for attempt in range(max_retries + 1):
        try:
            resp = requests.post(
                "https://api.x.ai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=(30, timeout),
            )
            if resp.status_code >= 300:
                err = RuntimeError(f"Grok API error {resp.status_code}: {resp.text}")
                if resp.status_code in (429, 500, 502, 503, 504) and attempt < max_retries:
                    last_err = err
                    delay = min(60.0, 2.0 ** attempt)
                    print(f"  Grok {resp.status_code} retry {attempt + 1}/{max_retries + 1} in {delay:.0f}s...")
                    time.sleep(delay)
                    continue
                raise err
            return resp.json()
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            last_err = e
            if attempt >= max_retries:
                raise
            time.sleep(min(60.0, 2.0 ** attempt))

    raise RuntimeError(f"Grok request failed: {last_err}")


def parse_json_with_repair(
    grok_api_key: str,
    raw_text: str,
    *,
    debug_tag: str,
    max_fix_attempts: int = 2,
    repair_model: str = "grok-4-1-fast-reasoning",
    repair_temperature: float = 0.0,
) -> Dict[str, Any]:
    os.makedirs(DEBUG_LLM_DIR, exist_ok=True)
    with open(os.path.join(DEBUG_LLM_DIR, f"{debug_tag}_raw.txt"), "w", encoding="utf-8") as f:
        f.write(raw_text or "")

    candidate = _extract_json_candidate(raw_text)
    try:
        return json.loads(candidate)
    except Exception as e:
        last_err = e

    fix_prompt = (
        "Repair the following malformed JSON. Return valid JSON only. "
        "Do not add explanations or markdown."
    )
    current_text = raw_text
    for i in range(max_fix_attempts):
        data = _grok_chat(
            grok_api_key,
            messages=[
                {"role": "system", "content": "You are a JSON repair engine. Return valid JSON only."},
                {"role": "user", "content": fix_prompt + "\n\n" + (current_text or "")},
            ],
            model=repair_model,
            temperature=repair_temperature,
            max_tokens=2500,
        )
        repaired = data["choices"][0]["message"]["content"]
        repaired_candidate = _extract_json_candidate(repaired)
        with open(os.path.join(DEBUG_LLM_DIR, f"{debug_tag}_repaired_attempt{i+1}.txt"), "w", encoding="utf-8") as f:
            f.write(repaired or "")
        try:
            return json.loads(repaired_candidate)
        except Exception as e:
            last_err = e
            current_text = repaired

    raise ValueError(f"Could not parse JSON after repair attempts: {last_err}")


def _load_docx_text(path: str) -> str:
    """Load text from docx with a pure-XML fallback (no lxml dependency)."""
    try:
        from docx import Document  # type: ignore

        doc = Document(path)
        parts: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for t in doc.tables:
            for row in t.rows:
                vals = [((c.text or "").strip().replace("\n", " ")) for c in row.cells]
                vals = [v for v in vals if v]
                if vals:
                    parts.append(" | ".join(vals))
        if parts:
            return "\n".join(parts)
    except Exception:
        pass

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    body = root.find(W + "body")
    parts = []
    if body is None:
        return ""

    def _text(el: ET.Element) -> str:
        return "".join((t.text or "") for t in el.iter(W + "t")).strip()

    for ch in body:
        if ch.tag == W + "p":
            t = _text(ch)
            if t:
                parts.append(t)
        elif ch.tag == W + "tbl":
            for tr in ch.findall(".//" + W + "tr"):
                row_vals: List[str] = []
                for tc in tr.findall(W + "tc"):
                    cell_ps = tc.findall(".//" + W + "p")
                    cell_text = " ".join([_text(p) for p in cell_ps if _text(p)])
                    if cell_text:
                        row_vals.append(cell_text)
                if row_vals:
                    parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def _slugify(s: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "_", (s or "").lower())
    return s.strip("_")[:50] or "criterion"


def parse_precis_rubric_criteria(docx_path: str) -> List[Dict[str, Any]]:
    """Parse criteria + marks from Precis Rubric.docx table."""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def _txt(el: ET.Element) -> str:
        return "".join((t.text or "") for t in el.iter(W + "t")).strip()

    rows: List[List[str]] = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            xml = z.read("word/document.xml")
        root = ET.fromstring(xml)
        body = root.find(W + "body")
        if body is None:
            return DEFAULT_PRECIS_CRITERIA

        for tbl in body.findall(W + "tbl"):
            for tr in tbl.findall(".//" + W + "tr"):
                cells: List[str] = []
                for tc in tr.findall(W + "tc"):
                    cell_parts = []
                    for p in tc.findall(".//" + W + "p"):
                        t = _txt(p)
                        if t:
                            cell_parts.append(t)
                    cells.append(" ".join(cell_parts).strip())
                if any(cells):
                    rows.append(cells)
    except Exception:
        return DEFAULT_PRECIS_CRITERIA

    parsed: List[Dict[str, Any]] = []
    for row in rows:
        if len(row) < 2:
            continue
        row_text = " ".join(row)
        if "criterion" in row_text.lower() and "marks" in row_text.lower():
            continue

        marks_raw = row[-1].strip() if row else ""
        m = re.search(r"(\d+(?:\.\d+)?)", marks_raw)
        if not m:
            continue
        marks_val = float(m.group(1))
        if marks_val <= 0:
            continue
        if marks_val > 5:
            if "subtotal" in row_text.lower() or "total" in row_text.lower():
                continue

        crit_raw = row[0].strip()
        if not crit_raw:
            continue
        crit = re.sub(r"^[IVXLCM]+\.?\s*", "", crit_raw, flags=re.IGNORECASE)
        crit = re.sub(r"^\d+\.?\s*", "", crit).strip()
        if not crit:
            continue

        parsed.append(
            {
                "id": _slugify(crit),
                "criterion": crit,
                "marks_allocated": int(round(marks_val)) if abs(marks_val - round(marks_val)) < 1e-6 else marks_val,
            }
        )

    if len(parsed) < 6:
        return DEFAULT_PRECIS_CRITERIA

    parsed = parsed[:8]
    for c in parsed:
        if c["criterion"].lower().startswith("title"):
            c["criterion"] = "Title"
            c["id"] = "title"
    return parsed


def load_environment(env_file: str) -> Tuple[str, DocumentAnalysisClient]:
    load_dotenv(env_file)

    grok_key = os.getenv("Grok_API")
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    azure_key = os.getenv("AZURE_KEY")

    missing = []
    if not grok_key:
        missing.append("Grok_API")
    if not azure_endpoint:
        missing.append("AZURE_ENDPOINT")
    if not azure_key:
        missing.append("AZURE_KEY")

    if missing:
        raise EnvironmentError(
            f"Missing env vars in {env_file}: {', '.join(missing)}"
        )

    doc_client = DocumentAnalysisClient(
        endpoint=azure_endpoint,
        credential=AzureKeyCredential(azure_key),
    )
    return grok_key, doc_client


def validate_input_paths(pdf_path: str, output_json_path: str, output_pdf_path: str) -> None:
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    with open(pdf_path, "rb") as f:
        if f.read(4) != b"%PDF":
            raise ValueError(f"Not a valid PDF: {pdf_path}")

    for outp in [output_json_path, output_pdf_path]:
        out_dir = os.path.dirname(outp)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)


def run_ocr_on_pdf(
    doc_client: DocumentAnalysisClient,
    pdf_path: str,
    *,
    workers: int = 2,
    render_dpi: int = 220,
) -> Dict[str, Any]:
    """Page-wise Azure OCR with retries on payload size."""

    def _encode_page(pil_img: Image.Image, scale: float, quality: int) -> bytes:
        img = pil_img.copy()
        if scale != 1.0:
            img = img.resize((max(1, int(img.width * scale)), max(1, int(img.height * scale))), Image.LANCZOS)
        if img.mode != "RGB":
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality, optimize=True)
        return buf.getvalue()

    def _analyze(img_bytes: bytes) -> Any:
        poller = doc_client.begin_analyze_document("prebuilt-read", document=img_bytes)
        return poller.result()

    doc = fitz.open(pdf_path)
    try:
        pil_pages: List[Tuple[int, Image.Image]] = []
        for idx in range(doc.page_count):
            pix = doc[idx].get_pixmap(dpi=render_dpi)
            pil_pages.append((idx + 1, Image.open(io.BytesIO(pix.tobytes("png")))))
    finally:
        doc.close()

    def _process(page_no: int, pil_img: Image.Image) -> Dict[str, Any]:
        attempts = [(1.0, 75), (0.85, 70), (0.7, 60)]
        result = None
        last_err: Optional[Exception] = None

        for scale, quality in attempts:
            try:
                img_bytes = _encode_page(pil_img, scale=scale, quality=quality)
                result = _analyze(img_bytes)
                break
            except HttpResponseError as e:
                last_err = e
                if "InvalidContentLength" in str(e):
                    continue
                raise

        if result is None:
            raise RuntimeError(f"OCR failed for page {page_no}: {last_err}")

        page_text_parts: List[str] = []
        lines_out: List[Dict[str, Any]] = []
        words_out: List[Dict[str, Any]] = []

        for p in result.pages or []:
            for w in p.words or []:
                wtxt = (w.content or "").strip()
                if not wtxt:
                    continue
                poly = []
                if w.polygon:
                    poly = [(int(pt.x), int(pt.y)) for pt in w.polygon]
                words_out.append(
                    {
                        "text": wtxt,
                        "bbox": poly,
                        "confidence": float(getattr(w, "confidence", 1.0) or 1.0),
                    }
                )
            for ln in p.lines or []:
                ltxt = (ln.content or "").strip()
                if not ltxt:
                    continue
                lpoly = []
                if ln.polygon:
                    lpoly = [(int(pt.x), int(pt.y)) for pt in ln.polygon]
                lines_out.append({"text": ltxt, "bbox": lpoly})
                page_text_parts.append(ltxt)

        return {
            "page_number": page_no,
            "ocr_page_text": " ".join(page_text_parts).strip(),
            "lines": lines_out,
            "words": words_out,
        }

    pages: List[Dict[str, Any]] = []
    worker_count = max(1, int(workers or 1))

    with ThreadPoolExecutor(max_workers=worker_count) as ex:
        futures = {ex.submit(_process, pno, img): pno for pno, img in pil_pages}
        for fut in as_completed(futures):
            pages.append(fut.result())

    pages.sort(key=lambda x: x.get("page_number", 0))
    full_text = "\n".join([p.get("ocr_page_text", "") for p in pages if p.get("ocr_page_text")]).strip()
    return {"pages": pages, "full_text": full_text}


def _is_noise_text(text: str, bbox: List[Tuple[int, int]], page_w: float, page_h: float) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    if len(t) <= 1:
        return True
    if len(re.sub(r"[^A-Za-z0-9]", "", t)) <= 1:
        return True
    if bbox:
        xs = [p[0] for p in bbox]
        ys = [p[1] for p in bbox]
        if xs and ys and page_w > 0 and page_h > 0:
            w = max(xs) - min(xs)
            h = max(ys) - min(ys)
            rel_w = w / max(1.0, page_w)
            rel_h = h / max(1.0, page_h)
            cx = (min(xs) + max(xs)) / 2.0
            cy = (min(ys) + max(ys)) / 2.0
            edge = (cx < page_w * 0.06 or cx > page_w * 0.94 or cy < page_h * 0.06 or cy > page_h * 0.94)
            if rel_w < 0.002 or rel_h < 0.002:
                return True
            if edge and len(t) < 20:
                return True
    return False


def split_extra_text(ocr_data: Dict[str, Any], pdf_path: str) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    """Remove likely watermark/camera/noise OCR lines and return cleaned OCR + extra text pack."""
    page_dims: Dict[int, Tuple[float, float]] = {}
    try:
        doc = fitz.open(pdf_path)
        for i, pg in enumerate(doc):
            r = pg.rect
            page_dims[i + 1] = (float(r.width), float(r.height))
        doc.close()
    except Exception:
        pass

    cleaned_pages: List[Dict[str, Any]] = []
    extra_items: List[Dict[str, Any]] = []

    for p in (ocr_data.get("pages") or []):
        pno = int(p.get("page_number") or 0)
        pw, ph = page_dims.get(pno, (0.0, 0.0))
        kept_lines: List[Dict[str, Any]] = []
        removed_lines: List[Dict[str, Any]] = []
        for ln in (p.get("lines") or []):
            ltxt = (ln.get("text") or "").strip()
            lbbox = ln.get("bbox") or []
            if _is_noise_text(ltxt, lbbox, pw, ph):
                removed_lines.append({"text": ltxt, "bbox": lbbox})
            else:
                kept_lines.append(ln)

        kept_text = " ".join([(x.get("text") or "").strip() for x in kept_lines if (x.get("text") or "").strip()]).strip()
        cp = dict(p)
        cp["lines"] = kept_lines
        cp["ocr_page_text"] = kept_text
        cleaned_pages.append(cp)

        if removed_lines:
            extra_items.append({"page_number": pno, "removed_lines": removed_lines})

    cleaned = {
        "pages": cleaned_pages,
        "full_text": "\n".join([p.get("ocr_page_text", "") for p in cleaned_pages if p.get("ocr_page_text")]).strip(),
    }
    extras = {
        "pdf_path": pdf_path,
        "removed_line_count": sum(len(x.get("removed_lines", [])) for x in extra_items),
        "pages": extra_items,
    }
    return cleaned, extras

def pdf_to_page_images_for_grok(
    pdf_path: str,
    max_pages: Optional[int] = None,
    output_dir: str = "grok_images_precis",
    max_dim: int = 850,
    max_total_base64_chars: int = 280_000,
) -> List[Dict[str, Any]]:
    """Render PDF pages to compact base64 JPEG chunks for Grok."""
    os.makedirs(output_dir, exist_ok=True)

    doc = fitz.open(pdf_path)
    try:
        total_pages = doc.page_count if max_pages is None else min(max_pages, doc.page_count)
        pil_pages: List[Image.Image] = []
        for idx in range(total_pages):
            pix = doc[idx].get_pixmap(dpi=200)
            pil_pages.append(Image.open(io.BytesIO(pix.tobytes("png"))))
    finally:
        doc.close()

    dim_candidates = [max_dim, 720, 640, 560, 512, 448, 384]
    quality_candidates = [68, 60, 52, 45]

    def _encode(dim: int, quality: int, write_files: bool) -> Tuple[List[Dict[str, Any]], int]:
        out: List[Dict[str, Any]] = []
        total = 0
        for i, pil_img in enumerate(pil_pages):
            img = pil_img.copy()
            img.thumbnail((dim, dim))
            if img.mode in ("RGBA", "LA", "P"):
                rgb = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                rgb.paste(img, mask=img.split()[-1] if img.mode in ("RGBA", "LA") else None)
                img = rgb
            elif img.mode != "RGB":
                img = img.convert("RGB")

            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=quality, optimize=True)
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            total += len(b64)

            file_path = None
            if write_files:
                file_path = os.path.join(output_dir, f"page_{i+1:03d}.jpg")
                with open(file_path, "wb") as f:
                    f.write(buf.getvalue())

            out.append({"page": i + 1, "image_base64": b64, "file_path": file_path, "truncated": False})
        return out, total

    best: Optional[Tuple[List[Dict[str, Any]], int, int, int]] = None
    for dim in dim_candidates:
        for quality in quality_candidates:
            pages, total = _encode(dim, quality, write_files=False)
            best = (pages, total, dim, quality)
            if total > max_total_base64_chars:
                continue
            final_pages, final_total = _encode(dim, quality, write_files=True)
            print(
                f"Saved {len(final_pages)} Grok page images to '{output_dir}/' "
                f"(dim={dim}, quality={quality}, chars={final_total})"
            )
            return final_pages

    if best:
        _, _, dim, quality = best
        final_pages, final_total = _encode(dim, quality, write_files=True)
        print(
            f"Saved {len(final_pages)} Grok page images using fallback settings "
            f"(dim={dim}, quality={quality}, chars={final_total})."
        )
        return final_pages

    return []


def pil_images_to_pdf_bytes(pages: List[Image.Image]) -> bytes:
    out = io.BytesIO()
    if not pages:
        return b""
    pages_rgb = [p.convert("RGB") for p in pages]
    pages_rgb[0].save(out, format="PDF", save_all=True, append_images=pages_rgb[1:])
    return out.getvalue()


def merge_report_and_annotated_answer(
    report_pdf_path: str,
    annotated_pages: List[Image.Image],
    output_pdf_path: str,
) -> None:
    out_doc = fitz.open()
    target_w = 595.0
    target_h = 842.0
    if os.path.exists(report_pdf_path):
        rdoc = fitz.open(report_pdf_path)
        if len(rdoc) > 0:
            r0 = rdoc[0].rect
            target_w, target_h = float(r0.width), float(r0.height)
        out_doc.insert_pdf(rdoc)
        rdoc.close()

    # Keep annotated-answer pages on the SAME page size as report pages.
    for img in annotated_pages:
        page = out_doc.new_page(width=target_w, height=target_h)
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="PNG")
        stream = buf.getvalue()
        # Preserve aspect ratio; top-align to avoid empty space above annotations.
        img_w, img_h = img.size
        if img_w <= 0 or img_h <= 0:
            continue
        scale = min(target_w / img_w, target_h / img_h)
        draw_w = img_w * scale
        draw_h = img_h * scale
        x0 = (target_w - draw_w) / 2.0
        y0 = 0.0
        rect = fitz.Rect(x0, y0, x0 + draw_w, y0 + draw_h)
        page.insert_image(rect, stream=stream)

    out_doc.save(output_pdf_path)
    out_doc.close()


def _normalize_rating(score: float, max_marks: float) -> str:
    ratio = 0.0 if max_marks <= 0 else score / max_marks
    if ratio >= 0.85:
        return "Excellent"
    if ratio >= 0.70:
        return "Good"
    if ratio >= 0.50:
        return "Average"
    return "Weak"


def _infer_length_status(original_words: int, required_words: int, student_words: int) -> str:
    if original_words <= 0 or required_words <= 0:
        return "Unknown"
    lower = int(round(required_words * 0.95))
    upper = int(round(required_words * 1.05))
    if lower <= student_words <= upper:
        return "Within +/-5%"
    if student_words < lower:
        return "Too Short"
    return "Too Long"


def call_grok_for_precis_grading(
    grok_api_key: str,
    rubric_text: str,
    criteria_template: List[Dict[str, Any]],
    ocr_data: Dict[str, Any],
    page_images: List[Dict[str, Any]],
    *,
    model: str = "grok-4-1-fast-reasoning",
    temperature: float = 0.10,
    repair_model: str = "grok-4-1-fast-reasoning",
    repair_temperature: float = 0.0,
) -> Dict[str, Any]:
    total_marks = int(round(sum(float(c.get("marks_allocated", 0)) for c in criteria_template)))

    schema_hint = {
        "topic": "",
        "student_title": "",
        "student_precis_text": "",
        "original_passage_word_count": 0,
        "required_precis_word_count": 0,
        "student_precis_word_count": 0,
        "length_status": "Within +/-5% | Too Long | Too Short | Unknown",
        "criteria": [
            {
                "id": c.get("id"),
                "criterion": c.get("criterion"),
                "marks_allocated": c.get("marks_allocated"),
                "marks_awarded": 0,
                "rating": "Weak",
                "key_comments": "",
            }
            for c in criteria_template
        ],
        "total_awarded": 0,
        "overall_rating": "Weak",
        "reasons_for_low_score": [""],
        "ideal_precis": {
            "title": "",
            "text": "",
        },
        "overall_remarks": "",
    }

    system = {
        "role": "system",
        "content": (
            "You are a strict CSS precis examiner. "
            "Return valid JSON only with no markdown or commentary."
        ),
    }

    instructions = (
        "Evaluate a 2-page precis submission.\n"
        "Document layout expectation:\n"
        "- Page 1 contains question/prompt and original source passage.\n"
        "- Page 2 contains student answer: title + one precis paragraph.\n"
        "Required length rule:\n"
        "- Required precis length is exactly one-third of original passage word count.\n"
        "- Also classify length_status using +/-5% tolerance around required length.\n"
        "Scoring rules:\n"
        "- Follow the provided rubric criteria and marks exactly.\n"
        "- For each criterion, give marks_awarded within [0, marks_allocated].\n"
        "- Add concise, evidence-based key_comments for each criterion.\n"
        "- Provide total_awarded as the sum of marks_awarded values.\n"
        "- Provide overall_rating from: Excellent, Good, Average, Weak.\n"
        "- reasons_for_low_score must contain only concrete weaknesses, not praise.\n"
        "- ideal_precis.text must be a high-quality improved precis for this same passage.\n"
        "- ideal_precis.title must be concise and relevant.\n"
        "Extraction rules:\n"
        "- topic should be the passage topic/theme.\n"
        "- student_title must be exactly what the student wrote when visible; else infer carefully.\n"
        "- student_precis_text must contain the student precis paragraph only.\n"
        "- Compute original_passage_word_count, required_precis_word_count, student_precis_word_count.\n"
        "- Ignore unrelated watermark/camera/footer artifacts or stray words not part of question/answer content.\n"
        "- Do not mention OCR or handwriting quality in comments.\n"
        "Return JSON strictly matching the schema."
    )

    payload = {
        "rubric_text": rubric_text,
        "criteria_template": criteria_template,
        "ocr_pages": ocr_data.get("pages", []),
        "ocr_full_text": ocr_data.get("full_text", ""),
        "page_images": page_images,
        "output_schema": schema_hint,
    }

    def _validate(parsed: Dict[str, Any]) -> bool:
        crit = parsed.get("criteria")
        if not isinstance(crit, list) or len(crit) != len(criteria_template):
            return False

        marks_sum = 0.0
        for i, c in enumerate(crit):
            alloc = float(criteria_template[i].get("marks_allocated", 0))
            try:
                aw = float(c.get("marks_awarded", 0))
            except Exception:
                return False
            aw = min(max(aw, 0.0), alloc)
            c["marks_awarded"] = round(aw, 2)
            if not c.get("rating"):
                c["rating"] = _normalize_rating(aw, alloc)
            marks_sum += aw

        declared_total = parsed.get("total_awarded")
        try:
            declared_total_f = float(declared_total)
        except Exception:
            declared_total_f = marks_sum

        if abs(declared_total_f - marks_sum) > 0.75:
            parsed["total_awarded"] = round(marks_sum, 2)
        else:
            parsed["total_awarded"] = round(declared_total_f, 2)

        parsed["total_awarded"] = max(0.0, min(float(total_marks), float(parsed["total_awarded"])))

        if parsed.get("overall_rating") not in ("Excellent", "Good", "Average", "Weak"):
            parsed["overall_rating"] = _normalize_rating(float(parsed["total_awarded"]), float(total_marks))

        ow = int(parsed.get("original_passage_word_count") or 0)
        rw = int(parsed.get("required_precis_word_count") or 0)
        sw = int(parsed.get("student_precis_word_count") or 0)
        if rw <= 0 and ow > 0:
            parsed["required_precis_word_count"] = int(round(ow / 3.0))
            rw = int(parsed["required_precis_word_count"])

        if not parsed.get("length_status"):
            parsed["length_status"] = _infer_length_status(ow, rw, sw)

        reasons = parsed.get("reasons_for_low_score")
        if not isinstance(reasons, list):
            parsed["reasons_for_low_score"] = []
        parsed["reasons_for_low_score"] = [str(x).strip() for x in parsed["reasons_for_low_score"] if str(x).strip()][:8]

        ideal = parsed.get("ideal_precis")
        if not isinstance(ideal, dict):
            parsed["ideal_precis"] = {"title": "", "text": ""}
        parsed["ideal_precis"]["title"] = str(parsed["ideal_precis"].get("title", "")).strip()
        parsed["ideal_precis"]["text"] = str(parsed["ideal_precis"].get("text", "")).strip()
        return True

    last_err: Optional[Exception] = None
    for attempt in range(4):
        print(f"  Precis grading attempt {attempt + 1}/4...")
        response = _grok_chat(
            grok_api_key,
            messages=[
                system,
                {
                    "role": "user",
                    "content": instructions + "\n\nDATA:\n" + json.dumps(payload, ensure_ascii=False),
                },
            ],
            model=model,
            temperature=temperature,
            max_tokens=4000,
        )
        content = response["choices"][0]["message"]["content"]
        try:
            parsed = parse_json_with_repair(
                grok_api_key,
                content,
                debug_tag="precis_grading",
                max_fix_attempts=3,
                repair_model=repair_model,
                repair_temperature=repair_temperature,
            )
        except Exception as e:
            last_err = e
            continue

        if _validate(parsed):
            print(f"  Precis grading validated on attempt {attempt + 1}.")
            return parsed
        last_err = ValueError("Invalid grading JSON")

    raise RuntimeError(f"Precis grading failed after retries: {last_err}")


def call_grok_for_precis_annotations(
    grok_api_key: str,
    annotations_rubric_text: str,
    ocr_data: Dict[str, Any],
    grading: Dict[str, Any],
    page_images: List[Dict[str, Any]],
    *,
    model: str = "grok-4-1-fast-reasoning",
    temperature: float = 0.15,
    repair_model: str = "grok-4-1-fast-reasoning",
    repair_temperature: float = 0.0,
) -> Dict[str, Any]:
    """
    Returns:
    {
      "page_suggestions": [{"page": 2, "suggestions": ["..."]}],
      "annotations": [ ... ],
      "errors": [ ... ]
    }
    """
    system = {
        "role": "system",
        "content": (
            "You generate precise, locatable annotations for handwritten precis scripts.\n"
            "Primary truth is page image; OCR is helper text only.\n"
            "Never mention OCR, scanning, or handwriting quality.\n"
            "Return JSON only."
        ),
    }

    schema_hint = {
        "page": 2,
        "page_suggestions": ["..."],
        "annotations": [
            {
                "page": 2,
                "type": "language_clarity",
                "rubric_point": "string",
                "anchor_quote": "EXACT substring from OCR_PAGE_TEXT",
                "correction": "string",
                "comment": "string",
            }
        ],
    }

    instructions = (
        "Using the provided annotation rubric text, generate concise actionable annotations for ONE page.\n"
        "Rules:\n"
        "- Prefer 2-5 annotations for the answer page.\n"
        "- Every annotation must be locatable from anchor_quote.\n"
        "- anchor_quote must be an exact contiguous substring from OCR_PAGE_TEXT.\n"
        "- If anchor cannot be found, skip that annotation.\n"
        "- Keep comments concise and corrective.\n"
        "- page_suggestions should be 2-4 concise actionable bullets for this page.\n"
        "- Ignore unrelated watermark/camera/footer artifacts or stray words not part of the answer.\n"
        "- Never mention OCR/scanning/handwriting.\n"
        "Return JSON matching schema."
    )

    os.makedirs(DEBUG_LLM_DIR, exist_ok=True)
    errors: List[Dict[str, Any]] = []
    annotations: List[Dict[str, Any]] = []
    page_suggestions: List[Dict[str, Any]] = []

    image_by_page = {p.get("page"): p for p in page_images}
    # For precis, prioritize answer page(s): page >= 2.
    ocr_pages = [p for p in (ocr_data.get("pages") or []) if int(p.get("page_number") or 0) >= 2]

    for page in ocr_pages:
        page_num = int(page.get("page_number") or 0)
        if page_num <= 0:
            continue
        ocr_page_text = (page.get("ocr_page_text") or "").strip()
        if not ocr_page_text:
            errors.append({"page": page_num, "error": "Missing ocr_page_text"})
            continue

        payload = {
            "annotations_rubric_text": annotations_rubric_text or "",
            "grading_summary": {
                "total_awarded": grading.get("total_awarded"),
                "criteria": grading.get("criteria", []),
                "reasons_for_low_score": grading.get("reasons_for_low_score", []),
            },
            "ocr_page": {
                "page_number": page_num,
                "ocr_page_text": ocr_page_text,
                "lines": page.get("lines", []),
            },
            "page_image": image_by_page.get(page_num),
            "output_schema": schema_hint,
        }

        parsed: Optional[Dict[str, Any]] = None
        last_err: Optional[Exception] = None
        for _ in range(3):
            try:
                resp = _grok_chat(
                    grok_api_key,
                    messages=[system, {"role": "user", "content": instructions + "\n\nDATA:\n" + json.dumps(payload, ensure_ascii=False)}],
                    model=model,
                    temperature=temperature,
                    max_tokens=2500,
                    timeout=200,
                    max_retries=4,
                )
                content = resp["choices"][0]["message"]["content"]
                parsed = parse_json_with_repair(
                    grok_api_key,
                    content,
                    debug_tag=f"precis_annotations_p{page_num}",
                    max_fix_attempts=2,
                    repair_model=repair_model,
                    repair_temperature=repair_temperature,
                )
                if not isinstance(parsed, dict):
                    raise ValueError("Annotation JSON is not object")
                if not isinstance(parsed.get("annotations"), list):
                    raise ValueError("Missing annotations list")
                if not isinstance(parsed.get("page_suggestions"), list):
                    raise ValueError("Missing page_suggestions list")
                break
            except Exception as e:
                last_err = e
                parsed = None

        if parsed is None:
            errors.append({"page": page_num, "error": str(last_err) if last_err else "unknown"})
            continue

        cleaned_ann: List[Dict[str, Any]] = []
        for a in (parsed.get("annotations") or []):
            if not isinstance(a, dict):
                continue
            aq = str(a.get("anchor_quote", "")).strip()
            if not aq or aq not in ocr_page_text:
                continue
            cleaned_ann.append(
                {
                    "page": page_num,
                    "type": str(a.get("type", "")).strip(),
                    "rubric_point": str(a.get("rubric_point", "")).strip(),
                    "anchor_quote": aq,
                    "target_word_or_sentence": "",
                    "context_before": "",
                    "context_after": "",
                    "correction": str(a.get("correction", "")).strip(),
                    "comment": str(a.get("comment", "")).strip(),
                }
            )
        annotations.extend(cleaned_ann)
        page_suggestions.append({"page": page_num, "suggestions": [str(x).strip() for x in (parsed.get("page_suggestions") or []) if str(x).strip()]})

    return {"annotations": annotations, "page_suggestions": page_suggestions, "errors": errors}

def _dominant_colors_from_scheme(image_path: str) -> Dict[str, Tuple[float, float, float]]:
    fallback = {
        "header_fill": (95 / 255.0, 110 / 255.0, 141 / 255.0),
        "header_text": (1.0, 1.0, 1.0),
        "row_alt": (0.94, 0.95, 0.97),
        "border": (0.35, 0.40, 0.50),
        "section_title": (0.16, 0.20, 0.30),
    }
    if not image_path or not os.path.exists(image_path):
        return fallback
    try:
        img = Image.open(image_path).convert("RGB")
        img.thumbnail((320, 320))
        px = list(img.getdata())

        def _is_whiteish(rgb: Tuple[int, int, int]) -> bool:
            return rgb[0] > 240 and rgb[1] > 240 and rgb[2] > 240

        non_white = [p for p in px if not _is_whiteish(p)]
        if not non_white:
            return fallback

        buckets: Dict[Tuple[int, int, int], int] = {}
        for r, g, b in non_white:
            key = (r // 16 * 16, g // 16 * 16, b // 16 * 16)
            buckets[key] = buckets.get(key, 0) + 1

        sorted_buckets = sorted(buckets.items(), key=lambda kv: kv[1], reverse=True)
        dark = next((k for k, _ in sorted_buckets if (k[0] + k[1] + k[2]) < 360), None)
        mid = next((k for k, _ in sorted_buckets if 360 <= (k[0] + k[1] + k[2]) <= 660), None)

        if dark:
            fallback["header_fill"] = tuple([v / 255.0 for v in dark])  # type: ignore
            fallback["section_title"] = tuple([max(0, min(1, (v - 25) / 255.0)) for v in dark])  # type: ignore
            fallback["border"] = tuple([max(0, min(1, (v - 15) / 255.0)) for v in dark])  # type: ignore
        if mid:
            fallback["row_alt"] = tuple([min(1.0, (v + 45) / 255.0) for v in mid])  # type: ignore
        return fallback
    except Exception:
        return fallback


def _wrap_lines(text: str, fontname: str, fontsize: float, max_width: float) -> List[str]:
    text = (text or "").strip()
    if not text:
        return [""]
    words = text.split()
    lines: List[str] = []
    cur = ""
    for w in words:
        test = (cur + " " + w).strip()
        if fitz.get_text_length(test, fontname=fontname, fontsize=fontsize) <= max_width or not cur:
            cur = test
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def _draw_wrapped_text(
    page: fitz.Page,
    x: float,
    y: float,
    text: str,
    *,
    fontname: str,
    fontsize: float,
    max_width: float,
    color: Tuple[float, float, float],
    line_gap: float = 1.3,
) -> float:
    lines = _wrap_lines(text, fontname=fontname, fontsize=fontsize, max_width=max_width)
    for ln in lines:
        page.insert_text((x, y), ln, fontname=fontname, fontsize=fontsize, color=color)
        y += fontsize * line_gap
    return y


def render_precis_report_pdf(
    grading: Dict[str, Any],
    output_pdf_path: str,
    *,
    colouring_scheme_image: str = "",
    max_pages: int = 2,
) -> None:
    """Render report on exactly one page by shrinking all text sizes if needed."""
    palette = _dominant_colors_from_scheme(colouring_scheme_image)
    W, H = 595.0, 842.0
    margin = 30.0

    total_awarded = grading.get("total_awarded", 0)
    total_marks = int(round(sum(float(c.get("marks_allocated", 0)) for c in grading.get("criteria", [])))) or 20
    fields: List[Tuple[str, str]] = [
        ("Student Title", str(grading.get("student_title", ""))),
        (
            "Word Counts",
            f"Original: {grading.get('original_passage_word_count', 0)} | Required (1/3): {grading.get('required_precis_word_count', 0)} | Student: {grading.get('student_precis_word_count', 0)}",
        ),
        ("Total Marks", f"{total_awarded}/{total_marks}"),
    ]

    criteria = grading.get("criteria") or []
    reasons = grading.get("reasons_for_low_score") or []
    if not reasons:
        reasons = ["No major weaknesses identified."]
    ideal = grading.get("ideal_precis") or {}
    ideal_title = str(ideal.get("title", "")).strip() or "(Not provided)"
    ideal_text = str(ideal.get("text", "")).strip() or "(Not provided)"

    def _sizes(shrink: int) -> Dict[str, float]:
        base = max(6.0, REPORT_BASE_TEXT_SIZE - float(shrink))
        return {
            "title": base,
            "field_label": base,
            "field_value": base,
            "table_header": base,
            "table_cell": base,
            "section": base,
            "bullet": base,
            "ideal_title": base,
        }

    def _render_once(doc: fitz.Document, shrink: int) -> bool:
        page = doc.new_page(width=W, height=H)
        s = _sizes(shrink)
        y = margin
        usable_h = H - margin

        def _need(h: float) -> bool:
            nonlocal y
            return (y + h) <= usable_h

        def _advance_after_text(font_size: float, extra: float = 0.0) -> None:
            nonlocal y
            y += font_size + extra

        # Title
        if not _need(s["title"] + 8):
            return False
        page.insert_text((margin, y + s["title"]), "Precis Evaluation Report", fontname="hebo", fontsize=s["title"], color=palette["section_title"])
        y += s["title"] + 8

        # Top fields (no bounding boxes)
        label_w = 120
        for label, value in fields:
            max_w = W - 2 * margin - label_w
            lines = _wrap_lines(str(value), "helv", s["field_value"], max_w)
            line_h = s["field_value"] * 1.2
            block_h = max(s["field_label"], line_h * max(1, len(lines))) + 4
            if not _need(block_h):
                return False
            page.insert_text((margin, y + s["field_label"]), f"{label}:", fontname="hebo", fontsize=s["field_label"], color=palette["section_title"])
            yy = y + s["field_value"]
            for ln in lines:
                page.insert_text((margin + label_w, yy), ln, fontname="helv", fontsize=s["field_value"], color=(0, 0, 0))
                yy += line_h
            y += block_h

        y += 4

        # Table
        headers = ["Criterion", "Marks Allocated", "Marks Awarded", "Key Comments"]
        col_w = [205.0, 78.0, 78.0, W - (margin * 2 + 205.0 + 78.0 + 78.0)]
        header_max_lines = max(len(_wrap_lines(h, "hebo", s["table_header"], col_w[i] - 6)) for i, h in enumerate(headers))
        header_h = max(16.0, header_max_lines * (s["table_header"] * 1.05) + 6)
        if not _need(header_h + 2):
            return False
        x = margin
        header_rect = fitz.Rect(margin, y, W - margin, y + header_h)
        page.draw_rect(header_rect, color=palette["border"], fill=palette["header_fill"], width=1)
        for i, htxt in enumerate(headers):
            h_lines = _wrap_lines(htxt, "hebo", s["table_header"], col_w[i] - 6)
            hy = y + s["table_header"] + 1
            for ln in h_lines[:2]:
                page.insert_text((x + 4, hy), ln, fontname="hebo", fontsize=s["table_header"], color=palette["header_text"])
                hy += s["table_header"] * 1.05
            x += col_w[i]
            if i < len(headers) - 1:
                page.draw_line((x, y), (x, y + header_h), color=palette["border"], width=1)
        y += header_h

        for idx, c in enumerate(criteria):
            crit = str(c.get("criterion", ""))
            alloc = str(c.get("marks_allocated", ""))
            award = str(c.get("marks_awarded", ""))
            comment = str(c.get("key_comments", ""))

            crit_lines = _wrap_lines(crit, "helv", s["table_cell"], col_w[0] - 6)
            cmt_lines = _wrap_lines(comment, "helv", s["table_cell"], col_w[3] - 6)
            line_h = s["table_cell"] * 1.2
            row_h = max(18.0, max(len(crit_lines), len(cmt_lines), 1) * line_h + 6)
            if not _need(row_h + 1):
                return False

            fill = palette["row_alt"] if idx % 2 == 0 else (1, 1, 1)
            row_rect = fitz.Rect(margin, y, W - margin, y + row_h)
            page.draw_rect(row_rect, color=palette["border"], fill=fill, width=0.7)
            x = margin

            yy = y + s["table_cell"] + 1
            for ln in crit_lines:
                page.insert_text((x + 4, yy), ln, fontname="helv", fontsize=s["table_cell"], color=(0, 0, 0))
                yy += line_h
            x += col_w[0]
            page.draw_line((x, y), (x, y + row_h), color=palette["border"], width=0.7)
            page.insert_text((x + 4, y + s["table_cell"] + 1), alloc, fontname="helv", fontsize=s["table_cell"], color=(0, 0, 0))
            x += col_w[1]
            page.draw_line((x, y), (x, y + row_h), color=palette["border"], width=0.7)
            page.insert_text((x + 4, y + s["table_cell"] + 1), award, fontname="helv", fontsize=s["table_cell"], color=(0, 0, 0))
            x += col_w[2]
            page.draw_line((x, y), (x, y + row_h), color=palette["border"], width=0.7)
            yy = y + s["table_cell"] + 1
            for ln in cmt_lines:
                page.insert_text((x + 4, yy), ln, fontname="helv", fontsize=s["table_cell"], color=(0, 0, 0))
                yy += line_h
            y += row_h

        y += 6

        # Reasons for low score
        if not _need(s["section"] + 6):
            return False
        page.insert_text((margin, y + s["section"]), "Reasons for Low Score", fontname="hebo", fontsize=s["section"], color=palette["section_title"])
        y += s["section"] + 4

        bullet_lh = s["bullet"] * 1.25
        for item in reasons[:8]:
            blines = _wrap_lines(f"- {str(item)}", "helv", s["bullet"], W - 2 * margin - 10)
            bh = max(1, len(blines)) * bullet_lh + 1
            if not _need(bh):
                return False
            yy = y + s["bullet"]
            for ln in blines:
                page.insert_text((margin + 8, yy), ln, fontname="helv", fontsize=s["bullet"], color=(0, 0, 0))
                yy += bullet_lh
            y += bh

        y += 4

        # Ideal precis (no bounding box)
        if not _need(s["section"] + 6):
            return False
        page.insert_text((margin, y + s["section"]), "Ideal Precis", fontname="hebo", fontsize=s["section"], color=palette["section_title"])
        y += s["section"] + 4

        title_lines = _wrap_lines(f"Title: {ideal_title}", "hebo", s["ideal_title"], W - 2 * margin - 8)
        title_lh = s["ideal_title"] * 1.2
        th = max(1, len(title_lines)) * title_lh + 2
        if not _need(th):
            return False
        yy = y + s["ideal_title"]
        for ln in title_lines:
            page.insert_text((margin + 4, yy), ln, fontname="hebo", fontsize=s["ideal_title"], color=(0, 0, 0))
            yy += title_lh
        y += th

        body_lines = _wrap_lines(ideal_text, "helv", s["bullet"], W - 2 * margin - 8)
        body_lh = s["bullet"] * 1.25
        bh = max(1, len(body_lines)) * body_lh + 2
        if not _need(bh):
            return False
        yy = y + s["bullet"]
        for ln in body_lines:
            page.insert_text((margin + 4, yy), ln, fontname="helv", fontsize=s["bullet"], color=(0, 0, 0))
            yy += body_lh
        return True

    # Try progressively smaller global text sizes until it fits one page.
    best_doc: Optional[fitz.Document] = None
    for shrink in range(0, 13):
        d = fitz.open()
        if _render_once(d, shrink):
            best_doc = d
            break
        d.close()

    if best_doc is None:
        # Last fallback at smallest size; keep first page only.
        best_doc = fitz.open()
        _render_once(best_doc, 12)

    os.makedirs(os.path.dirname(output_pdf_path) or ".", exist_ok=True)
    best_doc.save(output_pdf_path)
    best_doc.close()

def run_precis_grading(
    pdf_path: str,
    output_json_path: str,
    output_pdf_path: str,
    *,
    rubric_docx: str,
    annotations_rubric_docx: str,
    env_file: str,
    colouring_scheme_image: str,
    extra_json_path: str,
    report_only_pdf_path: Optional[str] = None,
    ocr_workers: int = 2,
    grading_model: str = DEFAULT_MODELS["grading"]["model"],
    grading_temperature: float = float(DEFAULT_MODELS["grading"]["temperature"]),
    annotations_model: str = DEFAULT_MODELS["annotations"]["model"],
    annotations_temperature: float = float(DEFAULT_MODELS["annotations"]["temperature"]),
    repair_model: str = DEFAULT_MODELS["json_repair"]["model"],
    repair_temperature: float = float(DEFAULT_MODELS["json_repair"]["temperature"]),
) -> Dict[str, Any]:
    validate_input_paths(pdf_path, output_json_path, output_pdf_path)
    grok_key, doc_client = load_environment(env_file)
    rubric_text = _load_docx_text(rubric_docx)
    annotations_rubric_text = _load_docx_text(annotations_rubric_docx)
    criteria_template = parse_precis_rubric_criteria(rubric_docx)

    timings: Dict[str, float] = {}
    t0_total = time.perf_counter()

    print("Running OCR on precis PDF...")
    t0 = time.perf_counter()
    ocr_data_raw = run_ocr_on_pdf(doc_client, pdf_path, workers=ocr_workers)
    timings["OCR"] = time.perf_counter() - t0
    print(f"OCR done in {_format_duration(timings['OCR'])}")

    t0 = time.perf_counter()
    ocr_data, extra_text_pack = split_extra_text(ocr_data_raw, pdf_path)
    os.makedirs(os.path.dirname(extra_json_path) or ".", exist_ok=True)
    with open(extra_json_path, "w", encoding="utf-8") as f:
        json.dump(extra_text_pack, f, indent=2, ensure_ascii=False)
    timings["Extra text filtering"] = time.perf_counter() - t0
    print(f"Extra text filtering done in {_format_duration(timings['Extra text filtering'])} "
          f"(removed {extra_text_pack.get('removed_line_count', 0)} lines)")

    print("Preparing page images for Grok...")
    t0 = time.perf_counter()
    page_images = pdf_to_page_images_for_grok(pdf_path, max_pages=2, output_dir=GROK_IMAGES_DIR)
    timings["Image prep"] = time.perf_counter() - t0
    print(f"Image prep done in {_format_duration(timings['Image prep'])}")

    print("Grading precis with rubric...")
    t0 = time.perf_counter()
    grading = call_grok_for_precis_grading(
        grok_key,
        rubric_text=rubric_text,
        criteria_template=criteria_template,
        ocr_data=ocr_data,
        page_images=page_images,
        model=grading_model,
        temperature=grading_temperature,
        repair_model=repair_model,
        repair_temperature=repair_temperature,
    )
    timings["LLM grading"] = time.perf_counter() - t0
    print(f"LLM grading done in {_format_duration(timings['LLM grading'])}")

    print("Generating precis annotations...")
    t0 = time.perf_counter()
    ann_pack = call_grok_for_precis_annotations(
        grok_key,
        annotations_rubric_text=annotations_rubric_text,
        ocr_data=ocr_data,
        grading=grading,
        page_images=page_images,
        model=annotations_model,
        temperature=annotations_temperature,
        repair_model=repair_model,
        repair_temperature=repair_temperature,
    )
    timings["Annotations"] = time.perf_counter() - t0
    print(f"Annotations done in {_format_duration(timings['Annotations'])}")

    annotations = ann_pack.get("annotations") or []
    page_suggestions = ann_pack.get("page_suggestions") or []
    ann_errors = ann_pack.get("errors") or []

    output = {
        "grading": grading,
        "criteria_template": criteria_template,
        "ocr_pages": len(ocr_data.get("pages", [])),
        "annotations": annotations,
        "page_suggestions": page_suggestions,
        "annotation_errors": ann_errors,
        "extra_text_json_path": extra_json_path,
        "model_config": {
            "grading": {"model": grading_model, "temperature": grading_temperature},
            "annotations": {"model": annotations_model, "temperature": annotations_temperature},
            "json_repair": {"model": repair_model, "temperature": repair_temperature},
        },
    }
    os.makedirs(os.path.dirname(output_json_path) or ".", exist_ok=True)
    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    print(f"Saved JSON -> {output_json_path}")

    print("Rendering precis report PDF...")
    t0 = time.perf_counter()
    report_tmp = report_only_pdf_path or os.path.join(os.path.dirname(output_pdf_path) or ".", "_precis_report_tmp.pdf")
    render_precis_report_pdf(
        grading,
        report_tmp,
        colouring_scheme_image=colouring_scheme_image,
        max_pages=1,
    )
    timings["PDF render"] = time.perf_counter() - t0
    print(f"PDF render done in {_format_duration(timings['PDF render'])}")

    print("Rendering annotated precis pages...")
    t0 = time.perf_counter()
    if annotate_pdf_essay_pages is None:
        raise RuntimeError("annotate_pdf_with_essay_rubric.py is required for annotation rendering.")

    annotated_pages = annotate_pdf_essay_pages(
        pdf_path=pdf_path,
        ocr_data=ocr_data,
        structure={"outline": {"present": False}, "paragraph_map": []},
        grading=grading,
        annotations=annotations,
        page_suggestions=page_suggestions,
        spelling_errors=None,
        max_callouts_per_page=8,
    )
    merge_report_and_annotated_answer(report_tmp, annotated_pages, output_pdf_path)
    if report_only_pdf_path is None:
        try:
            os.unlink(report_tmp)
        except Exception:
            pass
    timings["Merge output PDF"] = time.perf_counter() - t0
    print(f"Merge done in {_format_duration(timings['Merge output PDF'])}")

    total_time = time.perf_counter() - t0_total
    print("\n" + "=" * 60)
    print("PRECIS GRADING TIMING SUMMARY")
    print("=" * 60)
    for k, v in timings.items():
        print(f"  {k}: {_format_duration(v)}")
    print("-" * 60)
    print(f"  Total: {_format_duration(total_time)}")
    print("=" * 60)

    return {
        "status": "success",
        "json_path": output_json_path,
        "pdf_path": output_pdf_path,
        "grading": grading,
        "annotations": annotations,
        "extra_text": extra_text_pack,
        "timings": timings,
        "total_time": total_time,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Grade a precis PDF, render report, and append annotated precis pages.")
    parser.add_argument("--pdf", default=os.path.join("precis", "input.pdf"), help="Input precis PDF path")
    parser.add_argument("--output-json", default=os.path.join("precis", "precis_result.json"), help="Output JSON path")
    parser.add_argument("--output-pdf", default=os.path.join("precis", "output.pdf"), help="Output report PDF path")
    parser.add_argument("--rubric-docx", default=os.path.join("precis", "Precis Rubric.docx"), help="Precis rubric .docx path")
    parser.add_argument("--annotations-rubric-docx", default=os.path.join("precis", "ANNOTATIONS RUBRIC FOR PRECIS.docx"), help="Precis annotations rubric .docx path")
    parser.add_argument("--colouring-scheme-image", default=os.path.join("precis", "colouring_scheme.jpeg"), help="Colouring scheme image path")
    parser.add_argument("--env-file", default=os.path.join("precis", "env"), help="Env file path")
    parser.add_argument("--extra-json", default=os.path.join("precis", "extra_text.json"), help="Path to save removed extra/noise OCR text")
    parser.add_argument("--report-only-pdf", default="", help="Optional path to keep standalone report PDF")
    parser.add_argument("--ocr-workers", type=int, default=2, help="Parallel OCR worker count")
    parser.add_argument("--grading-model", default=DEFAULT_MODELS["grading"]["model"])
    parser.add_argument("--grading-temperature", type=float, default=float(DEFAULT_MODELS["grading"]["temperature"]))
    parser.add_argument("--annotations-model", default=DEFAULT_MODELS["annotations"]["model"])
    parser.add_argument("--annotations-temperature", type=float, default=float(DEFAULT_MODELS["annotations"]["temperature"]))
    parser.add_argument("--repair-model", default=DEFAULT_MODELS["json_repair"]["model"])
    parser.add_argument("--repair-temperature", type=float, default=float(DEFAULT_MODELS["json_repair"]["temperature"]))
    args = parser.parse_args()

    result = run_precis_grading(
        pdf_path=args.pdf,
        output_json_path=args.output_json,
        output_pdf_path=args.output_pdf,
        rubric_docx=args.rubric_docx,
        annotations_rubric_docx=args.annotations_rubric_docx,
        env_file=args.env_file,
        colouring_scheme_image=args.colouring_scheme_image,
        extra_json_path=args.extra_json,
        report_only_pdf_path=(args.report_only_pdf or None),
        ocr_workers=args.ocr_workers,
        grading_model=args.grading_model,
        grading_temperature=args.grading_temperature,
        annotations_model=args.annotations_model,
        annotations_temperature=args.annotations_temperature,
        repair_model=args.repair_model,
        repair_temperature=args.repair_temperature,
    )
    print(f"\nDone. Report PDF: {result['pdf_path']}")


if __name__ == "__main__":
    main()
