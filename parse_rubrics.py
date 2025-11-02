#!/usr/bin/env python3
"""Utility to parse rubric Word documents and extract structured data."""

from docx import Document
from pathlib import Path
import json

def parse_rubric_doc(docx_path: Path) -> dict:
    """Parse a rubric Word document and extract all content."""
    doc = Document(docx_path)

    result = {
        "subject": docx_path.stem,
        "file_path": str(docx_path),
        "paragraphs": [],
        "tables": [],
        "full_text": ""
    }

    # Extract all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            result["paragraphs"].append({
                "text": para.text.strip(),
                "style": para.style.name if para.style else "Normal"
            })

    # Extract all tables
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "table_number": table_idx + 1,
            "rows": []
        }
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data["rows"].append(row_data)
        result["tables"].append(table_data)

    # Full text
    result["full_text"] = "\n".join([p["text"] for p in result["paragraphs"]])

    return result

def analyze_rubric_structure(rubric_data: dict) -> dict:
    """Analyze the rubric structure to identify criteria, indicators, etc."""
    analysis = {
        "subject": rubric_data["subject"],
        "has_tables": len(rubric_data["tables"]) > 0,
        "num_tables": len(rubric_data["tables"]),
        "num_paragraphs": len(rubric_data["paragraphs"]),
        "structure": {}
    }

    # Look for common patterns
    full_text = rubric_data["full_text"].lower()

    patterns = {
        "has_indicators": "indicator" in full_text,
        "has_criteria": "criteri" in full_text or "rubric" in full_text,
        "has_marks": "marks" in full_text or "points" in full_text,
        "has_levels": "excellent" in full_text or "good" in full_text or "average" in full_text,
        "is_100_mark_paper": "100 marks paper" in full_text or "100 marks" in full_text,
        "is_20_mark_question": "20 marks" in full_text or "20-mark" in full_text,
    }

    analysis["patterns"] = patterns

    return analysis

if __name__ == "__main__":
    rubrics_dir = Path(__file__).parent / "Rubrics"

    # Find all rubric files
    rubric_files = list(rubrics_dir.glob("**/*.docx"))

    print(f"Found {len(rubric_files)} rubric files\n")

    # Parse first 3 for analysis
    for rubric_file in rubric_files[:3]:
        print(f"\n{'='*80}")
        print(f"PARSING: {rubric_file.stem}")
        print(f"{'='*80}\n")

        rubric_data = parse_rubric_doc(rubric_file)
        analysis = analyze_rubric_structure(rubric_data)

        print(f"Subject: {analysis['subject']}")
        print(f"Tables: {analysis['num_tables']}")
        print(f"Paragraphs: {analysis['num_paragraphs']}")
        print(f"\nPatterns:")
        for pattern, found in analysis['patterns'].items():
            print(f"  {pattern}: {found}")

        print(f"\n--- First 5 Paragraphs ---")
        for i, para in enumerate(rubric_data['paragraphs'][:5], 1):
            print(f"{i}. [{para['style']}] {para['text'][:150]}...")

        print(f"\n--- First Table (if exists) ---")
        if rubric_data['tables']:
            table = rubric_data['tables'][0]
            print(f"Rows: {len(table['rows'])}")
            for i, row in enumerate(table['rows'][:5], 1):
                print(f"Row {i}: {row[:3]}")  # First 3 cells

        print(f"\n--- Full Text Preview (first 500 chars) ---")
        print(rubric_data['full_text'][:500])
        print("\n")
