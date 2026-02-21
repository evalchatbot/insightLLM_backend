# grade_pdf_answer.py
#
# Simplified pipeline:
#   1) Google Vision OCR for text + bounding boxes.
#   2) Grok Prompt 1: Headings & structure detection (with PDF images).
#   3) Grok Prompt 2: Subject-wise marking (with PDF images + subject rubric DOCX).
#   4) Grok Prompt 3: Refined rubric annotations (with PDF images + refined rubric DOCX).
#   5) Render subject-wise report pages.
#   6) Render refined-rubric summary page.
#   7) Annotate answer pages according to simplified rules.
#   8) Merge all pages into final PDF.

import argparse
import base64
import io
import json
import os
import re
import sys
import tempfile
import time
import uuid
import shutil
from datetime import datetime
from functools import lru_cache
from typing import Any, Dict, List, Tuple, Optional, Set
import requests
from dotenv import load_dotenv
from google.api_core.client_options import ClientOptions
from google.cloud import vision
import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import cv2
try:
    from backend.ocr.annotate_pdf_with_rubric import annotate_pdf_answer_pages
except ImportError:
    from annotate_pdf_with_rubric import annotate_pdf_answer_pages


# -----------------------------
# UTILS & ENV
# -----------------------------




def debug_dump_sections(
    sections: List[Dict[str, Any]],
    output_path: str = "debug_sections.json",
) -> None:
    light_sections = []
    for idx, sec in enumerate(sections):
        light_sections.append(
            {
                "index": idx,
                "title": sec.get("title"),
                "exact_ocr_heading": sec.get("exact_ocr_heading"),
                "level": sec.get("level"),
                "page_numbers": sec.get("page_numbers"),
                "content_preview": (sec.get("content") or "")[:200],
                "comment": sec.get("comment"),
            }
        )

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(light_sections, f, ensure_ascii=False, indent=2)

    print("\n==== DETECTED HEADINGS / SECTIONS (from Grok) ====")
    for sec in light_sections:
        comment_display = f" | Comment: {sec['comment']}" if sec.get('comment') else ""
        print(
            f"[{sec['index']}] "
            f"Title: {sec['title']!r} | "
            f"exact_ocr_heading: {sec['exact_ocr_heading']!r} | "
            f"Level: {sec['level']} | "
            f"Pages: {sec['page_numbers']}"
            f"{comment_display}"
        )
    print(f"Saved detailed section info to {output_path}")
    print("=================================================\n")


def clean_json_from_llm(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z0-9]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
    return text.strip()


# -----------------------------
# JSON SCHEMA VALIDATION
# -----------------------------


def validate_refined_summary(summary_list: List[Dict[str, Any]]) -> bool:

    REQUIRED_KEYS = ["id", "name", "rating", "comment"]
    VALID_RATINGS = {"weak", "average", "good", "excellent"}
    VALID_IDS = {"argumentation_quality", "presentation",
                 "contemporary_relevance", "length_completeness"}

    for idx, item in enumerate(summary_list):
        # Check required keys
        missing = set(REQUIRED_KEYS) - set(item.keys())
        if missing:
            raise ValueError(f"refined_rubric_summary[{idx}] missing fields: {missing}")

        # Validate rating
        if item["rating"].lower() not in VALID_RATINGS:
            print(f"WARNING: Invalid rating in summary[{idx}]: {item['rating']} (expected: weak/average/good/excellent)")

        # Validate ID (warn but don't fail)
        if item["id"] not in VALID_IDS:
            print(f"WARNING: Unexpected summary ID in summary[{idx}]: {item['id']}")

    return True


def validate_annotation(annotation: Dict[str, Any], idx: int = 0) -> bool:
    """Validate single annotation schema."""
    REQUIRED_KEYS = ["type", "rubric_point", "page",
                     "target_word_or_sentence", "context_before",
                     "context_after", "correction", "comment"]

    missing = set(REQUIRED_KEYS) - set(annotation.keys())
    if missing:
        print(f"WARNING: Annotation[{idx}] missing fields: {missing}")
        return False

    # Validate page number
    if not isinstance(annotation["page"], int) or annotation["page"] < 1:
        print(f"WARNING: Invalid page number in annotation[{idx}]: {annotation['page']}")
        return False

    return True


def validate_input_paths(pdf_path: str, output_json_path: str, output_pdf_path: str) -> bool:
    """Validate all input/output paths before processing."""
    # Check PDF exists and is readable
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    try:
        with open(pdf_path, 'rb') as f:
            # Check file is not empty and starts with PDF header
            header = f.read(4)
            if header != b'%PDF':
                raise ValueError(f"File is not a valid PDF: {pdf_path}")
    except Exception as e:
        raise ValueError(f"Cannot read PDF {pdf_path}: {e}")

    # Check output paths are writable
    for path in [output_json_path, output_pdf_path]:
        try:
            dirname = os.path.dirname(path)
            if dirname and not os.path.exists(dirname):
                os.makedirs(dirname, exist_ok=True)
            # Test write access
            with open(path, 'w') as f:
                f.write("")
            os.remove(path)
        except Exception as e:
            raise ValueError(f"Cannot write to {path}: {e}")

    return True


def load_environment() -> Tuple[str, vision.ImageAnnotatorClient]:

    load_dotenv()
    grok_key = os.getenv("Grok_API")
    google_key = os.getenv("Google_cloud_key")
    missing = []
    if not grok_key:
        missing.append("Grok_API")
    if not google_key:
        missing.append("Google_cloud_key")
    if missing:
        raise EnvironmentError(
            f"Missing environment variable(s): {', '.join(missing)}. "
            "Please set them in your .env file."
        )

    # Validate API key formats
    if len(grok_key) < 20:
        raise ValueError(
            f"Invalid Grok_API key format: key is too short ({len(grok_key)} characters). "
            "Expected at least 20 characters."
        )
    if len(google_key) < 20:
        raise ValueError(
            f"Invalid Google_cloud_key format: key is too short ({len(google_key)} characters). "
            "Expected at least 20 characters."
        )

    client_options = ClientOptions(api_key=google_key)
    vision_client = vision.ImageAnnotatorClient(client_options=client_options)
    return grok_key, vision_client


# -----------------------------
# PDF → PAGE IMAGES (for Grok)
# -----------------------------


def pdf_to_page_images_for_grok(
    pdf_path: str,
    max_pages: int = 20,
    max_dim: int = 400,
    base64_cap: int = 12000,
    output_dir: str = "grok_images",
) -> List[Dict[str, Any]]:
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    doc = fitz.open(pdf_path)
    images = []
    for page in doc:
        # Use 200 DPI to match OCR processing (ensures coordinate alignment)
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        pil_img = Image.open(io.BytesIO(img_bytes))
        images.append(pil_img)

    page_images: List[Dict[str, Any]] = []

    for idx, img in enumerate(images):
        if idx >= max_pages:
            break
        resized = img.copy()
        # Reduced max dimension from 1200 to 800 (44% fewer pixels)
        resized.thumbnail((max_dim, max_dim))

        # Convert to RGB if necessary (JPEG doesn't support transparency)
        if resized.mode in ('RGBA', 'LA', 'P'):
            rgb_img = Image.new('RGB', resized.size, (255, 255, 255))
            if resized.mode == 'P':
                resized = resized.convert('RGBA')
            rgb_img.paste(resized, mask=resized.split()[-1] if resized.mode in ('RGBA', 'LA') else None)
            resized = rgb_img
        elif resized.mode != 'RGB':
            resized = resized.convert('RGB')

        buffer = io.BytesIO()
        # Changed from PNG to JPEG with 60% quality for reduced payload size
        resized.save(buffer, format="JPEG", quality=60, optimize=True)

        # Save image to disk
        file_path = os.path.join(output_dir, f"page_{idx + 1:03d}.jpg")
        resized.save(file_path, format="JPEG", quality=60, optimize=True)
        
        encoded = base64.b64encode(buffer.getvalue()).decode("utf-8")
        truncated = False
        if len(encoded) > base64_cap:
            encoded = encoded[:base64_cap]
            truncated = True
        page_images.append(
            {
                "page": idx + 1,
                "image_base64": encoded,
                "file_path": file_path,
                "truncated": truncated
            }
        )
    
    print(f"Saved {len(page_images)} page images to '{output_dir}/'")
    return page_images


def get_report_page_size(
    pdf_path: str,
    dpi: int = 200,
    margin_ratio: float = 0.40,
    min_height: int = 3500,
    fallback: Tuple[int, int] = (2977, 4211),
) -> Tuple[int, int]:
    """
    Match report page size to annotated answer pages:
    annotated width = orig_w + 2 * (margin_ratio * orig_w), height = orig_h.
    """
    doc = fitz.open(pdf_path)
    try:
        if doc.page_count == 0:
            return fallback
        pix = doc[0].get_pixmap(dpi=dpi)
        orig_w, orig_h = pix.width, pix.height
        margin = int(orig_w * margin_ratio)
        return (orig_w + 2 * margin, max(orig_h, min_height))
    except Exception:
        return fallback
    finally:
        doc.close()


# -----------------------------
# OCR WITH GOOGLE VISION
# -----------------------------


def _bbox_to_tuples(bbox) -> List[Tuple[int, int]]:
    return [(v.x, v.y) for v in bbox.vertices]


def _paragraph_text(paragraph) -> str:
    words = []
    for word in paragraph.words:
        symbols = "".join(symbol.text for symbol in word.symbols)
        words.append(symbols)
    return " ".join(words).strip()


def _is_noise_text(text: str, bbox: List[Tuple[int, int]], page_w: int, page_h: int) -> bool:
    """
    Filter out background noise from OCR results.
    Returns True if the text is likely noise.
    """
    if not text or not bbox:
        return True

    # Filter very short text (1-2 chars) that's likely noise
    if len(text.strip()) <= 2:
        return True

    # Calculate bbox dimensions
    xs = [p[0] for p in bbox]
    ys = [p[1] for p in bbox]
    if not xs or not ys:
        return True

    width = max(xs) - min(xs)
    height = max(ys) - min(ys)

    # Filter extremely small text (likely artifacts)
    if width < 10 or height < 10:
        return True

    # Filter text at extreme edges (often page numbers or noise)
    center_x = (min(xs) + max(xs)) / 2
    center_y = (min(ys) + max(ys)) / 2

    margin = 30  # pixels
    if center_x < margin or center_x > page_w - margin:
        if center_y < margin or center_y > page_h - margin:
            return True

    return False


def run_ocr_on_pdf(
    vision_client: vision.ImageAnnotatorClient, pdf_path: str
) -> Dict[str, Any]:
    doc = fitz.open(pdf_path)
    images = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        pil_img = Image.open(io.BytesIO(img_bytes))
        images.append(pil_img)

    pages_output: List[Dict[str, Any]] = []
    full_text_parts: List[str] = []

    for idx, img in enumerate(images):
        page_w, page_h = img.size
        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        vision_image = vision.Image(content=buffer.getvalue())
        response = vision_client.document_text_detection(image=vision_image)
        if response.error.message:
            raise RuntimeError(
                f"OCR failed on page {idx + 1}: {response.error.message}"
            )

        page_lines: List[Dict[str, Any]] = []
        annotation = response.full_text_annotation
        if annotation:
            full_text_parts.append(annotation.text.strip())
            for page in annotation.pages:
                for block in page.blocks:
                    for paragraph in block.paragraphs:
                        text = _paragraph_text(paragraph)
                        para_bbox = _bbox_to_tuples(paragraph.bounding_box)

                        # Filter noise
                        if _is_noise_text(text, para_bbox, page_w, page_h):
                            continue

                        word_entries: List[Dict[str, Any]] = []
                        for word in paragraph.words:
                            w_text = "".join(
                                symbol.text for symbol in word.symbols
                            ).strip()
                            if not w_text:
                                continue
                            word_bbox = _bbox_to_tuples(word.bounding_box)

                            # Filter noise words
                            if _is_noise_text(w_text, word_bbox, page_w, page_h):
                                continue

                            word_entries.append({
                                "text": w_text,
                                "bbox": word_bbox,
                            })

                        if word_entries:  # Only add paragraph if it has valid words
                            page_lines.append({
                                "text": text,
                                "bbox": para_bbox,
                                "words": word_entries,
                            })
        else:
            text_annotations = response.text_annotations
            if text_annotations:
                full_text_parts.append(text_annotations[0].description.strip())
                for ta in text_annotations[1:]:
                    ta_bbox = _bbox_to_tuples(ta.bounding_poly)
                    if not _is_noise_text(ta.description, ta_bbox, page_w, page_h):
                        page_lines.append({
                            "text": ta.description,
                            "bbox": ta_bbox,
                            "words": [],
                        })

        pages_output.append({"page_number": idx + 1, "lines": page_lines})

    return {"pages": pages_output, "full_text": "\n".join(full_text_parts).strip()}


# -----------------------------
# RUBRIC DOCX HELPERS
# -----------------------------


def _load_docx_text(path: str) -> str:
    try:
        doc = Document(path)
    except Exception as exc:
        return f"[Error reading DOCX at {path}: {exc}]"
    parts: List[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _normalize_subject_key(value: str) -> str:
    if not value:
        return ""
    key = value.lower().strip()
    key = re.sub(r"[\s_]+", "-", key)
    key = re.sub(r"[^a-z0-9-]", "", key)
    key = re.sub(r"-{2,}", "-", key)
    return key.strip("-")


def _subject_key_variants(value: str) -> Set[str]:

    base = _normalize_subject_key(value)
    variants: Set[str] = set()
    if base:
        variants.add(base)
        if base.endswith("-rubric"):
            trimmed = base[: -len("-rubric")].strip("-")
            if trimmed:
                variants.add(trimmed)
    return variants


def _subject_keys_match(candidate: str, target_variants: Set[str]) -> bool:
    candidate_variants = _subject_key_variants(candidate)
    return any(var in target_variants for var in candidate_variants)


def find_subject_rubric_path(subject: str) -> Optional[str]:

    base_dir = os.path.dirname(os.path.abspath(__file__))
    rubrics_root = os.path.join(base_dir, "20marks_Rubrics")
    if not os.path.isdir(rubrics_root):
        return None

    target_variants = _subject_key_variants(subject)
    if not target_variants:
        return None

    # Prefer matches where the directory name (subject folder) aligns with the
    # requested id; fall back to filename matches if needed.
    for entry in sorted(os.listdir(rubrics_root)):
        entry_path = os.path.join(rubrics_root, entry)
        if not os.path.isdir(entry_path):
            continue

        docx_files = sorted(
            f for f in os.listdir(entry_path) if f.lower().endswith(".docx")
        )
        if not docx_files:
            continue

        dir_matches = _subject_keys_match(entry, target_variants)
        for fname in docx_files:
            stem = os.path.splitext(fname)[0]
            if dir_matches or _subject_keys_match(stem, target_variants):
                return os.path.join(entry_path, fname)

    # Handle DOCX files that might live directly under the root.
    for fname in sorted(os.listdir(rubrics_root)):
        if not fname.lower().endswith(".docx"):
            continue
        stem = os.path.splitext(fname)[0]
        if _subject_keys_match(stem, target_variants):
            return os.path.join(rubrics_root, fname)

    return None


def load_subject_rubric_text(subject: str) -> Tuple[str, Optional[str]]:
    docx_path = find_subject_rubric_path(subject)
    if not docx_path:
        print(f"WARNING: No subject rubric DOCX found for '{subject}'.")
        return "", None
    text = _load_docx_text(docx_path)
    return text, docx_path


def load_refined_rubric_text() -> Tuple[str, Optional[str]]:

    base_dir = os.path.dirname(os.path.abspath(__file__))
    refined_path = os.path.join(base_dir, "REFINED RUBRIC.docx")
    if not os.path.isfile(refined_path):
        print("WARNING: Refined rubric DOCX not found.")
        return "", None
    text = _load_docx_text(refined_path)
    return text, refined_path


# -----------------------------
# GROK CALL 1: SECTION DETECTION
# -----------------------------

def call_grok_for_section_detection(
    grok_api_key: str,
    ocr_data: Dict[str, Any],
    page_images: List[Dict[str, Any]],
) -> Tuple[List[Dict[str, Any]], Dict[str, int]]:
    """
    Use Grok to detect headings, subheadings and content sections
    from the full OCR text + ALL page images.

    Returns a list of sections in the shape:
      {
        "title": str,
        "level": int,            # 1 = main heading, 2 = subheading
        "page_numbers": [int],
        "content": str,
        "line_indices": []       # kept for compatibility (not used)
      }

    IMPORTANT:
      - Grok is explicitly told to trust the PAGE IMAGES as the primary source
        and only use OCR as a helper.
    """
    system_msg = {
        "role": "system",
        "content": (
            "You are an expert at visually and logically segmenting handwritten exam answers.\n\n"
            "INPUT DATA YOU RECEIVE:\n"
            "- OCR text (approximate, may contain errors)\n"
            "- Per-page OCR text lines (text only, no bounding boxes or coordinates)\n"
            "- Base64-encoded page images of the handwritten script\n\n"
            "PRIMARY RULE:\n"
            "- The PAGE IMAGES are the primary source of truth\n"
            "- OCR text is only a helper for searching or clarifying words\n"
            "- If OCR and image disagree, ALWAYS trust the handwritten image\n\n"
            "YOUR TASK:\n"
            "Segment the answer into logical sections with this structure:\n"
            "Introduction → Main body sections (with optional subsections) → Conclusion\n\n"
            "STRICT REQUIREMENTS FOR SECTION DETECTION:\n\n"
            "1) INTRODUCTION (REQUIRED - MUST BE FIRST):\n"
            "   - If there is an explicit heading like 'Introduction', use that exact text as the title\n"
            "   - If no such heading exists, identify the first paragraph(s) that introduce the topic\n"
            "   - Set title to 'Introduction' for implicit introductions\n"
            "   - For 'exact_ocr_heading': copy the EXACT OCR text of the heading line (or first line if implicit)\n"
            "   - This MUST be the first section in your output\n\n"
            "2) BODY SECTIONS (IDENTIFY ALL MAJOR HEADINGS):\n"
            "   - Examine the page images carefully for ALL headings between Introduction and Conclusion\n"
            "   - Look for these VISUAL CUES in the images:\n"
            "     • Larger or bolder handwriting\n"
            "     • Underlined words or phrases\n"
            "     • Extra spacing above and/or below text\n"
            "     • Numbered headings: 1., 2., 3., or i., ii., iii., or (a), (b), (c)\n"
            "     • Short phrases at the start of a line that label the content below\n"
            "   - For each heading found:\n"
            "     • Create a section with that heading as the title\n"
            "     • Set 'level' = 1 for main topics, 'level' = 2 for subtopics under the previous main heading\n"
            "     • For 'exact_ocr_heading': copy the EXACT OCR text of that heading line (word-for-word, with any typos)\n"
            "     • Include ALL page numbers where this section's content appears\n"
            "   - DO NOT skip headings - find ALL of them\n"
            "   - DO NOT invent headings that don't exist visually\n\n"
            "3) CONCLUSION (REQUIRED - MUST BE LAST):\n"
            "   - If there is an explicit heading like 'Conclusion' or 'In conclusion', use that text\n"
            "   - If no such heading, identify the final paragraph(s) that summarize/wrap up the answer\n"
            "   - Set title to 'Conclusion' for implicit conclusions\n"
            "   - For 'exact_ocr_heading': copy the EXACT OCR text of the heading line (or last paragraph's first line if implicit)\n"
            "   - This MUST be the last section in your output\n\n"
            "SECTION DEFINITIONS:\n"
            "- A section is a continuous block of content belonging together under one heading/topic\n"
            "- Sections must appear in reading order (top to bottom, page by page)\n"
            "- Sections must NOT overlap - each line belongs to only one section\n"
            "- If you cannot confidently assign some lines, leave them out rather than forcing them\n\n"
            "CONTENT_TEXT REQUIREMENTS:\n"
            "- For each section, provide a concise summary of the student's actual content\n"
            "- Use your own words but stay faithful to what is written\n"
            "- Do NOT invent arguments or facts not present in the script\n"
            "- If handwriting is unclear, infer only what is reasonably supported\n\n"
            "HEADING QUALITY EVALUATION (NEW REQUIREMENT):\n"
            "For each heading/subheading (excluding 'Introduction' and 'Conclusion'), provide a 'comment' evaluating its quality:\n"
            "- Start with POSITIVE or NEGATIVE to clearly indicate the assessment\n"
            "- Evaluate based on these criteria:\n"
            "  • Is it self-explanatory? Does it immediately tell what point is being discussed?\n"
            "  • Is it directly relevant to the question and reflects exact themes from the question statement?\n"
            "  • Does it guide the reader by showing how this section contributes to answering the question?\n"
            "  • Or is it generic, vague, or acts as a decorative slogan rather than a meaningful signpost?\n"
            "- Examples:\n"
            "  • POSITIVE: This heading clearly identifies the specific aspect being analyzed and directly addresses a component of the question.\n"
            "  • NEGATIVE: This heading is too generic and doesn't clearly indicate what specific point will be discussed.\n"
            "- Keep comments concise (1-2 sentences)\n"
            "- For 'Introduction' and 'Conclusion' sections, you may set comment to empty string or a brief note\n\n"
            "OUTPUT FORMAT (CRITICAL):\n"
            "- Return ONLY valid JSON with structure: {\"sections\": [...]}\n"
            "- NO markdown formatting, NO code blocks, NO explanations\n"
            "- Each section object must have these EXACT fields:\n"
            "  • 'title': Clean, readable heading text\n"
            "  • 'exact_ocr_heading': EXACT OCR text with any typos/errors (used for precise location matching)\n"
            "  • 'level': integer (1 for main, 2 for sub)\n"
            "  • 'page_numbers': array of integers\n"
            "  • 'content_text': string summary\n"
            "  • 'comment': quality evaluation starting with POSITIVE or NEGATIVE (or empty for intro/conclusion)\n"
            "- NO extra fields, NO top-level keys besides 'sections'\n\n"
            "CONSISTENCY:\n"
            "- For the same input, produce consistent segmentation\n"
            "- Avoid random or arbitrary splits\n"
            "- Prefer fewer, well-justified sections over many uncertain micro-sections\n"
            "- Be thorough but not excessive\n"
        ),
    }


    # Sanitize OCR pages: remove any bounding-box or coordinate data before sending
    raw_pages = ocr_data.get("pages", [])
    sanitized_pages = []
    for p in raw_pages:
        lines = []
        for line in p.get("lines", []):
            # Keep only textual content (line text + list of word texts) — drop bbox info
            line_text = line.get("text", "")
            words = [w.get("text", "") for w in line.get("words", [])]
            lines.append({"text": line_text, "words": words})
        sanitized_pages.append({"page_number": p.get("page_number"), "lines": lines})

    user_payload = {
        "task": "Segment this handwritten exam answer into logical sections. Use the page images as primary source and OCR as helper.",
        "ocr_full_text": ocr_data.get("full_text", ""),
        "ocr_pages": sanitized_pages,
        "page_images_base64_png": page_images,
        "output_schema": {
            "sections": [
                {
                    "title": "string",
                    "exact_ocr_heading": "exact text from OCR",
                    "level": 1,
                    "page_numbers": [1],
                    "content_text": "string",
                    "comment": "POSITIVE/NEGATIVE: quality evaluation of heading",
                }
            ]
        },
    }

    user_msg = {
        "role": "user",
        "content": json.dumps(user_payload, ensure_ascii=False),
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {grok_api_key}",
    }

    payload = {
        "model": "grok-4-fast-reasoning",
        "messages": [system_msg, user_msg],
        "temperature": 0.1,
        "max_tokens": 4000,  # Sufficient for structure detection
    }

    max_retries = 3
    for attempt in range(max_retries):
        if attempt > 0:
            # Exponential backoff: 1s, 2s, 4s
            time.sleep(2 ** (attempt - 1))
            print(f"Retry attempt {attempt + 1}/{max_retries} for section detection...")

        try:
            resp = requests.post(
                "https://api.x.ai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=120,
            )

            if resp.status_code >= 300:
                if attempt < max_retries - 1:
                    print(f"API error {resp.status_code}, retrying...")
                    continue
                raise RuntimeError(f"Section detection API error {resp.status_code}: {resp.text}")

            data = resp.json()

            # Extract token usage from response
            usage = data.get("usage", {})
            token_usage = {
                "input_tokens": usage.get("prompt_tokens", 0),
                "output_tokens": usage.get("completion_tokens", 0),
            }

            # Print token usage for this Grok section-detection prompt
            try:
                total_tokens = int(token_usage.get("input_tokens", 0)) + int(
                    token_usage.get("output_tokens", 0)
                )
            except Exception:
                total_tokens = None
            print(
                f"[Grok Section Detection] tokens — input={token_usage['input_tokens']}, output={token_usage['output_tokens']}, total={total_tokens}"
            )

            # Check for truncation
            finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
            if finish_reason == "length" and attempt < max_retries - 1:
                payload["max_tokens"] = payload.get("max_tokens", 4000) + 2000
                print(f"Response truncated, increasing max_tokens to {payload['max_tokens']} and retrying...")
                continue

            content = data["choices"][0]["message"]["content"]
            cleaned = clean_json_from_llm(content)
            parsed = json.loads(cleaned)

            # Success
            if attempt > 0:
                print(f"Successfully parsed section detection on attempt {attempt + 1}")

            break  # Exit retry loop on success

        except (requests.Timeout, requests.ConnectionError) as e:
            if attempt < max_retries - 1:
                print(f"Network error, retrying...")
                continue
            raise RuntimeError(f"Section detection network error: {e}") from e

        except json.JSONDecodeError as e:
            if attempt < max_retries - 1:
                print(f"JSON parse error, retrying...")
                continue
            # Fallback on final attempt
            return [
                {
                    "title": "Section Detection Error",
                    "level": 1,
                    "page_numbers": [1],
                    "content_text": f"Grok section detection failed: {e}\nRaw: {content[:400]}",
                }
            ], token_usage

        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Unexpected error: {exc}, retrying...")
                continue
            raise RuntimeError(f"Unexpected section detection error: {exc}") from exc

    # If we get here, parsing succeeded

    raw_sections = parsed.get("sections", []) or []

    sections: List[Dict[str, Any]] = []
    for sec in raw_sections:
        title = (sec.get("title") or "UNSPECIFIED").strip()
        exact_ocr_heading = (sec.get("exact_ocr_heading") or title).strip()
        level = sec.get("level") or 1
        pages = sec.get("page_numbers") or []
        content_text = sec.get("content_text") or sec.get("content") or ""
        comment = (sec.get("comment") or "").strip()

        sections.append(
            {
                "title": title,
                "exact_ocr_heading": exact_ocr_heading,  # Store exact OCR text
                "level": int(level) if isinstance(level, (int, float)) else 1,
                "page_numbers": sorted(
                    set(int(p) for p in pages if isinstance(p, (int, float)))
                ),
                "content": content_text,
                "comment": comment,  # Quality evaluation of heading
                "line_indices": [],  # not used downstream, kept for compatibility
            }
        )

    return sections, token_usage


# -----------------------------
# GROK CALL 2: SUBJECT-WISE MARKING
# -----------------------------


def build_grok_payload_for_grading(
    subject: str,
    subject_rubric_text: str,
    ocr_data: Dict[str, Any],
    sections: List[Dict[str, Any]],
    page_images: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """
    Build payload for Grok subject-wise grading.

    We send:
      - subject,
      - subject rubric text (DOCX),
      - OCR full text,
      - sections structure,
      - page images.
    """
    schema_hint = {
        "subject": subject,
        "max_marks": 20,
        "total_marks_awarded": 0,
        "question_statement": "",
        "question_expectation": [
            "Bullet point 1: Key theme/concept expected",
            "Bullet point 2: Important theory or framework",
            "Bullet point 3: Historical context or period"
        ],
        "criteria": [
            {
                "id": "knowledge_accuracy",
                "name": "Knowledge & Accuracy",
                "max": 8,
                "awarded": 5,
                "strengths": ["..."],
                "weaknesses": ["..."],
            }
        ],
        "overall_comment": "",
    }

    instructions = (
        "You are an experienced strict CSS examiner. "
        "Using ONLY the provided subject-wise rubric text, you must grade the student's answer with STRICT marking. "
        "IMPORTANT: Maximum marks awarded should NOT exceed 14 out of 20. "
        "Average/acceptable answers should score LESS than 10 marks. "
        "Only exceptional answers should approach 14 marks. "
        "Derive criteria and marks from the rubric text. Return STRICT JSON.\n\n"
        "Required fields:\n"
        "  - subject\n"
        "  - max_marks: always 20\n"
        "  - total_marks_awarded (cap at 14 maximum)\n"
        "  - question_statement: the exam question as written by the student\n"
        "  - question_expectation: MUST be an array of 3-5 short, specific bullet points describing what an excellent answer should cover according to the subject rubric. Each bullet should be one clear sentence focusing on key themes, concepts, theories, or historical periods expected.\n"
        "  - criteria[]: each criterion with id, name, max, awarded, strengths[], weaknesses[]\n"
        "  - overall_comment: 3-5 sentence holistic evaluation.\n"
    )

    content_payload = {
        "subject": subject,
        "rubric_text": subject_rubric_text,
        # Send full OCR text (no truncation)
        "ocr_full_text": ocr_data.get("full_text", ""),
        "sections": sections,
        "page_images_base64_png": page_images,
        "output_schema": schema_hint,
    }

    return {
        "model": "grok-4-fast-reasoning",
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an expert CSS examiner. "
                    "You produce detailed, rubric-based marking reports and respond in JSON only. "
                    "IMPORTANT: Do NOT treat events from 2025 or later years as speculation. "
                    "If you encounter dates/events you don't have knowledge about, ignore them and focus on grading based on the rubric criteria. "
                    "Never comment on whether information is speculative based on your knowledge cutoff."
                ),
            },
            {
                "role": "user",
                "content": instructions
                + "\n\nDATA:\n"
                + json.dumps(content_payload, ensure_ascii=False),
            },
        ],
        "temperature": 0.15,
        "max_tokens": 8000,  # Increased to allow longer responses
    }


def call_grok_for_grading(
    grok_api_key: str,
    subject: str,
    subject_rubric_text: str,
    ocr_data: Dict[str, Any],
    sections: List[Dict[str, Any]],
    page_images: List[Dict[str, Any]],
    max_retries: int = 3,
) -> Tuple[Dict[str, Any], Dict[str, int]]:
    payload = build_grok_payload_for_grading(
        subject, subject_rubric_text, ocr_data, sections, page_images
    )
    headers = {
        "Authorization": f"Bearer {grok_api_key}",
        "Content-Type": "application/json",
    }

    for attempt in range(max_retries):
        if attempt > 0:
            print(f"Retry attempt {attempt + 1}/{max_retries} for Grok grading...")

        resp = requests.post(
            "https://api.x.ai/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=150,
        )
        if resp.status_code >= 300:
            if attempt < max_retries - 1:
                print(f"API error {resp.status_code}, retrying...")
                continue
            raise RuntimeError(f"Grok grading error {resp.status_code}: {resp.text}")

        try:
            data = resp.json()
        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Response JSON parse error, retrying...")
                continue
            raise RuntimeError(f"Grok grading JSON parse error: {exc}") from exc

        # Extract token usage from response
        usage = data.get("usage", {})
        token_usage = {
            "input_tokens": usage.get("prompt_tokens", 0),
            "output_tokens": usage.get("completion_tokens", 0),
        }

        try:
            content = data["choices"][0]["message"]["content"]
        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Missing content in response, retrying...")
                continue
            raise RuntimeError(f"Unexpected Grok grading response: {data}") from exc

        # Check if response was truncated due to finish_reason
        finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
        if finish_reason == "length":
            print(f"WARNING: Grok response was truncated due to max token limit!")
            print(f"Response length: {len(content)} characters")
            if attempt < max_retries - 1:
                # Increase max_tokens and retry
                payload["max_tokens"] = payload.get("max_tokens", 8000) + 2000
                print(f"Increasing max_tokens to {payload['max_tokens']} and retrying...")
                continue

        try:
            cleaned = clean_json_from_llm(content)
            parsed = json.loads(cleaned)
            # Success! Return the result
            if attempt > 0:
                print(f"Successfully parsed JSON on attempt {attempt + 1}")
            return parsed, token_usage
        except json.JSONDecodeError as exc:
            # Save malformed JSON to file for debugging (with unique timestamp)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            error_file = f"grok_error_response_{timestamp}_{attempt + 1}.txt"
            with open(error_file, "w", encoding="utf-8") as f:
                f.write(f"=== FULL RESPONSE (length: {len(content)} chars) ===\n")
                f.write(content)
                f.write(f"\n\n=== CLEANED (length: {len(cleaned)} chars) ===\n")
                f.write(cleaned)
                f.write(f"\n\n=== ERROR ===\n{exc}\n")

            print(f"\nDEBUG: Full content length: {len(content)} characters")
            print(f"DEBUG: Finish reason: {finish_reason}")
            print(f"DEBUG: JSON parse error at position {exc.pos}: {exc.msg}")
            print(f"DEBUG: Saved full response to {error_file}")
            print(f"DEBUG: First 300 chars: {content[:300]}")
            print(f"DEBUG: Last 300 chars: {content[-300:]}")

            if attempt < max_retries - 1:
                print(f"Malformed JSON, retrying...")
                continue

            raise RuntimeError(
                f"Grok grading returned malformed JSON after {max_retries} attempts. "
                f"Error: {exc.msg} at position {exc.pos}. "
                f"Full response saved to {error_file}"
            ) from exc

    raise RuntimeError(f"Failed to get valid response after {max_retries} attempts")


# -----------------------------
# GROK CALL 3: REFINED RUBRIC ANNOTATIONS
# -----------------------------


def call_grok_for_refined_rubric_annotations(
    grok_api_key: str,
    refined_rubric_text: str,
    ocr_data: Dict[str, Any],
    sections: List[Dict[str, Any]],
    page_images: List[Dict[str, Any]],
    max_retries: int = 3,
    output_json_path: str = "refined_rubric_result.json",
) -> Tuple[Dict[str, Any], Dict[str, int]]:
    """
    Ask Grok to apply the refined generic rubric and produce:
      - annotations[] (for drawing boxes / comments),
      - refined_rubric_summary[] (one entry per rubric point).

    Output shape:

      {
        "annotations": [
          { ... see annotate_pdf_with_rubric.py docstring ... }
        ],
        "refined_rubric_summary": [
          {
            "id": "introduction_quality",
            "name": "Introduction Quality",
            "rating": "weak/average/good/excellent",
            "comment": "..."
          },
          ...
        ]
      }
    """
    system_msg = {
        "role": "system",
        "content": (
            "You are an expert examiner and annotation assistant.\n"
            "You apply a refined generic rubric to handwritten exam answers and output:\n"
            "- concrete annotations (where to draw boxes and what feedback to write), and\n"
            "- a short summary for each rubric point.\n\n"
            "PRIMARY INPUTS YOU RECEIVE:\n"
            "- refined_rubric_text: the generic rubric you must apply;\n"
            "- ocr_full_text + ocr_pages: approximate OCR text and per-page structure;\n"
            "- sections: JSON from a previous step that segments the answer into logical sections;\n"
            "- page_images_base64_png: base64 PNG page images (the handwritten script).\n\n"
            "GROUND TRUTH PRIORITY:\n"
            "- The PAGE IMAGES are the ultimate source of truth.\n"
            "- OCR is only a helper for locating text; if OCR and image disagree, trust the image.\n"
            "- Never annotate OCR artifacts or OCR mistakes. If you suspect OCR is wrong, skip that annotation.\n"
            "- The SECTIONS JSON gives you stable titles and page ranges. Use it whenever you need\n"
            "  to refer to headings or sections (via section_id or target_section_id).\n\n"
            "YOUR GOAL:\n"
            "- Produce a JSON object with:\n"
            "    annotations[]: ALL important issues and observations you can detect,\n"
            "    refined_rubric_summary[]: one item per rubric point with rating + comment.\n"
            "- Use the refined rubric text for every judgment and comment; keep all evaluations consistent with it.\n"
            "- Be especially thorough for:\n"
            "    * incorrect or weak headings (missing clarity / relevance), and\n"
            "    * factual inaccuracies (wrong dates, wrong names, wrong causal claims, etc.).\n"
            "- It is acceptable to be conservative for spelling/grammar due to OCR noise.\n"
            "IMPORTANT: Do NOT treat events from 2025 or later years as speculation. "
            "If you encounter dates/events you don't have knowledge about, ignore them and focus on structural/heading/factual issues. "
            "Never comment on whether information is speculative based on your knowledge cutoff.\n\n"
            "STRICT OUTPUT FORMAT (IMPORTANT):\n"
            "- Return ONLY valid JSON (no markdown, no commentary).\n"
            "- Top-level keys allowed: 'annotations' and 'refined_rubric_summary'.\n"
            "- annotations[] entries can have additional fields, but MUST include the fields\n"
            "  required in the provided output_schema examples for each type.\n"
            "- The JSON must be parseable by a strict JSON parser.\n"
        ),
    }

    instructions = (
        "Use the refined generic rubric text to evaluate this answer strictly. "
        "You must obey these annotation rules:\n\n"
        "⚠️ ABSOLUTE REQUIREMENT - INTRODUCTION COMMENT:\n"
        "- THE VERY FIRST ANNOTATION IN YOUR OUTPUT MUST BE type='introduction_comment'\n"
        "- THIS IS NON-NEGOTIABLE - EVERY ANSWER MUST HAVE EXACTLY ONE INTRODUCTION COMMENT\n"
        "- PLACE IT AS THE FIRST ITEM IN THE annotations[] ARRAY\n"
        "- IF YOU DO NOT INCLUDE THIS, YOUR RESPONSE WILL BE REJECTED\n\n"
        "CRITICAL REQUIREMENT FOR ALL ANNOTATIONS:\n"
        "- For ALL annotation types, you MUST copy text EXACTLY as it appears in the OCR text.\n"
        "- NEVER paraphrase, reword, or correct the text when filling target_word_or_sentence.\n"
        "- ALWAYS provide context_before (3-5 words immediately before) and context_after (3-5 words immediately after).\n"
        "- Copy these contexts EXACTLY from the OCR text with original spelling and punctuation.\n\n"
        "- If the OCR text appears wrong compared to the page image, SKIP the annotation entirely.\n"
        "- Do NOT annotate errors that are caused by OCR mistakes.\n\n"
        "ALL ANNOTATIONS MUST HAVE THIS UNIFIED SCHEMA:\n"
        "  type: string (introduction_comment/heading_issue/factual_error/grammar_language/repetition)\n"
        "  rubric_point: string (e.g., 'introduction_quality', 'headings_subheadings', 'factual_accuracy', 'grammar_language')\n"
        "  page: integer (page number where the annotation appears)\n"
        "  target_word_or_sentence: string (EXACT text from OCR - the word, phrase, or sentence being annotated)\n"
        "  context_before: string (EXACT 3-5 words from OCR that appear immediately before the target)\n"
        "  context_after: string (EXACT 3-5 words from OCR that appear immediately after the target)\n"
        "  correction: string (the correct version, or suggestion for improvement)\n"
        "  comment: string (explanation of the issue)\n"
        "  sentiment: string (required for heading_issue: 'positive' or 'negative')\n\n"
        "1) Introduction:\n"
        "   - Use the refined rubric for all introduction judgments and comments.\n"
        "   - MANDATORY: You MUST ALWAYS create exactly ONE annotation of type 'introduction_comment'.\n"
        "   - This annotation is REQUIRED for every answer, regardless of quality.\n"
        "   - Decide if introduction is weak/average/good/excellent and be strict about it.\n"
        "   - Create type 'introduction_comment' with:\n"
        "       rubric_point = 'introduction_quality',\n"
        "       page = first page where introduction appears,\n"
        "       target_word_or_sentence = EXACT first sentence or opening phrase from OCR,\n"
        "       context_before = '' (empty for first sentence),\n"
        "       context_after = EXACT next 3-5 words from OCR after the target,\n"
        "       correction = '' (not applicable for introduction comments),\n"
        "       comment = a detailed 3–5 sentence evaluation of ONLY the introduction using the refined generic rubric.\n"
        "   - DO NOT SKIP THIS ANNOTATION. It is MANDATORY.\n\n"
        "2) Headings and subheadings:\n"
        "   - Use the refined rubric for all heading/subheading judgments and comments.\n"
        "   - MANDATORY: You MUST evaluate EVERY SINGLE heading and subheading detected in the sections[] array.\n"
        "   - For EACH heading/subheading found, you MUST create ONE 'heading_issue' annotation.\n"
        "   - DO NOT skip any headings. If you detect 5 headings, you MUST create 5 heading_issue annotations.\n"
        "   - IMPORTANT: DO NOT evaluate spelling, grammar, or OCR errors in headings. These are handled separately.\n"
        "   - ONLY evaluate heading CONTENT: relevance, clarity, self-explanatory nature.\n"
        "   - For CORRECT headings (self-explanatory, relevant, clear), add annotation type 'heading_issue' with:\n"
        "       rubric_point = 'headings_subheadings',\n"
        "       page = page number of that heading,\n"
        "       target_word_or_sentence = EXACT heading text from OCR,\n"
        "       context_before = EXACT 3-5 words from OCR before the heading,\n"
        "       context_after = EXACT 3-5 words from OCR after the heading,\n"
        "       sentiment = 'positive',\n"
        "       correction = '' (empty for positive headings),\n"
        "       comment = 'Correct heading' or 'Good heading - clear and relevant'.\n"
        "   - For INCORRECT/PROBLEMATIC headings, create type 'heading_issue' with:\n"
        "       rubric_point = 'headings_subheadings',\n"
        "       page = page number of that heading,\n"
        "       target_word_or_sentence = EXACT heading text from OCR (even if misspelled),\n"
        "       context_before = EXACT 3-5 words from OCR before the heading,\n"
        "       context_after = EXACT 3-5 words from OCR after the heading,\n"
        "       sentiment = 'negative',\n"
        "       correction = a better alternate heading that would be more self-explanatory and relevant,\n"
        "       comment = short explanation of the issue (NEVER mention spelling/grammar/OCR errors).\n"
        "   - Focus ONLY on: not being self-explanatory, not directly relevant, vague, unclear, irrelevant.\n"
        "   - IGNORE: spelling mistakes, grammar errors, OCR misreads in headings.\n\n"
        "3) Factual inaccuracies:\n"
        "   - Use the refined rubric for all factual judgments and comments.\n"
        "   - CRITICAL: Do NOT create annotations for CORRECT facts. ONLY annotate ACTUAL ERRORS.\n"
        "   - Do NOT mark spelling mistakes as factual errors.\n"
        "   - If a date/fact is correct, DO NOT create any annotation for it.\n"
        "   - Only flag actual factual mistakes (wrong dates, wrong facts, incorrect information).\n"
        "   - Keep target_word_or_sentence VERY SHORT (1-10 words max) containing only the error.\n"
        "   - For each ACTUAL factual mistake, create type 'factual_error' with:\n"
        "       rubric_point = 'factual_accuracy',\n"
        "       page = page where the error appears,\n"
        "       target_word_or_sentence = EXACT SHORT PHRASE containing the WRONG fact (e.g., '1944' when it should be '1945'),\n"
        "       context_before = EXACT 3-5 words immediately before the error in OCR,\n"
        "       context_after = EXACT 3-5 words immediately after the error in OCR,\n"
        "       correction = the CORRECT fact (e.g., '1945'),\n"
        "       comment = short explanation (e.g., 'Year should be 1945 not 1944').\n"
        "   - NEVER copy full sentences - only the specific phrase with the error.\n"
        "   - Examples:\n"
        "       * WRONG: target='1944', correction='1944', comment='correct' (DO NOT DO THIS!)\n"
        "       * WRONG: target='1707', correction='1707', comment='Year is correct' (DO NOT DO THIS!)\n"
        "       * CORRECT: target='1944', correction='1945', comment='Year should be 1945 not 1944'\n"
        "       * CORRECT: target='World War I', correction='World War II', comment='Should be WWII not WWI'\n\n"
        "4) Spelling only (no grammar):\n"
        "   - Use the refined rubric for all spelling judgments and comments.\n"
        "   - Focus ONLY on clear spelling mistakes (wrongly spelled words).\n"
        "   - Do NOT correct grammar, sentence structure, style, or phrasing.\n"
        "   - For each spelling issue, create type 'grammar_language' with:\n"
        "       rubric_point = 'grammar_language',\n"
        "       page = page where the misspelled word appears,\n"
        "       target_word_or_sentence = EXACT misspelled word or very short span (1-3 words) from OCR,\n"
        "       context_before = EXACT 3-5 words from OCR immediately before the misspelled word,\n"
        "       context_after = EXACT 3-5 words from OCR immediately after the misspelled word,\n"
        "       correction = the correctly spelled word or very short corrected phrase,\n"
        "       comment = brief note like 'spelling error'.\n"
        "   - Always cross-check using BOTH OCR text AND the page image:\n"
        "       * Use ocr_full_text to locate the word,\n"
        "       * Then visually verify the spelling directly on the page image.\n"
        "       * If OCR and the image disagree, TRUST THE IMAGE and do NOT flag a spelling error.\n"
        "   - Do not send entire paragraphs or long sentences as target_word_or_sentence.\n"
        "   - If the same misspelling occurs multiple times on the same page, create a SEPARATE annotation for EACH occurrence,\n"
        "     with different context_before and context_after for each instance.\n\n"
        "5) Repetition:\n"
        "   - Use the refined rubric for all repetitiveness judgments and comments.\n"
        "   - If content is repeated across pages, create type 'repetition' with:\n"
        "       rubric_point = 'repetitiveness',\n"
        "       page = the page where the repeated content appears again,\n"
        "       target_word_or_sentence = EXACT repeated phrase or sentence from OCR,\n"
        "       context_before = EXACT 3-5 words from OCR before the repeated text,\n"
        "       context_after = EXACT 3-5 words from OCR after the repeated text,\n"
        "       correction = suggestion like 'Remove repetition' or 'Already mentioned on page X',\n"
        "       comment = note indicating where it was first mentioned.\n\n"
        "Additionally, build refined_rubric_summary[]:\n"
        "   - Use the refined rubric for all summary ratings and comments.\n"
        "   - ONLY include these 4 rubric points:\n"
        "     1. argumentation_quality (name: 'Argumentation Quality')\n"
        "     2. presentation (name: 'Presentation Quality')\n"
        "     3. contemporary_relevance (name: 'Contemporary Relevance')\n"
        "     4. length_completeness (name: 'Length & Completeness')\n"
        "   - Each entry: id, name, rating (weak/average/good/excellent), comment (1 sentence max, very concise and brief).\n"
    )


    # Sanitize OCR pages (remove bounding boxes) and send full OCR text
    raw_pages = ocr_data.get("pages", [])
    sanitized_pages = []
    for p in raw_pages:
        lines = []
        for line in p.get("lines", []):
            line_text = line.get("text", "")
            words = [w.get("text", "") for w in line.get("words", [])]
            lines.append({"text": line_text, "words": words})
        sanitized_pages.append({"page_number": p.get("page_number"), "lines": lines})

    user_payload = {
        "refined_rubric_text": refined_rubric_text,
        "ocr_full_text": ocr_data.get("full_text", ""),
        "ocr_pages": sanitized_pages,
        "sections": sections,
        "page_images_base64_png": page_images,
        "output_schema": {
            "annotations": [
                {
                    "type": "introduction_comment",
                    "rubric_point": "introduction_quality",
                    "page": 1,
                    "target_word_or_sentence": "First sentence of introduction (EXACT from OCR)",
                    "context_before": "",
                    "context_after": "Next 3-5 words after (EXACT from OCR)",
                    "correction": "",
                    "comment": "Detailed evaluation of introduction quality (3-5 sentences)",
                },
                {
                    "type": "heading_issue",
                    "rubric_point": "headings_subheadings",
                    "page": 1,
                    "target_word_or_sentence": "EXACT heading text from OCR",
                    "context_before": "EXACT 3-5 words before",
                    "context_after": "EXACT 3-5 words after",
                    "correction": "Better heading suggestion or empty if positive",
                    "comment": "Explanation of issue or 'Correct heading'",
                    "sentiment": "positive/negative",
                },
                {
                    "type": "factual_error",
                    "rubric_point": "factual_accuracy",
                    "page": 2,
                    "target_word_or_sentence": "SHORT PHRASE with error (EXACT from OCR)",
                    "context_before": "EXACT 3-5 words before",
                    "context_after": "EXACT 3-5 words after",
                    "correction": "Correct fact",
                    "comment": "Explanation of error",
                }
            ],
            "refined_rubric_summary": [
                {
                    "id": "argumentation_quality",
                    "name": "Argumentation Quality",
                    "rating": "weak/average/good/excellent",
                    "comment": "string",
                },
                {
                    "id": "presentation",
                    "name": "Presentation Quality",
                    "rating": "weak/average/good/excellent",
                    "comment": "string",
                },
                {
                    "id": "contemporary_relevance",
                    "name": "Contemporary Relevance",
                    "rating": "weak/average/good/excellent",
                    "comment": "string",
                },
                {
                    "id": "length_completeness",
                    "name": "Length & Completeness",
                    "rating": "weak/average/good/excellent",
                    "comment": "string",
                }
            ],
        },
    }

    user_msg = {
        "role": "user",
        "content": instructions + "\n\nDATA:\n" + json.dumps(user_payload, ensure_ascii=False),
    }

    headers = {
        "Authorization": f"Bearer {grok_api_key}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": "grok-4-fast-reasoning",
        "messages": [system_msg, user_msg],
        "temperature": 0.1,
        "max_tokens": 6000,  # Increased for refined rubric annotations
    }

    def _save_refined_result(result: Dict[str, Any]) -> None:
        if not output_json_path:
            return
        try:
            with open(output_json_path, "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            print(f"Saved refined rubric JSON to {output_json_path}")
        except Exception as exc:
            print(f"WARNING: Failed to save refined rubric JSON to {output_json_path}: {exc}")

    for attempt in range(max_retries):
        if attempt > 0:
            print(f"Retry attempt {attempt + 1}/{max_retries} for refined rubric annotations...")

        resp = requests.post(
            "https://api.x.ai/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=150,
        )

        if resp.status_code >= 300:
            if attempt < max_retries - 1:
                print(f"API error {resp.status_code}, retrying...")
                continue
            # Fallback on final failure
            fallback = {
                "annotations": [
                    {
                        "type": "introduction_comment",
                        "rubric_point": "introduction_quality",
                        "page": 1,
                        "target_word_or_sentence": "",
                        "context_before": "",
                        "context_after": "",
                        "correction": "",
                        "comment": f"Error: API returned {resp.status_code}",
                    }
                ],
                "refined_rubric_summary": [],
            }
            _save_refined_result(fallback)
            return fallback, {"input_tokens": 0, "output_tokens": 0}

        try:
            data = resp.json()
        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Response JSON parse error, retrying...")
                continue
            # Fallback on final failure
            fallback = {
                "annotations": [
                    {
                        "type": "introduction_comment",
                        "rubric_point": "introduction_quality",
                        "page": 1,
                        "target_word_or_sentence": "",
                        "context_before": "",
                        "context_after": "",
                        "correction": "",
                        "comment": f"Error parsing API response: {exc}",
                    }
                ],
                "refined_rubric_summary": [],
            }
            _save_refined_result(fallback)
            return fallback, {"input_tokens": 0, "output_tokens": 0}

        # Extract token usage from response
        usage = data.get("usage", {})
        token_usage = {
            "input_tokens": usage.get("prompt_tokens", 0),
            "output_tokens": usage.get("completion_tokens", 0),
        }

        try:
            content = data["choices"][0]["message"]["content"]
        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Missing content in response, retrying...")
                continue
            # Fallback on final failure
            fallback = {
                "annotations": [
                    {
                        "type": "introduction_comment",
                        "rubric_point": "introduction_quality",
                        "page": 1,
                        "target_word_or_sentence": "",
                        "context_before": "",
                        "context_after": "",
                        "correction": "",
                        "comment": f"Error: Unexpected API response structure",
                    }
                ],
                "refined_rubric_summary": [],
            }
            _save_refined_result(fallback)
            return fallback, token_usage

        # Check if truncated
        finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
        if finish_reason == "length":
            print(f"WARNING: Refined rubric response truncated!")
            if attempt < max_retries - 1:
                payload["max_tokens"] = payload.get("max_tokens", 6000) + 2000
                print(f"Increasing max_tokens to {payload['max_tokens']} and retrying...")
                continue

        try:
            cleaned = clean_json_from_llm(content)
            parsed = json.loads(cleaned)
        except json.JSONDecodeError as exc:
            # Save error for debugging (with unique timestamp)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            error_file = f"grok_refined_error_{timestamp}_{attempt + 1}.txt"
            with open(error_file, "w", encoding="utf-8") as f:
                f.write(f"=== FULL RESPONSE (length: {len(content)} chars) ===\n")
                f.write(content)
                f.write(f"\n\n=== ERROR ===\n{exc}\n")

            print(f"\nDEBUG: JSON parse error at position {exc.pos}: {exc.msg}")
            print(f"DEBUG: Saved full response to {error_file}")

            if attempt < max_retries - 1:
                print(f"Malformed JSON, retrying...")
                continue

            # Fallback on final failure
            fallback = {
                "annotations": [
                    {
                        "type": "introduction_comment",
                        "rubric_point": "introduction_quality",
                        "page": 1,
                        "target_word_or_sentence": "",
                        "context_before": "",
                        "context_after": "",
                        "correction": "",
                        "comment": f"Error parsing refined rubric response: {exc}",
                    }
                ],
                "refined_rubric_summary": [],
            }
            _save_refined_result(fallback)
            return fallback, token_usage

        # Success! Normalize and validate
        parsed.setdefault("annotations", [])
        parsed.setdefault("refined_rubric_summary", [])

        # ENFORCE: Ensure introduction_comment always exists as first annotation
        annotations = parsed.get("annotations", [])
        has_intro = any(a.get("type") == "introduction_comment" for a in annotations)

        if not has_intro:
            # Inject introduction_comment as first annotation
            intro_text = ocr_data.get("full_text", "")[:200].split('\n')[0] if ocr_data.get("full_text") else "Introduction"
            intro_annotation = {
                "type": "introduction_comment",
                "rubric_point": "introduction_quality",
                "page": 1,
                "target_word_or_sentence": intro_text,
                "context_before": "",
                "context_after": "",
                "correction": "",
                "comment": "Introduction evaluation (auto-generated due to missing annotation)"
            }
            parsed["annotations"].insert(0, intro_annotation)
            print("⚠️  Warning: Introduction comment was missing. Auto-injected.")

        if attempt > 0:
            print(f"Successfully parsed refined rubric annotations on attempt {attempt + 1}")

        _save_refined_result(parsed)
        return parsed, token_usage

    # Should never reach here
    fallback = {
        "annotations": [
            {
                "type": "introduction_comment",
                "rubric_point": "introduction_quality",
                "page": 1,
                "target_word_or_sentence": "",
                "context_before": "",
                "context_after": "",
                "correction": "",
                "comment": f"Failed after {max_retries} attempts",
            }
        ],
        "refined_rubric_summary": [],
    }
    _save_refined_result(fallback)
    return fallback, {"input_tokens": 0, "output_tokens": 0}


# -----------------------------
# GROK CALL 4: PAGE-WISE IMPROVEMENT SUGGESTIONS
# -----------------------------


def call_grok_for_page_wise_suggestions(
    grok_api_key: str,
    subject: str,
    subject_rubric_text: str,
    ocr_data: Dict[str, Any],
    sections: List[Dict[str, Any]],) -> Tuple[Dict[str, Any], Dict[str, int]]:
    """
    Ask Grok to provide specific, actionable improvement suggestions for each page.

    Output shape:
      {
        "page_suggestions": [
          {
            "page": 1,
            "suggestions": [
              "Add comparison with Kant's categorical imperative",
              "Include the 1945 UN Charter establishment date",
              "Reference Rawls' theory of justice for stronger argumentation",
              "Add empirical evidence from 2020 Climate Summit"
            ]
          },
          ...
        ]
      }
    """
    system_msg = {
        "role": "system",
        "content": (
            "You are an expert CSS examiner focused on helping students improve their answers.\n"
            "You receive:\n"
            "- The subject rubric (detailed criteria)\n"
            "- OCR text from student's answer (page-wise)\n"
            "- Section structure (headings detected)\n\n"
            "YOUR GOAL:\n"
            "For each page, provide 3-6 specific, actionable suggestions for improvement.\n"
            "Focus on VALUE ADDITIONS that would strengthen the answer.\n"
            "IMPORTANT: Do NOT include grammar or spelling suggestions. These are handled separately.\n"
            "Focus only on content additions: theories, facts, evidence, comparisons, critical perspectives, and contemporary relevance.\n"
            "IMPORTANT: Do NOT treat events from 2025 or later years as speculation. "
            "If you encounter dates/events you don't have knowledge about, simply ignore them and focus on what the student can add to improve the answer. "
            "Never comment on whether content is speculative based on your knowledge cutoff.\n\n"
            "TYPES OF SUGGESTIONS TO PROVIDE:\n"
            "1. Theoretical additions: 'Add comparison with X philosopher/theorist'\n"
            "2. Factual additions: 'Include the date: [specific event] occurred in [year]'\n"
            "3. Evidence additions: 'Add empirical data from [specific study/report]'\n"
            "4. Comparative analysis: 'Compare with [country/era/policy]'\n"
            "5. Critical perspectives: 'Include critique from [scholar/school of thought]'\n"
            "6. Contemporary relevance: 'Link to recent event: [specific event in year]'\n\n"
            "STRICT OUTPUT FORMAT:\n"
            "- Return ONLY valid JSON\n"
            "- Top-level key: 'page_suggestions' with array of page objects\n"
            "- Each page object has: 'page' (integer) and 'suggestions' (array of strings)\n"
            "- Each suggestion must be a single, specific, actionable statement\n"
            "- No markdown, no commentary, just JSON\n"
        ),
    }

    instructions = (
        "Analyze this student's answer page by page.\n"
        "For each page, identify 3-6 specific additions that would improve the answer quality.\n\n"
        "REQUIREMENTS:\n"
        "1. Suggestions must be SPECIFIC, not vague\n"
        "   ❌ Bad: 'Add more theories'\n"
        "   ✅ Good: 'Add Foucault's concept of biopower (1976)'\n\n"
        "2. Focus on what to ADD, not what's wrong\n"
        "   ❌ Bad: 'This argument is weak'\n"
        "   ✅ Good: 'Strengthen argument by adding Weber's bureaucracy theory'\n\n"
        "3. Include specific names, dates, events, theories\n"
        "   ❌ Bad: 'Reference a philosopher'\n"
        "   ✅ Good: 'Reference Mill's harm principle (On Liberty, 1859)'\n\n"
        "4. Suggestions should align with the subject rubric criteria\n\n"
        "5. Each page should have 3-6 suggestions maximum\n\n"
        "OUTPUT:\n"
        "Return JSON with this exact structure:\n"
        "{\n"
        "  \"page_suggestions\": [\n"
        "    {\n"
        "      \"page\": 1,\n"
        "      \"suggestions\": [\n"
        "        \"Add comparison with Locke's social contract theory (Two Treatises, 1689)\",\n"
        "        \"Include the Magna Carta signing date (1215) as historical precedent\",\n"
        "        \"Reference the 1948 Universal Declaration of Human Rights\"\n"
        "      ]\n"
        "    },\n"
        "    {\n"
        "      \"page\": 2,\n"
        "      \"suggestions\": [\n"
        "        \"Add Amartya Sen's capability approach for economic analysis\",\n"
        "        \"Include World Bank poverty data (2021 report)\",\n"
        "        \"Compare with China's economic reforms post-1978\"\n"
        "      ]\n"
        "    }\n"
        "  ]\n"
        "}\n"
    )

    # Simplified page data (text only, no images or bounding boxes to reduce tokens)
    ocr_pages_minimal = [
        {
            "page": p.get("page", idx + 1),
            "text": p.get("text", "")[:300000]  # Limit per-page text
        }
        for idx, p in enumerate(ocr_data.get("pages", []))
    ]

    user_payload = {
        "subject": subject,
        "rubric_text": subject_rubric_text,
        "ocr_full_text": ocr_data.get("full_text", "")[:15000],
        "ocr_pages": ocr_pages_minimal,
        "sections": sections,
        "output_schema": {
            "page_suggestions": [
                {
                    "page": 1,
                    "suggestions": ["string", "string"]
                }
            ]
        },
    }

    user_msg = {
        "role": "user",
        "content": instructions + "\n\nDATA:\n" + json.dumps(user_payload, ensure_ascii=False),
    }

    headers = {
        "Authorization": f"Bearer {grok_api_key}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": "grok-4-fast-reasoning",
        "messages": [system_msg, user_msg],
        "temperature": 0.2,
        "max_tokens": 4000,  # Sufficient for page suggestions
    }

    max_retries = 3
    for attempt in range(max_retries):
        if attempt > 0:
            # Exponential backoff: 1s, 2s, 4s
            time.sleep(2 ** (attempt - 1))
            print(f"Retry attempt {attempt + 1}/{max_retries} for page suggestions...")

        try:
            resp = requests.post(
                "https://api.x.ai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=150,
            )

            if resp.status_code >= 300:
                if attempt < max_retries - 1:
                    print(f"API error {resp.status_code}, retrying...")
                    continue
                # Fallback on final attempt
                return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}

            data = resp.json()

            # Extract token usage
            usage = data.get("usage", {})
            token_usage = {
                "input_tokens": usage.get("prompt_tokens", 0),
                "output_tokens": usage.get("completion_tokens", 0),
            }

            # Check for truncation
            finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
            if finish_reason == "length" and attempt < max_retries - 1:
                payload["max_tokens"] = payload.get("max_tokens", 4000) + 2000
                print(f"Response truncated, increasing max_tokens to {payload['max_tokens']} and retrying...")
                continue

            content = data["choices"][0]["message"]["content"]
            cleaned = clean_json_from_llm(content)
            parsed = json.loads(cleaned)

            # Success
            if attempt > 0:
                print(f"Successfully parsed page suggestions on attempt {attempt + 1}")

            return parsed, token_usage

        except (requests.Timeout, requests.ConnectionError) as e:
            if attempt < max_retries - 1:
                print(f"Network error, retrying...")
                continue
            # Fallback on final attempt
            return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}

        except json.JSONDecodeError as e:
            if attempt < max_retries - 1:
                print(f"JSON parse error, retrying...")
                continue
            # Fallback on final attempt
            return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}

        except Exception as exc:
            if attempt < max_retries - 1:
                print(f"Unexpected error: {exc}, retrying...")
                continue
            # Fallback on final attempt
            return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}

    # Should never reach here, but fallback just in case
    return {"page_suggestions": []}, {"input_tokens": 0, "output_tokens": 0}


# -----------------------------
# REPORT RENDERING (SUBJECT MARKING)
# -----------------------------


_MODULE_DIR = os.path.dirname(os.path.abspath(__file__))
_FONTS_DIR = os.path.join(_MODULE_DIR, "fonts")
_FONT_CANDIDATES = [
    os.environ.get("OCR_FONT_PATH"),
    os.path.join(_FONTS_DIR, "ReportFont.ttf"),
    os.path.join(_FONTS_DIR, "NotoSans-Regular.ttf"),
    os.path.join(_FONTS_DIR, "DejaVuSans.ttf"),
    "arial.ttf",
    "Arial.ttf",
    "LiberationSans-Regular.ttf",
    "DejaVuSans.ttf",
]


def _iter_font_candidates() -> List[str]:
    seen: Set[str] = set()
    ordered: List[str] = []
    for candidate in _FONT_CANDIDATES:
        if not candidate:
            continue
        norm = os.path.normpath(candidate) if os.path.isabs(candidate) else candidate
        if norm in seen:
            continue
        seen.add(norm)
        ordered.append(candidate)
    return ordered


@lru_cache(maxsize=32)
def _get_font(size: int) -> ImageFont.FreeTypeFont:
    for candidate in _iter_font_candidates():
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _extract_expectation_bullets(expectation_text: Any) -> List[str]:
    bullet_points: List[str] = []
    if isinstance(expectation_text, list):
        bullet_points = [str(p).strip() for p in expectation_text if p]
    elif expectation_text:
        expectation_str = str(expectation_text).strip()
        sentences = re.split(r"(?<=[.!?])\s+", expectation_str)
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence:
                if sentence.endswith((".", "!", "?")):
                    sentence = sentence[:-1].strip()
                bullet_points.append(sentence)
        if not bullet_points:
            comma_parts = [p.strip() for p in expectation_str.split(",")]
            bullet_points = [p for p in comma_parts if p]
    if not bullet_points:
        bullet_points = ["No specific expectations provided"]
    bullet_points = [pt for pt in bullet_points if pt]
    return bullet_points[:4]


def _wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.FreeTypeFont,
    max_width: int,
) -> List[str]:
    words = text.split()
    lines: List[str] = []
    current = ""
    for w in words:
        trial = (current + " " + w).strip()
        if not trial:
            continue
        width = draw.textlength(trial, font=font)
        if width <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines

def render_subject_report_pages(
    grading_result: Dict[str, Any],
    page_size: Tuple[int, int] = (2977, 4211),  # 200 DPI: Width x Height
) -> List[Image.Image]:
    """
    Render the subject-wise marking report on exactly 2 pages max.
    If content exceeds 2 pages, font sizes are reduced until it fits.
      - SUBJECT NAME (at top)
      - TOTAL MARKS
      - QUESTION STATEMENT
      - WHAT QUESTION EXPECTS (right after question statement, as single paragraph)
      - CRITERIA breakdown (with strengths & weaknesses)
      (Note: high-scoring outline is on a separate dedicated page)
    """
    # Try rendering with progressively smaller fonts until it fits in 2 pages
    font_scale = 1.0
    max_attempts = 10

    for attempt in range(max_attempts):
        pages = _render_subject_report_with_scale(grading_result, page_size, font_scale)

        if len(pages) <= 2:
            # Success! Fits in 2 pages or less
            if attempt > 0:
                print(f"Subject report fit in {len(pages)} pages after reducing font to {font_scale:.1%} of original size")
            return pages

        # Too many pages, reduce font size
        print(f"Subject report has {len(pages)} pages (attempt {attempt+1}), reducing font size...")
        font_scale *= 0.85  # Reduce by 15% each iteration

    # If we still can't fit after max attempts, return what we have
    print(f"WARNING: Subject report still has {len(pages)} pages after {max_attempts} attempts to reduce font size")
    return pages


def _render_subject_report_with_scale(
    grading_result: Dict[str, Any],
    page_size: Tuple[int, int],
    font_scale: float = 1.0,
) -> List[Image.Image]:
    """
    Internal helper to render subject report with a given font scale.
    """
    W, H = page_size
    margin = int(W * 0.07)
    line_spacing = 1.4

    # Base font sizes (scaled by font_scale parameter)
    title_font = _get_font(int(90 * font_scale))
    subject_font = _get_font(int(74 * font_scale))
    section_heading_font = _get_font(int(78 * font_scale))
    question_text_font = _get_font(int(62 * font_scale))
    criteria_heading_font = _get_font(int(70 * font_scale))
    body_font = _get_font(int(58 * font_scale))
    total_heading_font = _get_font(int(106 * font_scale))
    total_marks_font = _get_font(int(140 * font_scale))

    pages: List[Image.Image] = []

    def new_page() -> Tuple[Image.Image, ImageDraw.ImageDraw, int]:
        img = Image.new("RGB", (W, H), "white")
        draw = ImageDraw.Draw(img)
        return img, draw, margin

    img, draw, y = new_page()

    def ensure_space(font_obj: ImageFont.FreeTypeFont, needed_lines: int):
        nonlocal img, draw, y
        line_h = font_obj.getbbox("Ag")[3] - font_obj.getbbox("Ag")[1]
        if y + line_h * needed_lines * line_spacing > H - margin:
            pages.append(img)
            img, draw, y = new_page()

    def draw_bold_text(
        text: str,
        font_obj: ImageFont.FreeTypeFont,
        position: Tuple[int, int],
        fill: str,
    ) -> None:
        offsets = [(0, 0), (1, 0), (0, 1), (1, 1)]
        for dx, dy in offsets:
            draw.text((position[0] + dx, position[1] + dy), text, font=font_obj, fill=fill)

    max_text_width = W - 2 * margin

    # SUBJECT NAME (at top)
    subject = grading_result.get("subject", "")
    if subject:
        ensure_space(subject_font, 2)
        draw.text((margin, y), f"Subject: {subject}", font=subject_font, fill="black")
        line_h_subj = subject_font.getbbox("Ag")[3] - subject_font.getbbox("Ag")[1]
        y += int(line_h_subj * line_spacing * 1.5)

    # TOTAL MARKS
    total = grading_result.get("total_marks_awarded", 0)
    maximum = grading_result.get("max_marks", 20)

    ensure_space(total_heading_font, 2)
    draw_bold_text("TOTAL MARKS", total_heading_font, (margin, y), "black")
    line_h = total_heading_font.getbbox("Ag")[3] - total_heading_font.getbbox("Ag")[1]
    y += int(line_h * line_spacing)
    draw_bold_text(f"{total} / {maximum}", total_marks_font, (margin, y), "#B22222")  # Keep red color
    line_h2 = total_marks_font.getbbox("Ag")[3] - total_marks_font.getbbox("Ag")[1]
    y += int(line_h2 * line_spacing * 2)

    # QUESTION STATEMENT
    question = grading_result.get("question_statement", "")
    ensure_space(section_heading_font, 3)
    # Center the heading
    heading_text = "QUESTION STATEMENT"
    heading_width = draw.textlength(heading_text, font=section_heading_font)
    heading_x = (W - heading_width) // 2
    draw_bold_text(heading_text, section_heading_font, (heading_x, y), "black")
    line_h_section = section_heading_font.getbbox("Ag")[3] - section_heading_font.getbbox("Ag")[1]
    y += int(line_h_section * 1.5)  # Increased spacing between heading and question text
    for line in _wrap_text(draw, question, question_text_font, max_text_width):
        ensure_space(question_text_font, 1)
        line_hq = question_text_font.getbbox("Ag")[3] - question_text_font.getbbox("Ag")[1]
        draw_bold_text(line, question_text_font, (margin, y), "black")  # Bold question text
        y += int(line_hq * line_spacing)
    y += int(question_text_font.getbbox("Ag")[3] - question_text_font.getbbox("Ag")[1])

    # WHAT QUESTION EXPECTS (right after question statement, as a single paragraph)
    expectation = grading_result.get("question_expectation", "")
    expectation_bullets = _extract_expectation_bullets(expectation) if expectation else []
    if expectation_bullets:
        ensure_space(section_heading_font, 2)
        # Center the heading
        heading_text = "WHAT THE QUESTION EXPECTS"
        heading_width = draw.textlength(heading_text, font=section_heading_font)
        heading_x = (W - heading_width) // 2
        draw_bold_text(heading_text, section_heading_font, (heading_x, y), "black")
        line_h_section = section_heading_font.getbbox("Ag")[3] - section_heading_font.getbbox("Ag")[1]
        y += int(line_h_section * line_spacing)
        line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
        for bullet in expectation_bullets:
            wrapped = _wrap_text(draw, bullet, body_font, max_text_width - int(0.05 * W))
            ensure_space(body_font, len(wrapped) + 1)
            for idx, line in enumerate(wrapped):
                if idx == 0:
                    draw.text((margin, y), f"- {line}", font=body_font, fill="black")
                else:
                    indent_x = margin + int(0.04 * W)
                    draw.text((indent_x, y), line, font=body_font, fill="black")
                y += int(line_hb * line_spacing)
            y += int(line_hb * 0.5)
        y += int(line_hb)

        # CRITERIA BREAKDOWN
    ensure_space(section_heading_font, 2)
    # Center the heading
    heading_text = "MARKS BREAKDOWN"
    heading_width = draw.textlength(heading_text, font=section_heading_font)
    heading_x = (W - heading_width) // 2
    draw_bold_text(heading_text, section_heading_font, (heading_x, y), "black")
    line_h_section = section_heading_font.getbbox("Ag")[3] - section_heading_font.getbbox("Ag")[1]
    y += int(line_h_section * line_spacing * 1.5)

    for crit in grading_result.get("criteria", []):
        name = crit.get("name", "")
        awarded = crit.get("awarded", 0)
        max_marks = crit.get("max", 0)
        header = f"{name}  —  {awarded}/{max_marks}"

        ensure_space(criteria_heading_font, 2)
        line_h_criteria = criteria_heading_font.getbbox("Ag")[3] - criteria_heading_font.getbbox("Ag")[1]
        draw_bold_text(header, criteria_heading_font, (margin, y), "black")
        y += int(line_h_criteria * line_spacing)

        # Strengths
        strengths = crit.get("strengths") or []
        if strengths:
            ensure_space(body_font, len(strengths) + 1)
            line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
            draw.text((margin, y), "Strengths:", font=body_font, fill="darkgreen")
            y += int(line_hb * line_spacing)
            for s in strengths:
                for line in _wrap_text(draw, f"• {s}", body_font, max_text_width):
                    ensure_space(body_font, 1)
                    draw.text((margin, y), line, font=body_font, fill="black")
                    y += int(line_hb * line_spacing)

        # Weaknesses
        weaknesses = crit.get("weaknesses") or []
        if weaknesses:
            ensure_space(body_font, len(weaknesses) + 1)
            line_hb = body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1]
            draw.text((margin, y), "Weaknesses:", font=body_font, fill="darkred")
            y += int(line_hb * line_spacing)
            for wtxt in weaknesses:
                for line in _wrap_text(draw, f"• {wtxt}", body_font, max_text_width):
                    ensure_space(body_font, 1)
                    draw.text((margin, y), line, font=body_font, fill="black")
                    y += int(line_hb * line_spacing)

        y += int(body_font.getbbox("Ag")[3] - body_font.getbbox("Ag")[1])

    pages.append(img)
    return pages




# -----------------------------------------
# PAGE: WHAT THE QUESTION EXPECTS (BULLETS)
# -----------------------------------------


# -----------------------------
# MAIN PIPELINE
# -----------------------------


def grade_pdf_answer(
    pdf_path: str,
    subject: str,
    output_json_path: str,
    output_pdf_path: str,
    user_id: Optional[str] = None,
) -> None:
    # Validate all inputs before processing
    print("Validating inputs...")
    validate_input_paths(pdf_path, output_json_path, output_pdf_path)

    # Validate subject name
    if not subject or len(subject.strip()) == 0:
        raise ValueError("Subject name cannot be empty")

    # Create unique output directory per request to prevent concurrent process conflicts
    unique_id = uuid.uuid4().hex[:8]
    output_dir = os.path.join(tempfile.gettempdir(), f"grok_images_{unique_id}")
    os.makedirs(output_dir, exist_ok=True)
    print(f"Created unique temp directory: {output_dir}")

    try:
        grok_key, vision_client = load_environment()

        print("Step 1: Converting PDF pages to images (for Grok)...")
        page_images = pdf_to_page_images_for_grok(pdf_path, output_dir=output_dir)

        print("Step 2: Running OCR on PDF (Google Vision)...")
        ocr_data = run_ocr_on_pdf(vision_client, pdf_path)

        print("Step 3: Detecting sections/headings with Grok...")
        sections, section_token_usage = call_grok_for_section_detection(
            grok_api_key=grok_key,
            ocr_data=ocr_data,
            page_images=page_images,
        )

        # Debug dump (only if DEBUG_SECTIONS environment variable is set)
        # Disabled by default in production to prevent file accumulation and OOM errors
        if os.getenv("DEBUG_SECTIONS", "").lower() in ("true", "1", "yes"):
            debug_dump_sections(sections, output_path="debug_sections.json")

        # Track total token usage
        total_input_tokens = section_token_usage.get("input_tokens", 0)
        total_output_tokens = section_token_usage.get("output_tokens", 0)

        print("Step 4: Loading subject-wise rubric DOCX...")
        subject_rubric_text, subject_rubric_path = load_subject_rubric_text(subject)
        if subject_rubric_path:
            print(f"Using subject rubric file: {subject_rubric_path}")
        else:
            print("Warning: No subject rubric file found; grading will be weaker.")

        print("Step 5: Calling Grok for subject-wise grading...")
        grading_result, grading_token_usage = call_grok_for_grading(
            grok_api_key=grok_key,
            subject=subject,
            subject_rubric_text=subject_rubric_text,
            ocr_data=ocr_data,
            sections=sections,
            page_images=page_images,
        )

        # Accumulate token usage
        total_input_tokens += grading_token_usage.get("input_tokens", 0)
        total_output_tokens += grading_token_usage.get("output_tokens", 0)

        grading_result.setdefault("subject", subject)
        if not grading_result.get("max_marks"):
            grading_result["max_marks"] = 20

        # Add token usage to grading result before saving
        grading_result["token_usage"] = {
            "input_tokens": total_input_tokens,
            "output_tokens": total_output_tokens,
            "total_tokens": total_input_tokens + total_output_tokens
        }
        print(f"OCR completed. Token usage: Input: {total_input_tokens}, Output: {total_output_tokens}")

        print("Saving grading JSON...")
        with open(output_json_path, "w", encoding="utf-8") as f:
            json.dump(grading_result, f, ensure_ascii=False, indent=2)
        print(f"Grading JSON saved to {output_json_path}")

        print("Step 6: Rendering subject-wise report pages...")
        report_page_size = get_report_page_size(pdf_path)
        subject_report_pages = render_subject_report_pages(
            grading_result,
            page_size=report_page_size,
        )

        print("Step 7: Loading refined rubric DOCX...")
        refined_rubric_text, refined_rubric_path = load_refined_rubric_text()
        if refined_rubric_path:
            print(f"Using refined rubric file: {refined_rubric_path}")
        else:
            print("Warning: No refined rubric file found; annotations will be weaker.")

        print("Step 8: Calling Grok for refined rubric annotations...")
        refined_result, refined_token_usage = call_grok_for_refined_rubric_annotations(
            grok_api_key=grok_key,
            refined_rubric_text=refined_rubric_text,
            ocr_data=ocr_data,
            sections=sections,
            page_images=page_images,
        )

        # Accumulate token usage
        total_input_tokens += refined_token_usage.get("input_tokens", 0)
        total_output_tokens += refined_token_usage.get("output_tokens", 0)

        annotations = refined_result.get("annotations", []) or []
        refined_summary = refined_result.get("refined_rubric_summary", []) or []

        # Validate refined summary schema
        if refined_summary:
            try:
                validate_refined_summary(refined_summary)
                print(f"Validated {len(refined_summary)} refined rubric summary items")
            except ValueError as e:
                print(f"WARNING: Refined summary validation failed: {e}")

        # Validate annotations schema
        valid_annotations = []
        for idx, ann in enumerate(annotations):
            if not validate_annotation(ann, idx):
                continue
            if (ann.get("type") or "").lower() == "factual_error":
                target = (ann.get("target_word_or_sentence") or "").strip()
                correction = (ann.get("correction") or "").strip()
                if not correction or target == correction:
                    continue
            valid_annotations.append(ann)
        if len(valid_annotations) < len(annotations):
            print(f"WARNING: {len(annotations) - len(valid_annotations)} annotations failed validation and were skipped")
        annotations = valid_annotations

        print("Step 9: Calling Grok for page-wise improvement suggestions...")
        page_suggestions_result, page_suggestions_token_usage = call_grok_for_page_wise_suggestions(
            grok_api_key=grok_key,
            subject=subject,
            subject_rubric_text=subject_rubric_text,
            ocr_data=ocr_data,
            sections=sections,
        )

        # Accumulate token usage
        total_input_tokens += page_suggestions_token_usage.get("input_tokens", 0)
        total_output_tokens += page_suggestions_token_usage.get("output_tokens", 0)

        page_suggestions = page_suggestions_result.get("page_suggestions", []) or []

        print("Step 10: Annotating answer pages with improvement suggestions...")
        annotated_answer_pages = annotate_pdf_answer_pages(
            pdf_path=pdf_path,
            ocr_data=ocr_data,
            sections=sections,
            annotations=annotations,
            page_suggestions=page_suggestions,
            refined_summary=refined_summary,
        )

        # Assemble final PDF
        all_pages: List[Image.Image] = []
        all_pages.extend(subject_report_pages)
        all_pages.extend(annotated_answer_pages)

        first = all_pages[0]
        rest = all_pages[1:]

        print(f"Step 11: Writing final PDF to {output_pdf_path} ...")
        first.save(
            output_pdf_path,
            "PDF",
            resolution=300.0,
            save_all=True,
            append_images=rest,
        )
        print("Done.")

    finally:
        # Clean up unique temp directory
        if os.path.exists(output_dir):
            try:
                shutil.rmtree(output_dir)
                print(f"Cleaned up temp directory: {output_dir}")
            except Exception as e:
                print(f"WARNING: Failed to remove temp directory {output_dir}: {e}")

# -----------------------------
# CLI
# -----------------------------


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Grade a handwritten PDF answer using Grok + Google Vision OCR, "
        "generate a subject-wise report, annotate the answer pages, "
        "and add a refined-rubric summary page."
    )
    parser.add_argument("--pdf_path", required=True, help="Path to the PDF file to grade.")
    parser.add_argument("--subject", required=True, help="Subject name, e.g., 'British History'.")
    parser.add_argument(
        "--output_json",
        default="grading_result.json",
        help="Path to write the grading JSON output.",
    )
    parser.add_argument(
        "--output_pdf",
        default="result.pdf",
        help="Path to write the final PDF (report + annotated pages + rubric summary).",
    )
    return parser.parse_args()


if __name__ == "__main__":
    try:
        args = parse_args()
        grade_pdf_answer(
            pdf_path=args.pdf_path,
            subject=args.subject,
            output_json_path=args.output_json,
            output_pdf_path=args.output_pdf,
        )
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        sys.exit(1)
